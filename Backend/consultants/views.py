from rest_framework import status
from rest_framework.decorators import api_view, parser_classes, renderer_classes
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from rest_framework.renderers import JSONRenderer
from rest_framework.response import Response
from django.shortcuts import get_object_or_404
from django.views.decorators.csrf import csrf_exempt
from django.core.files.storage import default_storage
from django.db import transaction
from django.db.models import Q
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from django.middleware.csrf import get_token
from django.views.decorators.csrf import ensure_csrf_cookie
from django.views.decorators.http import require_http_methods


from django.http import FileResponse
import mimetypes
from functools import lru_cache
import hashlib
import threading
from django.core.mail import send_mail
from django.conf import settings
from django.contrib.auth.tokens import default_token_generator
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.utils.encoding import force_bytes, force_str
from .models import Consultant, Competence, AppelOffre, User, CriteresEvaluation, MatchingResult, Notification
from .models import DocumentGED, DocumentCategory, DocumentVersion, DocumentAccess, Document 
from .serializers import ConsultantSerializer, CompetenceSerializer, AppelOffreSerializer, CriteresEvaluationSerializer
from .serializers import DocumentGEDSerializer, DocumentCategorySerializer
from .email_service import send_registration_email, send_validation_email
from .competences_data import ALL_SKILLS,FINANCE_BANKING_SKILLS, ENERGY_TRANSITION_SKILLS, \
    INDUSTRY_MINING_SKILLS
import fitz  # PyMuPDF
from PIL import Image, ImageEnhance
import pytesseract
import spacy
import os
import re
from dotenv import load_dotenv
from django.utils.timezone import now
from decimal import Decimal
from datetime import datetime
import logging

# Configurer le logging
logger = logging.getLogger(__name__)

# Chargement des variables d'environnement
load_dotenv()

# Chargement du mod√®le NLP
try:
    nlp = spacy.load("en_core_web_sm")
except Exception as e:
    logger.error(f"Erreur lors du chargement du mod√®le spaCy: {e}")
    nlp = None


@api_view(['POST'])
def cleanup_orphaned_users(request):
    """Nettoie les utilisateurs qui n'ont plus de consultants associ√©s"""
    try:
        # R√©cup√©rer tous les IDs des utilisateurs associ√©s √† des consultants
        consultant_user_ids = set(Consultant.objects.values_list('user_id', flat=True))

        # R√©cup√©rer tous les utilisateurs avec le r√¥le CONSULTANT qui ne sont pas associ√©s √† un consultant
        orphaned_users = User.objects.filter(role='CONSULTANT').exclude(id__in=consultant_user_ids)

        count = orphaned_users.count()
        orphaned_users.delete()

        return Response({"message": f"{count} utilisateurs orphelins supprim√©s avec succ√®s"})
    except Exception as e:
        return Response({"error": f"Erreur lors du nettoyage des utilisateurs: {str(e)}"}, status=500)
@ensure_csrf_cookie
@require_http_methods(["GET"])
@api_view(['GET'])
def get_csrf_token(request):
    """
    Endpoint pour r√©cup√©rer le token CSRF n√©cessaire pour les formulaires
    """
    try:
        csrf_token = get_token(request)
        logger.info("Token CSRF g√©n√©r√© avec succ√®s")
        
        return Response({
            'success': True,
            'csrf_token': csrf_token,
            'message': 'Token CSRF r√©cup√©r√© avec succ√®s'
        })
        
    except Exception as e:
        logger.error(f"Erreur lors de la g√©n√©ration du token CSRF: {str(e)}")
        return Response({
            'success': False,
            'error': 'Erreur lors de la g√©n√©ration du token CSRF'
        }, status=500)

@api_view(['GET'])
def api_public_consultants(request):
    """R√©cup√®re la liste des consultants disponibles publiquement"""
    consultants = Consultant.objects.all()
    serializer = ConsultantSerializer(consultants, many=True)
    return Response(serializer.data)



    """Liste ou cr√©e des consultants (acc√®s admin)"""
    try:
        # R√©cup√®re seulement les consultants valid√©s
        if request.method == 'GET':
            consultants = Consultant.objects.filter(is_validated=True)
            serializer = ConsultantSerializer(consultants, many=True)
            return Response({"success": True, "data": serializer.data})

        if request.method == 'POST':
            serializer = ConsultantSerializer(data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des consultants: {str(e)}")
        return Response({"success": False, "error": str(e)}, status=500)


# Remplacez la fonction admin_consultant_detail dans votre fichier views.py
@api_view(['GET'])
def consultant_matches(request, consultant_id):
    """
    Endpoint pour r√©cup√©rer les meilleurs matchings pour un consultant
    """
    limit = int(request.query_params.get('limit', 5))

    try:
        # V√©rifier que le consultant existe
        consultant = get_object_or_404(Consultant, id=consultant_id)
        
        # R√©cup√©rer les matchings pour ce consultant
        matchings = MatchingResult.objects.filter(
            consultant_id=consultant_id
        ).select_related('appel_offre').order_by('-score')[:limit]

        results = []
        for matching in matchings:
            appel_offre = matching.appel_offre
            
            results.append({
                'id': matching.id,
                'appel_offre_id': appel_offre.id,
                'appel_offre_name': getattr(appel_offre, 'titre', getattr(appel_offre, 'nom_projet', 'Titre non d√©fini')),
                'client': getattr(appel_offre, 'client', 'Client non sp√©cifi√©'),
                'date_debut': getattr(appel_offre, 'date_debut', getattr(appel_offre, 'date_de_publication', None)),
                'date_fin': getattr(appel_offre, 'date_fin', getattr(appel_offre, 'date_limite', None)),
                'score': float(matching.score),
                'is_validated': matching.is_validated
            })

        return Response({
            'success': True,
            'matches': results
        })

    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des matchings pour le consultant {consultant_id}: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)
@api_view(['GET'])
def admin_pending_consultants_fixed(request):
    """
    Version corrig√©e de la r√©cup√©ration des consultants en attente
    Compatible avec AdminLayout.tsx et AdminSidebar.tsx
    """
    try:
        logger.info("R√©cup√©ration des consultants en attente - version corrig√©e")
        
        # R√©cup√©rer les consultants en attente
        consultants = Consultant.objects.filter(is_validated=False)
        
        logger.info(f"Nombre de consultants en attente trouv√©s: {consultants.count()}")
        
        consultants_data = []
        for consultant in consultants:
            try:
                # Construction s√©curis√©e des donn√©es
                consultant_data = {
                    'id': consultant.id,
                    'nom': getattr(consultant, 'nom', ''),
                    'prenom': getattr(consultant, 'prenom', ''),
                    'firstName': getattr(consultant, 'prenom', ''),
                    'lastName': getattr(consultant, 'nom', ''),
                    'email': getattr(consultant, 'email', ''),
                    'telephone': getattr(consultant, 'telephone', ''),
                    'phone': getattr(consultant, 'telephone', ''),
                    'pays': getattr(consultant, 'pays', ''),
                    'country': getattr(consultant, 'pays', ''),
                    'ville': getattr(consultant, 'ville', ''),
                    'city': getattr(consultant, 'ville', ''),
                    
                    # Gestion s√©curis√©e des dates
                    'date_debut_dispo': consultant.date_debut_dispo.isoformat() if consultant.date_debut_dispo else None,
                    'date_fin_dispo': consultant.date_fin_dispo.isoformat() if consultant.date_fin_dispo else None,
                    'startAvailability': consultant.date_debut_dispo.isoformat() if consultant.date_debut_dispo else None,
                    'endAvailability': consultant.date_fin_dispo.isoformat() if consultant.date_fin_dispo else None,
                    
                    # Autres champs avec valeurs par d√©faut
                    'domaine_principal': getattr(consultant, 'domaine_principal', 'DIGITAL'),
                    'specialite': getattr(consultant, 'specialite', ''),
                    'expertise': getattr(consultant, 'expertise', 'D√©butant'),
                    'statut': getattr(consultant, 'statut', 'En_attente'),
                    'is_validated': getattr(consultant, 'is_validated', False),
                    
                    # Fichiers (gestion s√©curis√©e)
                    'cv': consultant.cv.url if getattr(consultant, 'cv', None) else None,
                    'photo': consultant.photo.url if getattr(consultant, 'photo', None) else None,
                    
                    # Comp√©tences
                    'skills': '',
                    
                    # M√©tadonn√©es
                    'created_at': consultant.created_at.isoformat() if consultant.created_at else None,
                    'updated_at': consultant.updated_at.isoformat() if consultant.updated_at else None,
                }
                
                # R√©cup√©rer les comp√©tences si possible
                try:
                    competences = Competence.objects.filter(consultant=consultant)
                    skills_list = [comp.nom_competence for comp in competences]
                    consultant_data['skills'] = ', '.join(skills_list)
                except Exception:
                    consultant_data['skills'] = ''
                
                consultants_data.append(consultant_data)
                
            except Exception as e:
                logger.error(f"Erreur traitement consultant {consultant.id}: {str(e)}")
                continue
        
        logger.info(f"Donn√©es finales pr√©par√©es: {len(consultants_data)} consultants")
        
        return Response({
            "success": True,
            "data": consultants_data,
            "count": len(consultants_data)
        })
        
    except Exception as e:
        logger.error(f"Erreur critique dans admin_pending_consultants_fixed: {str(e)}")
        return Response({
            "success": False,
            "error": str(e),
            "message": "Erreur lors de la r√©cup√©ration des consultants en attente"
        }, status=500)
        
# Ajoutez cette fonction corrig√©e dans views.py pour remplacer matching_for_offer_updated

@api_view(['GET', 'POST'])
def matching_for_offer_updated(request, appel_offre_id):
    """
    Endpoint pour g√©n√©rer ou r√©cup√©rer des matchings pour un appel d'offre scrap√©
    VERSION CORRIG√âE avec gestion d'erreur robuste
    """
    logger.info(f"Requ√™te {request.method} re√ßue pour appel d'offre scrap√© ID {appel_offre_id}")
    
    try:
        appel_offre_id = int(appel_offre_id)
    except ValueError:
        return Response({
            'success': False,
            'error': f"ID d'appel d'offre invalide: {appel_offre_id}"
        }, status=400)

    # V√©rifier que l'appel d'offre existe
    try:
        appel_offre = AppelOffre.objects.get(id=appel_offre_id)
        logger.info(f"Appel d'offre trouv√©: {appel_offre.titre}")
    except AppelOffre.DoesNotExist:
        logger.error(f"Appel d'offre avec ID {appel_offre_id} introuvable")
        return Response({
            'success': False,
            'error': f"Appel d'offre avec ID {appel_offre_id} introuvable"
        }, status=404)
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration de l'AO {appel_offre_id}: {str(e)}")
        return Response({
            'success': False,
            'error': f"Erreur lors de la r√©cup√©ration de l'appel d'offre: {str(e)}"
        }, status=500)

    if request.method == 'GET':
        try:
            # R√©cup√©rer tous les matchings pour cet appel d'offre
            matches = MatchingResult.objects.filter(
                appel_offre_id=appel_offre_id
            ).order_by('-score')

            logger.info(f"{matches.count()} matchings trouv√©s pour l'AO scrap√© {appel_offre_id}")

            if not matches.exists():
                return Response({
                    'success': True,
                    'matches': []
                })

            result = []
            for match in matches:
                consultant = match.consultant

                try:
                    # Recalculer les scores si n√©cessaire avec gestion d'erreur
                    try:
                        date_score = calculate_date_match_score_updated(consultant, match.appel_offre)
                        skills_score = calculate_skills_match_score_updated(consultant, match.appel_offre)
                        calculated_score = (date_score * 0.3) + (skills_score * 0.7)
                        stored_score = float(match.score)
                        
                        # V√©rifier la coh√©rence
                        if abs(calculated_score - stored_score) > 5:
                            match.score = Decimal(str(calculated_score))
                            match.save(update_fields=['score'])
                            stored_score = calculated_score
                            
                    except Exception as score_error:
                        logger.warning(f"Erreur lors du recalcul des scores: {str(score_error)}")
                        date_score = 0
                        skills_score = 0
                        stored_score = float(match.score)

                    stored_score = max(0, min(100, float(stored_score)))

                    # R√©cup√©rer les comp√©tences de mani√®re s√©curis√©e
                    try:
                        top_skills = get_top_skills_updated(consultant)
                    except Exception as skills_error:
                        logger.warning(f"Erreur r√©cup√©ration comp√©tences: {str(skills_error)}")
                        top_skills = []

                    result.append({
                        'id': match.id,
                        'consultant_id': consultant.id,
                        'consultant_name': f"{consultant.prenom or ''} {consultant.nom or ''}".strip(),
                        'consultant_expertise': consultant.expertise or "D√©butant",
                        'email': consultant.email or '',
                        'domaine_principal': consultant.domaine_principal or 'DIGITAL',
                        'specialite': consultant.specialite or "",
                        'top_skills': top_skills,
                        'date_match_score': round(date_score, 1),
                        'skills_match_score': round(skills_score, 1),
                        'score': stored_score,
                        'is_validated': match.is_validated
                    })

                except Exception as match_error:
                    logger.error(f"Erreur traitement match {match.id}: {str(match_error)}")
                    continue

            sorted_result = sorted(result, key=lambda x: x['score'], reverse=True)

            return Response({
                'success': True,
                'matches': sorted_result
            })
            
        except Exception as e:
            logger.error(f"Erreur lors de la r√©cup√©ration des matchings: {str(e)}")
            return Response({
                'success': False,
                'error': f"Erreur lors de la r√©cup√©ration des matchings: {str(e)}"
            }, status=500)

    elif request.method == 'POST':
        try:
            logger.info(f"G√©n√©ration de nouveaux matchings pour l'AO scrap√© {appel_offre_id}")
            
            # V√©rifier qu'il y a des consultants disponibles
            consultants = Consultant.objects.filter(
                is_validated=True,
                statut='Actif'
            ).exclude(
                date_debut_dispo=None
            ).exclude(
                date_fin_dispo=None
            )
            
            if not consultants.exists():
                logger.warning(f"Aucun consultant disponible pour le matching")
                return Response({
                    'success': False,
                    'error': "Aucun consultant disponible pour le matching"
                }, status=404)
            
            # Vider les anciens matchings
            deleted_count, _ = MatchingResult.objects.filter(appel_offre=appel_offre).delete()
            logger.info(f"{deleted_count} anciens matchings supprim√©s pour l'AO {appel_offre_id}")
            
            # Vider le cache
            try:
                clear_score_cache()
            except Exception as cache_error:
                logger.warning(f"Erreur nettoyage cache: {str(cache_error)}")
            
            results = []
            score_stats = {"min": 100, "max": 0, "total": 0, "count": 0}
            
            # Calculer les scores pour chaque consultant
            for consultant in consultants:
                try:
                    # Calculer les scores avec gestion d'erreur
                    try:
                        date_score = calculate_date_match_score_updated(consultant, appel_offre)
                    except Exception as date_error:
                        logger.warning(f"Erreur calcul date score pour consultant {consultant.id}: {str(date_error)}")
                        date_score = 0
                    
                    try:
                        skills_score = calculate_skills_match_score_updated(consultant, appel_offre)
                    except Exception as skills_error:
                        logger.warning(f"Erreur calcul skills score pour consultant {consultant.id}: {str(skills_error)}")
                        skills_score = 0
                    
                    # Pond√©ration: 30% date, 70% comp√©tences
                    final_score = (date_score * 0.3) + (skills_score * 0.7)
                    final_score = min(100, max(0, final_score))
                    
                    # Mettre √† jour les statistiques
                    if score_stats["count"] == 0:
                        score_stats["min"] = final_score
                        score_stats["max"] = final_score
                    else:
                        score_stats["min"] = min(score_stats["min"], final_score)
                        score_stats["max"] = max(score_stats["max"], final_score)
                    
                    score_stats["total"] += final_score
                    score_stats["count"] += 1
                    
                    logger.debug(f"Score final pour {consultant.nom}: {final_score:.2f}%")
                    
                    # Enregistrer le r√©sultat
                    try:
                        matching = MatchingResult.objects.create(
                            consultant=consultant,
                            appel_offre=appel_offre,
                            score=Decimal(str(final_score)),
                            is_validated=False
                        )
                        
                        # R√©cup√©rer les comp√©tences de mani√®re s√©curis√©e
                        try:
                            top_skills = get_top_skills_updated(consultant)
                        except Exception:
                            top_skills = []
                        
                        # Ajouter √† la liste de r√©sultats
                        results.append({
                            'id': matching.id,
                            'consultant_id': consultant.id,
                            'consultant_name': f"{consultant.prenom or ''} {consultant.nom or ''}".strip(),
                            'consultant_expertise': consultant.expertise or "D√©butant",
                            'email': consultant.email or '',
                            'domaine_principal': consultant.domaine_principal or 'DIGITAL',
                            'specialite': consultant.specialite or "",
                            'top_skills': top_skills,
                            'date_match_score': float(date_score),
                            'skills_match_score': float(skills_score),
                            'score': float(final_score),
                            'is_validated': False
                        })
                        
                    except Exception as save_error:
                        logger.error(f"Erreur sauvegarde matching pour consultant {consultant.id}: {str(save_error)}")
                        continue
                    
                except Exception as e:
                    logger.error(f"Erreur lors du calcul pour le consultant {consultant.id}: {str(e)}")
                    continue
            
            # Calculer les statistiques finales
            if results:
                score_stats["avg"] = score_stats["total"] / score_stats["count"]
                logger.info(f"Matchings g√©n√©r√©s: {len(results)}, score moyen: {score_stats['avg']:.2f}%")
            else:
                score_stats["avg"] = 0
                logger.warning("Aucun matching g√©n√©r√©")
            
            # Trier les r√©sultats par score d√©croissant
            sorted_results = sorted(results, key=lambda x: x['score'], reverse=True)
            
            return Response({
                'success': True,
                'matches': sorted_results,
                'stats': score_stats,
                'appel_offre_info': {
                    'id': appel_offre.id,
                    'titre': appel_offre.titre,
                    'client': appel_offre.client,
                    'has_description': bool(appel_offre.description),
                    'has_criteria': CriteresEvaluation.objects.filter(appel_offre=appel_offre).exists()
                }
            })
            
        except Exception as e:
            logger.error(f"Exception lors de la g√©n√©ration des matchings: {str(e)}")
            return Response({
                'success': False,
                'error': f"Erreur lors de la g√©n√©ration des matchings: {str(e)}"
            }, status=500)
    else:
        return Response({
            'success': False,
            'error': f"M√©thode {request.method} non support√©e"
        }, status=405)


# Fonction utilitaire pour nettoyer le cache
def clear_score_cache():
    """
    Nettoie le cache des scores de matching
    """
    global _score_cache
    try:
        _score_cache.clear()
        logger.info("Cache des scores nettoy√©")
    except Exception as e:
        logger.warning(f"Erreur lors du nettoyage du cache: {str(e)}")


# Fonction get_top_skills_updated s√©curis√©e
def get_top_skills_updated(consultant, limit=5):
    """
    R√©cup√®re les comp√©tences principales d'un consultant de mani√®re s√©curis√©e
    """
    try:
        top_skills = Competence.objects.filter(consultant=consultant).order_by('-niveau')[:limit]
        return [skill.nom_competence for skill in top_skills if skill.nom_competence]
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des comp√©tences: {str(e)}")
        return []
    
@api_view(['PUT', 'DELETE'])
@parser_classes([MultiPartParser, FormParser, JSONParser])
def admin_consultant_detail(request, pk):
    """
    G√®re un consultant sp√©cifique - ENDPOINT PRINCIPAL CORRIG√â
    Supprime le conflit entre les d√©corateurs
    """
    if request.method == 'DELETE':
        # CORRECTION: Appel direct de la logique de suppression dans la m√™me fonction
        return handle_consultant_deletion(request, pk)
        
    elif request.method == 'PUT':
        # Logique de mise √† jour existante
        try:
            consultant = get_object_or_404(Consultant, pk=pk)
            
            # Utiliser le serializer pour la mise √† jour
            serializer = ConsultantSerializer(consultant, data=request.data, partial=True)
            
            if serializer.is_valid():
                updated_consultant = serializer.save()
                
                # Mettre √† jour l'utilisateur associ√© si n√©cessaire
                if 'email' in request.data and consultant.user:
                    try:
                        consultant.user.email = request.data['email']
                        consultant.user.username = request.data['email']
                        consultant.user.save()
                        logger.info(f"Utilisateur mis √† jour pour consultant {pk}")
                    except Exception as user_error:
                        logger.error(f"Erreur mise √† jour utilisateur: {user_error}")
                
                # Pr√©parer la r√©ponse
                response_serializer = ConsultantSerializer(updated_consultant)
                logger.info(f"Consultant {pk} modifi√© avec succ√®s")
                
                return Response({
                    'success': True,
                    'data': response_serializer.data,
                    'message': f'Consultant {updated_consultant.prenom} {updated_consultant.nom} modifi√© avec succ√®s'
                }, status=200)
                
            else:
                logger.error(f"Erreurs de validation: {serializer.errors}")
                return Response({
                    'success': False,
                    'errors': serializer.errors,
                    'message': 'Erreurs de validation'
                }, status=400)
                
        except Exception as e:
            logger.error(f"Erreur mise √† jour consultant {pk}: {str(e)}")
            return Response({
                'success': False,
                'error': str(e)
            }, status=500)
    else:  # GET
        try:
            consultant = get_object_or_404(Consultant, pk=pk)
            serializer = ConsultantSerializer(consultant)
            return Response({
                'success': True,
                'data': serializer.data
            })
        except Exception as e:
            logger.error(f"Erreur r√©cup√©ration consultant {pk}: {str(e)}")
            return Response({
                'success': False,
                'error': str(e)
            }, status=500)


def handle_consultant_deletion(request, pk):
    """
    Logique de suppression des consultants - FONCTION INTERNE
    """
    try:
        consultant = get_object_or_404(Consultant, pk=pk)
        
        logger.info(f"=== D√âBUT SUPPRESSION CONSULTANT {pk} ===")
        logger.info(f"Consultant: {consultant.prenom} {consultant.nom}")
        
        # Stocker les informations avant suppression
        user = consultant.user
        consultant_name = f"{consultant.prenom} {consultant.nom}"
        consultant_id = consultant.id
        consultant_email = consultant.email
        
        # Utiliser une transaction pour garantir la coh√©rence
        with transaction.atomic():
            
            # 1. Supprimer les notifications li√©es
            try:
                notifications_deleted = 0
                try:
                    from .models import Notification
                    notifications = Notification.objects.filter(consultant=consultant)
                    notifications_deleted = notifications.count()
                    notifications.delete()
                    logger.info(f"‚úÖ {notifications_deleted} notifications supprim√©es")
                except ImportError:
                    logger.info("‚ÑπÔ∏è Mod√®le Notification non disponible")
            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Erreur suppression notifications: {e}")
                notifications_deleted = 0

            # 2. Supprimer les matchings li√©s
            try:
                matchings_deleted = 0
                try:
                    from .models import MatchingResult
                    matchings = MatchingResult.objects.filter(consultant=consultant)
                    matchings_deleted = matchings.count()
                    matchings.delete()
                    logger.info(f"‚úÖ {matchings_deleted} matchings supprim√©s")
                except ImportError:
                    logger.info("‚ÑπÔ∏è Mod√®le MatchingResult non disponible")
            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Erreur suppression matchings: {e}")
                matchings_deleted = 0

            # 3. Supprimer les comp√©tences
            try:
                competences = Competence.objects.filter(consultant=consultant)
                competences_deleted = competences.count()
                competences.delete()
                logger.info(f"‚úÖ {competences_deleted} comp√©tences supprim√©es")
            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Erreur suppression comp√©tences: {e}")
                competences_deleted = 0

            # 4. Mettre √† jour les documents GED (ne pas supprimer, juste d√©lier)
            try:
                doc_ged_updated = 0
                try:
                    from .models import DocumentGED
                    documents = DocumentGED.objects.filter(consultant=consultant)
                    doc_ged_updated = documents.update(consultant=None)
                    logger.info(f"‚úÖ {doc_ged_updated} documents GED d√©li√©s")
                except ImportError:
                    logger.info("‚ÑπÔ∏è Mod√®le DocumentGED non disponible")
            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Erreur mise √† jour documents GED: {e}")
                doc_ged_updated = 0

            # 5. Supprimer les missions li√©es
            try:
                missions_deleted = 0
                try:
                    from .models import Mission
                    missions = Mission.objects.filter(consultant=consultant)
                    missions_deleted = missions.count()
                    missions.delete()
                    logger.info(f"‚úÖ {missions_deleted} missions supprim√©es")
                except ImportError:
                    logger.info("‚ÑπÔ∏è Mod√®le Mission non disponible")
            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Erreur suppression missions: {e}")
                missions_deleted = 0

            # 6. Supprimer les fichiers physiques (CV et photo)
            files_deleted = []
            
            # Suppression du CV
            if consultant.cv and consultant.cv.name:
                try:
                    cv_path = consultant.cv.path
                    if os.path.exists(cv_path):
                        os.remove(cv_path)
                        files_deleted.append(f"CV: {cv_path}")
                        logger.info(f"‚úÖ Fichier CV supprim√©: {cv_path}")
                except Exception as e:
                    logger.warning(f"‚ö†Ô∏è Erreur suppression CV: {e}")
            
            # Suppression de la photo
            if consultant.photo and consultant.photo.name:
                try:
                    photo_path = consultant.photo.path
                    if os.path.exists(photo_path):
                        os.remove(photo_path)
                        files_deleted.append(f"Photo: {photo_path}")
                        logger.info(f"‚úÖ Fichier photo supprim√©: {photo_path}")
                except Exception as e:
                    logger.warning(f"‚ö†Ô∏è Erreur suppression photo: {e}")

            # 7. Supprimer les CV standardis√©s/Richat
            try:
                cv_standardises_deleted = 0
                cv_dir = os.path.join(settings.MEDIA_ROOT, 'standardized_cvs')
                if os.path.exists(cv_dir):
                    # Patterns de recherche pour les CV de ce consultant
                    patterns = [
                        f"standardized_cv_{consultant_id}_",
                        f"CV_Richat_{consultant.prenom}_{consultant.nom}",
                        f"CV_Richat_{consultant.nom}_{consultant.prenom}"
                    ]
                    
                    for filename in os.listdir(cv_dir):
                        if filename.endswith('.pdf'):
                            if any(pattern in filename for pattern in patterns if pattern):
                                try:
                                    file_path = os.path.join(cv_dir, filename)
                                    os.remove(file_path)
                                    cv_standardises_deleted += 1
                                    files_deleted.append(f"CV standardis√©: {filename}")
                                    logger.info(f"‚úÖ CV standardis√© supprim√©: {filename}")
                                except Exception as e:
                                    logger.warning(f"‚ö†Ô∏è Erreur suppression CV standardis√© {filename}: {e}")
                
                logger.info(f"‚úÖ {cv_standardises_deleted} CV standardis√©s supprim√©s")
            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Erreur suppression CV standardis√©s: {e}")
                cv_standardises_deleted = 0

            # 8. Supprimer le consultant de la base de donn√©es
            try:
                consultant.delete()
                logger.info(f"‚úÖ Consultant {consultant_id} supprim√© de la base de donn√©es")
            except Exception as e:
                logger.error(f"‚ùå Erreur critique suppression consultant: {e}")
                raise e

            # 9. Supprimer l'utilisateur associ√© (si existe et n'est li√© √† rien d'autre)
            user_deleted = False
            if user:
                try:
                    user_id = user.id
                    user_username = user.username
                    
                    # V√©rifier si l'utilisateur a d'autres relations
                    has_other_relations = False
                    
                    # V√©rifier s'il y a d'autres consultants avec ce user
                    other_consultants = Consultant.objects.filter(user=user).exists()
                    if other_consultants:
                        has_other_relations = True
                        logger.info(f"‚ÑπÔ∏è Utilisateur {user_id} a d'autres consultants, conservation")
                    
                    # Si pas d'autres relations, supprimer l'utilisateur
                    if not has_other_relations:
                        user.delete()
                        user_deleted = True
                        logger.info(f"‚úÖ Utilisateur {user_id} ({user_username}) supprim√©")
                    else:
                        logger.info(f"‚ÑπÔ∏è Utilisateur {user_id} conserv√© (autres relations)")
                        
                except Exception as e:
                    logger.error(f"‚ùå Erreur suppression utilisateur: {e}")
                    # Ne pas faire √©chouer toute la transaction pour l'utilisateur
                    user_deleted = False

        # 10. R√©sum√© des suppressions
        deletion_summary = {
            'consultant_id': consultant_id,
            'consultant_name': consultant_name,
            'consultant_email': consultant_email,
            'notifications_deleted': notifications_deleted,
            'matchings_deleted': matchings_deleted,
            'competences_deleted': competences_deleted,
            'missions_deleted': missions_deleted,
            'documents_ged_unlinked': doc_ged_updated,
            'files_deleted': files_deleted,
            'cv_standardises_deleted': cv_standardises_deleted,
            'user_deleted': user_deleted,
            'total_files_deleted': len(files_deleted)
        }

        logger.info(f"‚úÖ SUPPRESSION TERMIN√âE AVEC SUCC√àS")
        logger.info(f"üìä R√©sum√©: {deletion_summary}")

        return Response({
            'success': True,
            'message': f'Consultant "{consultant_name}" supprim√© avec succ√®s',
            'deletion_summary': deletion_summary
        }, status=200)

    except Exception as e:
        logger.error(f"‚ùå ERREUR CRITIQUE lors de la suppression du consultant {pk}: {str(e)}")
        return Response({
            'success': False,
            'error': f'Erreur lors de la suppression: {str(e)}',
            'consultant_id': pk
        }, status=500)


# FONCTION UTILITAIRE S√âPAR√âE POUR NETTOYER LES DONN√âES ORPHELINES
@api_view(['POST'])
def cleanup_orphaned_data(request):
    """
    Nettoie les donn√©es orphelines apr√®s suppressions
    """
    try:
        if not request.user.is_authenticated or not request.user.is_staff:
            return Response({
                'error': 'Permission refus√©e'
            }, status=403)

        cleanup_results = {
            'orphaned_users': 0,
            'orphaned_files': 0,
            'orphaned_matchings': 0,
            'errors': []
        }

        # 1. Nettoyer les utilisateurs orphelins
        try:
            # Utilisateurs CONSULTANT qui n'ont plus de consultant associ√©
            consultant_user_ids = set(Consultant.objects.values_list('user_id', flat=True))
            orphaned_users = User.objects.filter(
                role='CONSULTANT'
            ).exclude(id__in=consultant_user_ids)
            
            cleanup_results['orphaned_users'] = orphaned_users.count()
            orphaned_users.delete()
            
        except Exception as e:
            cleanup_results['errors'].append(f"Erreur nettoyage utilisateurs: {str(e)}")

        # 2. Nettoyer les matchings orphelins
        try:
            from .models import MatchingResult
            # Matchings dont le consultant n'existe plus
            existing_consultant_ids = set(Consultant.objects.values_list('id', flat=True))
            orphaned_matchings = MatchingResult.objects.exclude(
                consultant_id__in=existing_consultant_ids
            )
            
            cleanup_results['orphaned_matchings'] = orphaned_matchings.count()
            orphaned_matchings.delete()
            
        except Exception as e:
            cleanup_results['errors'].append(f"Erreur nettoyage matchings: {str(e)}")

        # 3. Nettoyer les fichiers orphelins
        try:
            cv_dir = os.path.join(settings.MEDIA_ROOT, 'standardized_cvs')
            if os.path.exists(cv_dir):
                existing_consultant_ids = list(Consultant.objects.values_list('id', flat=True))
                
                for filename in os.listdir(cv_dir):
                    if filename.endswith('.pdf'):
                        # Extraire l'ID du consultant du nom de fichier
                        import re
                        match = re.search(r'standardized_cv_(\d+)_', filename)
                        if match:
                            consultant_id = int(match.group(1))
                            if consultant_id not in existing_consultant_ids:
                                try:
                                    file_path = os.path.join(cv_dir, filename)
                                    os.remove(file_path)
                                    cleanup_results['orphaned_files'] += 1
                                except Exception as e:
                                    cleanup_results['errors'].append(f"Erreur suppression {filename}: {str(e)}")
                                    
        except Exception as e:
            cleanup_results['errors'].append(f"Erreur nettoyage fichiers: {str(e)}")

        return Response({
            'success': True,
            'message': 'Nettoyage termin√©',
            'results': cleanup_results
        })

    except Exception as e:
        logger.error(f"Erreur lors du nettoyage: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)


# FONCTION DE TEST POUR V√âRIFIER QUE TOUT FONCTIONNE
@api_view(['GET'])
def test_consultant_deletion_endpoint(request, pk):
    """
    Fonction de test pour v√©rifier l'√©tat d'un consultant avant suppression
    """
    try:
        consultant = get_object_or_404(Consultant, pk=pk)
        
        # Compter les relations
        competences_count = Competence.objects.filter(consultant=consultant).count()
        
        try:
            from .models import MatchingResult
            matchings_count = MatchingResult.objects.filter(consultant=consultant).count()
        except ImportError:
            matchings_count = 0
            
        try:
            from .models import Notification
            notifications_count = Notification.objects.filter(consultant=consultant).count()
        except ImportError:
            notifications_count = 0
        
        return Response({
            'consultant': {
                'id': consultant.id,
                'nom': f"{consultant.prenom} {consultant.nom}",
                'email': consultant.email,
                'user_id': consultant.user.id if consultant.user else None
            },
            'relations': {
                'competences': competences_count,
                'matchings': matchings_count,
                'notifications': notifications_count
            },
            'files': {
                'has_cv': bool(consultant.cv),
                'has_photo': bool(consultant.photo)
            }
        })
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=500)
    """
    G√®re un consultant sp√©cifique - ENDPOINT PRINCIPAL CORRIG√â
    Supprime le conflit entre les d√©corateurs
    """
    if request.method == 'DELETE':
        # CORRECTION: Appel direct de la logique de suppression dans la m√™me fonction
        return handle_consultant_deletion(request, pk)
        
    elif request.method == 'PUT':
        # Logique de mise √† jour existante
        try:
            consultant = get_object_or_404(Consultant, pk=pk)
            
            # Utiliser le serializer pour la mise √† jour
            serializer = ConsultantSerializer(consultant, data=request.data, partial=True)
            
            if serializer.is_valid():
                updated_consultant = serializer.save()
                
                # Mettre √† jour l'utilisateur associ√© si n√©cessaire
                if 'email' in request.data and consultant.user:
                    try:
                        consultant.user.email = request.data['email']
                        consultant.user.username = request.data['email']
                        consultant.user.save()
                        logger.info(f"Utilisateur mis √† jour pour consultant {pk}")
                    except Exception as user_error:
                        logger.error(f"Erreur mise √† jour utilisateur: {user_error}")
                
                # Pr√©parer la r√©ponse
                response_serializer = ConsultantSerializer(updated_consultant)
                logger.info(f"Consultant {pk} modifi√© avec succ√®s")
                
                return Response({
                    'success': True,
                    'data': response_serializer.data,
                    'message': f'Consultant {updated_consultant.prenom} {updated_consultant.nom} modifi√© avec succ√®s'
                }, status=200)
                
            else:
                logger.error(f"Erreurs de validation: {serializer.errors}")
                return Response({
                    'success': False,
                    'errors': serializer.errors,
                    'message': 'Erreurs de validation'
                }, status=400)
                
        except Exception as e:
            logger.error(f"Erreur mise √† jour consultant {pk}: {str(e)}")
            return Response({
                'success': False,
                'error': str(e)
            }, status=500)
    else:  # GET
        try:
            consultant = get_object_or_404(Consultant, pk=pk)
            serializer = ConsultantSerializer(consultant)
            return Response({
                'success': True,
                'data': serializer.data
            })
        except Exception as e:
            logger.error(f"Erreur r√©cup√©ration consultant {pk}: {str(e)}")
            return Response({
                'success': False,
                'error': str(e)
            }, status=500)
@api_view(['GET'])
def consultant_competences(request, consultant_id):
    """R√©cup√®re les comp√©tences d'un consultant sp√©cifique"""
    consultant = get_object_or_404(Consultant, id=consultant_id)
    competences = Competence.objects.filter(consultant=consultant)
    serializer = CompetenceSerializer(competences, many=True)
    return Response(serializer.data)


@api_view(['POST'])
def add_consultant_competence(request, consultant_id):
    """Ajoute une comp√©tence √† un consultant"""
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        nom_competence = request.data.get('nom_competence')
        niveau = request.data.get('niveau', 1)

        if not nom_competence:
            return Response({"error": "Nom de comp√©tence requis"}, status=400)

        # V√©rifier si la comp√©tence existe d√©j√†
        if Competence.objects.filter(consultant=consultant, nom_competence__iexact=nom_competence).exists():
            return Response({"error": "Cette comp√©tence existe d√©j√† pour ce consultant"}, status=400)

        # Cr√©er la comp√©tence
        competence = Competence.objects.create(
            consultant=consultant,
            nom_competence=nom_competence,
            niveau=niveau
        )

        # Mise √† jour du niveau d'expertise
        nb = Competence.objects.filter(consultant=consultant).count()
        consultant.expertise = "Expert" if nb >= 10 else "Interm√©diaire" if nb >= 5 else "D√©butant"
        consultant.save()

        return Response(CompetenceSerializer(competence).data, status=201)
    except Exception as e:
        return Response({"error": str(e)}, status=500)


@api_view(['DELETE'])
def delete_consultant_competence(request, consultant_id, competence_id):
    """Supprime une comp√©tence d'un consultant"""
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        competence = get_object_or_404(Competence, id=competence_id, consultant=consultant)

        competence.delete()

        # Mise √† jour du niveau d'expertise
        nb = Competence.objects.filter(consultant=consultant).count()
        consultant.expertise = "Expert" if nb >= 10 else "Interm√©diaire" if nb >= 5 else "D√©butant"
        consultant.save()

        return Response(status=204)
    except Exception as e:
        return Response({"error": str(e)}, status=500)


@api_view(['GET'])
def get_competences_by_domain(request, domain):
    """R√©cup√®re toutes les comp√©tences d'un domaine sp√©cifique"""
    if domain not in ALL_SKILLS:
        return Response({"error": "Domaine invalide"}, status=400)

    competences = ALL_SKILLS[domain]
    return Response({"domain": domain, "competences": competences})


@api_view(['GET'])
def get_all_domains(request):
    """R√©cup√®re la liste des domaines avec un √©chantillon de comp√©tences"""
    domains = []
    for domain, skills in ALL_SKILLS.items():
        # R√©cup√©rer les 5 premi√®res comp√©tences comme exemple
        sample_skills = skills[:5]
        domains.append({
            "code": domain,
            "name": dict(Consultant.SPECIALITES_CHOICES)[domain],
            "sample_skills": sample_skills
        })

    return Response({"domains": domains})


@api_view(['GET'])
def dashboard_stats(request):
    """
    R√©cup√®re les statistiques pour le tableau de bord admin - VERSION CORRIG√âE
    """
    try:
        from django.utils import timezone
        
        consultants_count = Consultant.objects.count()
        appels_total = AppelOffre.objects.count()
        
        # CORRECTION: Le nouveau mod√®le AppelOffre n'a pas de champ 'statut'
        # Utiliser les dates pour d√©terminer les offres actives
        now = timezone.now().date()
        
        # Offres actives = celles dont la date limite n'est pas d√©pass√©e
        offres_actives = AppelOffre.objects.filter(
            date_limite__gte=now
        ).count()
        
        # Offres expir√©es = celles dont la date limite est d√©pass√©e
        offres_expirees = AppelOffre.objects.filter(
            date_limite__lt=now
        ).count()
        
        # Offres sans date limite = consid√©r√©es comme actives
        offres_sans_date = AppelOffre.objects.filter(
            date_limite__isnull=True
        ).count()
        
        # Ajouter les offres sans date aux actives
        offres_actives += offres_sans_date

        # Derniers consultants
        derniers_consultants = Consultant.objects.order_by('-created_at')[:3]
        
        # Derniers appels d'offres (utiliser date_de_publication ou created_at)
        derniers_appels = AppelOffre.objects.order_by('-date_de_publication', '-created_at')[:3]

        data = {
            "consultants_count": consultants_count,
            "appels_total": appels_total,
            "offres_actives": offres_actives,
            "offres_expirees": offres_expirees,
            "derniers_consultants": [
                {
                    "nom": f"{c.prenom or ''} {c.nom or ''}".strip(),
                    "specialite": c.specialite or c.domaine_principal or "Non sp√©cifi√©",
                    "date": c.created_at.strftime('%d/%m/%Y') if c.created_at else "Non d√©fini"
                } for c in derniers_consultants
            ],
            "derniers_appels": [
                {
                    "title": a.titre or "Titre non d√©fini",
                    "client": a.client or "Client non sp√©cifi√©", 
                    "date": a.date_de_publication.strftime('%d/%m/%Y') if a.date_de_publication else a.created_at.strftime('%d/%m/%Y') if a.created_at else "Non d√©fini"
                } for a in derniers_appels
            ]
        }
        
        return Response(data)
        
    except Exception as e:
        logger.error(f"Erreur dans dashboard_stats: {str(e)}")
        return Response({
            "error": "Erreur lors de la r√©cup√©ration des statistiques",
            "detail": str(e)
        }, status=500)

@api_view(['GET'])
def admin_appels_offres(request):
    """
    Liste tous les appels d'offres - VERSION CORRIG√âE
    """
    try:
        logger.info("R√©cup√©ration des appels d'offres - D√âBUT")
        
        # R√©cup√©rer tous les appels d'offres
        appels = AppelOffre.objects.all().order_by('-date_de_publication', '-created_at')
        
        logger.info(f"Nombre d'appels trouv√©s: {appels.count()}")
        
        # S√©rialiser avec gestion des erreurs
        appels_data = []
        for appel in appels:
            try:
                appel_data = {
                    'id': appel.id,
                    'titre': appel.titre or '',
                    'date_de_publication': appel.date_de_publication.isoformat() if appel.date_de_publication else None,
                    'date_limite': appel.date_limite.isoformat() if appel.date_limite else None,
                    'client': appel.client or '',
                    'type_d_appel_d_offre': appel.type_d_appel_d_offre or '',
                    'description': appel.description or '',
                    'critere_evaluation': appel.critere_evaluation or '',
                    'documents': appel.documents or '',
                    'lien_site': appel.lien_site or '',
                    'created_at': appel.created_at.isoformat() if appel.created_at else None,
                    'updated_at': appel.updated_at.isoformat() if appel.updated_at else None,
                    # Propri√©t√©s calcul√©es
                    'is_expired': appel.is_expired,
                    'days_remaining': appel.days_remaining
                }
                appels_data.append(appel_data)
            except Exception as e:
                logger.error(f"Erreur s√©rialisation appel {appel.id}: {str(e)}")
                continue
        
        logger.info(f"Donn√©es s√©rialis√©es: {len(appels_data)} appels")
        return Response(appels_data)
        
    except Exception as e:
        logger.error(f"Erreur dans admin_appels_offres: {str(e)}")
        return Response({
            'error': 'Erreur lors de la r√©cup√©ration des appels d\'offres',
            'detail': str(e)
        }, status=500)
        
@api_view(['GET', 'PUT', 'DELETE'])
@parser_classes([JSONParser, MultiPartParser, FormParser])
def admin_appel_offre_detail(request, pk):
    """
    G√®re un appel d'offre sp√©cifique - ADAPT√â AU NOUVEAU MOD√àLE
    Compatible avec le dialog d'enrichissement d'AppelsOffres.tsx
    """
    try:
        appel = get_object_or_404(AppelOffre, pk=pk)

        if request.method == 'GET':
            # R√©cup√©rer les d√©tails complets
            appel_data = {
                'id': appel.id,
                'titre': appel.titre,
                'date_de_publication': appel.date_de_publication.isoformat() if appel.date_de_publication else None,
                'date_limite': appel.date_limite.isoformat() if appel.date_limite else None,
                'client': appel.client,
                'type_d_appel_d_offre': appel.type_d_appel_d_offre,
                'description': appel.description,
                'critere_evaluation': appel.critere_evaluation,
                'documents': appel.documents,
                'lien_site': appel.lien_site,
                'created_at': appel.created_at.isoformat(),
                'updated_at': appel.updated_at.isoformat(),
                'is_expired': appel.is_expired,
                'days_remaining': appel.days_remaining
            }
            
            # Informations d'enrichissement pour le frontend
            appel_data['enrichment_details'] = {
                'has_description': bool(appel.description and len(appel.description) > 50),
                'has_criteria': bool(appel.critere_evaluation and len(appel.critere_evaluation) > 20),
                'has_type': bool(appel.type_d_appel_d_offre),
                'structured_criteria': CriteresEvaluation.objects.filter(appel_offre=appel).count()
            }
            
            # Stats de matching
            matching_count = MatchingResult.objects.filter(appel_offre=appel).count()
            validated_matching_count = MatchingResult.objects.filter(appel_offre=appel, is_validated=True).count()
            
            appel_data['matching_stats'] = {
                'total_matches': matching_count,
                'validated_matches': validated_matching_count,
                'has_matches': matching_count > 0
            }
            
            return Response(appel_data)

        elif request.method == 'PUT':
            # Enrichissement des donn√©es - SEULS LES CHAMPS MODIFIABLES
            logger.info(f"Enrichissement de l'appel d'offres {pk}")
            
            # Champs que le frontend peut modifier (pas les donn√©es scrap√©es de base)
            modifiable_fields = ['description', 'critere_evaluation', 'type_d_appel_d_offre']
            
            # Extraire seulement les champs modifiables des donn√©es re√ßues
            update_data = {}
            for field in modifiable_fields:
                if field in request.data:
                    value = request.data[field]
                    if value is not None:  # Permet les cha√Ænes vides pour effacer un champ
                        update_data[field] = value
            
            if not update_data:
                return Response({
                    'error': 'Aucun champ modifiable fourni',
                    'modifiable_fields': modifiable_fields,
                    'received_fields': list(request.data.keys())
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Sauvegarder les modifications
            try:
                with transaction.atomic():
                    for field, value in update_data.items():
                        setattr(appel, field, value)
                    
                    appel.save(update_fields=list(update_data.keys()) + ['updated_at'])
                    
                    logger.info(f"Appel d'offres {pk} enrichi avec succ√®s: {list(update_data.keys())}")
                    
                    # Retourner les donn√©es mises √† jour
                    response_data = {
                        'id': appel.id,
                        'titre': appel.titre,
                        'date_de_publication': appel.date_de_publication.isoformat() if appel.date_de_publication else None,
                        'date_limite': appel.date_limite.isoformat() if appel.date_limite else None,
                        'client': appel.client,
                        'type_d_appel_d_offre': appel.type_d_appel_d_offre,
                        'description': appel.description,
                        'critere_evaluation': appel.critere_evaluation,
                        'documents': appel.documents,
                        'lien_site': appel.lien_site,
                        'created_at': appel.created_at.isoformat(),
                        'updated_at': appel.updated_at.isoformat(),
                        'is_expired': appel.is_expired,
                        'days_remaining': appel.days_remaining
                    }
                    
                    response_data['message'] = 'Appel d\'offres enrichi avec succ√®s'
                    response_data['updated_fields'] = list(update_data.keys())
                    
                    return Response(response_data)
                    
            except Exception as save_error:
                logger.error(f"Erreur sauvegarde AO {pk}: {str(save_error)}")
                return Response({
                    'error': f'Erreur lors de la sauvegarde: {str(save_error)}'
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        elif request.method == 'DELETE':
            # Suppression avec nettoyage des relations
            logger.warning(f"Suppression de l'appel d'offres {pk}")
            
            try:
                with transaction.atomic():
                    # Supprimer les relations d√©pendantes
                    matchings_deleted = MatchingResult.objects.filter(appel_offre=appel).delete()[0]
                    criteria_deleted = CriteresEvaluation.objects.filter(appel_offre=appel).delete()[0]
                    notifications_deleted = Notification.objects.filter(related_appel=appel).delete()[0]
                    
                    # Supprimer l'appel d'offres
                    appel_title = appel.titre
                    appel.delete()
                    
                    logger.info(f"Appel d'offres '{appel_title}' supprim√© avec ses relations")
                    
                    return Response({
                        'message': f'Appel d\'offres "{appel_title}" supprim√© avec succ√®s',
                        'deleted_relations': {
                            'matchings': matchings_deleted,
                            'criteria': criteria_deleted,
                            'notifications': notifications_deleted
                        }
                    }, status=status.HTTP_204_NO_CONTENT)
                    
            except Exception as delete_error:
                logger.error(f"Erreur suppression AO {pk}: {str(delete_error)}")
                return Response({
                    'error': f'Erreur lors de la suppression: {str(delete_error)}'
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
    except Exception as e:
        logger.error(f"Erreur dans admin_appel_offre_detail pour AO {pk}: {str(e)}")
        return Response({
            'error': f'Erreur lors de la gestion de l\'appel d\'offre: {str(e)}'
        }, status=500)


def extract_and_save_competences_async(file_path, consultant):
    """
    Fonction asynchrone pour extraire les comp√©tences d'un CV et les sauvegarder
    """
    try:
        logger.info(f"D√©but d'extraction des comp√©tences pour {consultant.nom} depuis {file_path}")

        # Extraction des comp√©tences
        competences, primary_domain = extract_competences_from_cv(file_path)
        logger.info(f"Comp√©tences trouv√©es pour {consultant.nom}: {competences}")
        logger.info(f"Domaine principal d√©tect√©: {primary_domain}")

        # Mettre √† jour le domaine principal si d√©tect√©
        if primary_domain:
            consultant.domaine_principal = primary_domain
            consultant.save()
            logger.info(f"Domaine principal mis √† jour pour {consultant.nom}: {primary_domain}")

        # Enregistrement des comp√©tences
        added_names = set()
        for nom in competences:
            nom_clean = nom.strip().strip(":,.\u2022¬∞").title()

            if nom_clean.lower() in added_names:
                continue

            if not Competence.objects.filter(nom_competence__iexact=nom_clean, consultant=consultant).exists():
                Competence.objects.create(
                    consultant=consultant,
                    nom_competence=nom_clean,
                    niveau=1
                )
                added_names.add(nom_clean.lower())
                logger.info(f"Comp√©tence ajout√©e pour {consultant.nom}: {nom_clean}")

        # Mise √† jour du niveau d'expertise
        nb = Competence.objects.filter(consultant=consultant).count()
        consultant.expertise = "Expert" if nb >= 10 else "Interm√©diaire" if nb >= 5 else "D√©butant"
        consultant.save()

        logger.info(f"Extraction termin√©e pour {consultant.nom} - {len(added_names)} comp√©tences ajout√©es")
    except Exception as e:
        logger.error(f"Erreur compl√®te lors de l'extraction pour {consultant.nom}: {e}")

# views.py - CORRECTION pour l'erreur des champs startAvailability/endAvailability

@api_view(['GET'])
def admin_pending_consultants_fixed(request):
    """
    Version corrig√©e de la r√©cup√©ration des consultants en attente
    """
    try:
        logger.info("R√©cup√©ration des consultants en attente - version s√©curis√©e")
        
        # R√©cup√©rer les consultants en attente
        consultants = Consultant.objects.filter(is_validated=False)
        
        logger.info(f"Nombre de consultants en attente trouv√©s: {consultants.count()}")
        
        consultants_data = []
        for consultant in consultants:
            try:
                # Construction s√©curis√©e des donn√©es
                consultant_data = {
                    'id': consultant.id,
                    'nom': safe_get_field_value(consultant, 'nom', default=''),
                    'prenom': safe_get_field_value(consultant, 'prenom', default=''),
                    'email': safe_get_field_value(consultant, 'email', default=''),
                    'telephone': safe_get_field_value(consultant, 'telephone', default=''),
                    'pays': safe_get_field_value(consultant, 'pays', default=''),
                    'ville': safe_get_field_value(consultant, 'ville', default=''),
                    
                    # Champs avec fallback vers les champs principaux
                    'firstName': safe_get_field_value(consultant, 'firstName', 'prenom', ''),
                    'lastName': safe_get_field_value(consultant, 'lastName', 'nom', ''),
                    'phone': safe_get_field_value(consultant, 'phone', 'telephone', ''),
                    'country': safe_get_field_value(consultant, 'country', 'pays', ''),
                    'city': safe_get_field_value(consultant, 'city', 'ville', ''),
                    
                    # Gestion s√©curis√©e des dates
                    'date_debut_dispo': consultant.date_debut_dispo.isoformat() if consultant.date_debut_dispo else None,
                    'date_fin_dispo': consultant.date_fin_dispo.isoformat() if consultant.date_fin_dispo else None,
                    'startAvailability': safe_get_field_value(
                        consultant, 'startAvailability', 'date_debut_dispo'
                    ),
                    'endAvailability': safe_get_field_value(
                        consultant, 'endAvailability', 'date_fin_dispo'
                    ),
                    
                    # Autres champs avec valeurs par d√©faut
                    'domaine_principal': safe_get_field_value(consultant, 'domaine_principal', default='DIGITAL'),
                    'specialite': safe_get_field_value(consultant, 'specialite', default=''),
                    'expertise': safe_get_field_value(consultant, 'expertise', default='D√©butant'),
                    'statut': safe_get_field_value(consultant, 'statut', default='En_attente'),
                    'is_validated': safe_get_field_value(consultant, 'is_validated', default=False),
                    
                    # Nouveaux champs d'expertise avec valeurs par d√©faut
                    'annees_experience': safe_get_field_value(consultant, 'annees_experience', default=0),
                    'formation_niveau': safe_get_field_value(consultant, 'formation_niveau', default='BAC+3'),
                    'certifications_count': safe_get_field_value(consultant, 'certifications_count', default=0),
                    'projets_realises': safe_get_field_value(consultant, 'projets_realises', default=0),
                    'leadership_experience': safe_get_field_value(consultant, 'leadership_experience', default=False),
                    'international_experience': safe_get_field_value(consultant, 'international_experience', default=False),
                    'expertise_score': safe_get_field_value(consultant, 'expertise_score', default=None),
                    
                    # Fichiers avec gestion d'erreur
                    'cv': consultant.cv.url if consultant.cv else None,
                    'photo': consultant.photo.url if consultant.photo else None,
                    'cvFilename': safe_get_field_value(consultant, 'cvFilename', default=None),
                    'standardizedCvFilename': safe_get_field_value(consultant, 'standardizedCvFilename', default=None),
                    'profileImage': safe_get_field_value(consultant, 'profileImage', default=None),
                    
                    # Gestion des comp√©tences
                    'skills': safe_get_field_value(consultant, 'skills', default=''),
                    
                    # M√©tadonn√©es
                    'created_at': consultant.created_at.isoformat() if consultant.created_at else None,
                    'updated_at': consultant.updated_at.isoformat() if consultant.updated_at else None,
                }
                
                # Convertir les dates pour les champs startAvailability et endAvailability
                if consultant_data['startAvailability'] and hasattr(consultant_data['startAvailability'], 'isoformat'):
                    consultant_data['startAvailability'] = consultant_data['startAvailability'].isoformat()
                elif consultant_data['date_debut_dispo']:
                    consultant_data['startAvailability'] = consultant_data['date_debut_dispo']
                
                if consultant_data['endAvailability'] and hasattr(consultant_data['endAvailability'], 'isoformat'):
                    consultant_data['endAvailability'] = consultant_data['endAvailability'].isoformat()
                elif consultant_data['date_fin_dispo']:
                    consultant_data['endAvailability'] = consultant_data['date_fin_dispo']
                
                # R√©cup√©rer les comp√©tences si pas encore dans skills
                if not consultant_data['skills']:
                    try:
                        competences = Competence.objects.filter(consultant=consultant)
                        skills_list = [comp.nom_competence for comp in competences]
                        consultant_data['skills'] = ', '.join(skills_list)
                    except Exception:
                        consultant_data['skills'] = ''
                
                consultants_data.append(consultant_data)
                
            except Exception as e:
                logger.error(f"Erreur traitement consultant {consultant.id}: {str(e)}")
                # Ajouter quand m√™me avec des donn√©es minimales
                consultants_data.append({
                    'id': consultant.id,
                    'nom': getattr(consultant, 'nom', 'Erreur'),
                    'prenom': getattr(consultant, 'prenom', 'Chargement'),
                    'email': getattr(consultant, 'email', ''),
                    'error': str(e)
                })
        
        logger.info(f"Donn√©es finales pr√©par√©es: {len(consultants_data)} consultants")
        
        return Response({
            "success": True,
            "data": consultants_data,
            "count": len(consultants_data),
            "note": "Version s√©curis√©e avec gestion des champs manquants"
        })
        
    except Exception as e:
        logger.error(f"Erreur critique dans admin_pending_consultants_fixed: {str(e)}")
        return Response({
            "success": False,
            "error": str(e),
            "message": "Erreur lors de la r√©cup√©ration des consultants en attente"
        }, status=500)


# Remplacement temporaire de la fonction actuelle

@api_view(['GET'])
def admin_consultants_corrected(request):
    """
    VERSION FINALE - Liste les consultants valid√©s avec gestion des champs
    """
    try:
        from django.db import connection
        
        # V√©rifier quels champs existent r√©ellement
        with connection.cursor() as cursor:
            cursor.execute("SHOW COLUMNS FROM consultants_consultant")
            existing_columns = {row[0] for row in cursor.fetchall()}
        
        consultants = Consultant.objects.filter(is_validated=True)
        consultants_data = []
        
        for consultant in consultants:
            try:
                consultant_data = {
                    'id': consultant.id,
                    'nom': consultant.nom or '',
                    'prenom': consultant.prenom or '',
                    'email': consultant.email or '',
                    'telephone': consultant.telephone or '',
                    'pays': consultant.pays or '',
                    'ville': consultant.ville or '',
                    'domaine_principal': consultant.domaine_principal or 'DIGITAL',
                    'specialite': consultant.specialite or '',
                    'expertise': consultant.expertise or 'D√©butant',
                    'statut': consultant.statut or 'En_attente',
                    'is_validated': consultant.is_validated,
                    
                    # Dates de base
                    'date_debut_dispo': consultant.date_debut_dispo.isoformat() if consultant.date_debut_dispo else None,
                    'date_fin_dispo': consultant.date_fin_dispo.isoformat() if consultant.date_fin_dispo else None,
                    
                    # Alias pour le frontend
                    'firstName': consultant.prenom or '',
                    'lastName': consultant.nom or '',
                    'phone': consultant.telephone or '',
                    'country': consultant.pays or '',
                    'city': consultant.ville or '',
                    'startAvailability': consultant.date_debut_dispo.isoformat() if consultant.date_debut_dispo else None,
                    'endAvailability': consultant.date_fin_dispo.isoformat() if consultant.date_fin_dispo else None,
                }
                
                # Ajouter les autres champs conditionnellement
                for field_name, default_value in [
                    ('annees_experience', 0),
                    ('formation_niveau', 'BAC+3'),
                    ('certifications_count', 0),
                    ('projets_realises', 0),
                    ('leadership_experience', False),
                    ('international_experience', False),
                    ('expertise_score', None),
                    ('cvFilename', None),
                    ('standardizedCvFilename', None),
                    ('profileImage', None),
                ]:
                    if field_name in existing_columns:
                        consultant_data[field_name] = getattr(consultant, field_name, default_value)
                    else:
                        consultant_data[field_name] = default_value
                
                # Fichiers
                try:
                    consultant_data['cv'] = consultant.cv.url if consultant.cv else None
                    consultant_data['photo'] = consultant.photo.url if consultant.photo else None
                except Exception:
                    consultant_data['cv'] = None
                    consultant_data['photo'] = None
                
                # Comp√©tences
                try:
                    competences = Competence.objects.filter(consultant=consultant)
                    skills_list = [comp.nom_competence for comp in competences]
                    consultant_data['skills'] = ', '.join(skills_list)
                except Exception:
                    consultant_data['skills'] = ''
                
                consultant_data['created_at'] = consultant.created_at.isoformat() if consultant.created_at else None
                consultant_data['updated_at'] = consultant.updated_at.isoformat() if consultant.updated_at else None
                
                consultants_data.append(consultant_data)
                
            except Exception as e:
                logger.error(f"Erreur traitement consultant {consultant.id}: {str(e)}")
                continue
        
        return Response({"success": True, "data": consultants_data})
            
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des consultants: {str(e)}")
        return Response({"success": False, "error": str(e)}, status=500)

    """
    SOLUTION D√âFINITIVE - R√©cup√®re les consultants en attente SANS utiliser les champs alias
    """
    try:
        logger.info("R√©cup√©ration des consultants en attente - version corrig√©e")
        
        # R√©cup√©rer SEULEMENT les consultants en attente SANS .only() pour √©viter les erreurs de champs
        consultants = Consultant.objects.filter(is_validated=False)
        
        logger.info(f"Nombre de consultants en attente trouv√©s: {consultants.count()}")
        
        consultants_data = []
        for consultant in consultants:
            try:
                # Construction des donn√©es SANS les champs probl√©matiques
                consultant_data = {
                    'id': consultant.id,
                    'nom': getattr(consultant, 'nom', '') or '',
                    'prenom': getattr(consultant, 'prenom', '') or '',
                    'firstName': getattr(consultant, 'prenom', '') or '',  # Utilise prenom, pas firstName
                    'lastName': getattr(consultant, 'nom', '') or '',     # Utilise nom, pas lastName
                    'email': getattr(consultant, 'email', '') or '',
                    'telephone': getattr(consultant, 'telephone', '') or '',
                    'phone': getattr(consultant, 'telephone', '') or '',  # Utilise telephone, pas phone
                    'pays': getattr(consultant, 'pays', '') or '',
                    'country': getattr(consultant, 'pays', '') or '',     # Utilise pays, pas country
                    'ville': getattr(consultant, 'ville', '') or '',
                    'city': getattr(consultant, 'ville', '') or '',       # Utilise ville, pas city
                    
                    # CORRECTION PRINCIPALE : Utiliser SEULEMENT les champs qui existent
                    'date_debut_dispo': consultant.date_debut_dispo.isoformat() if consultant.date_debut_dispo else None,
                    'date_fin_dispo': consultant.date_fin_dispo.isoformat() if consultant.date_fin_dispo else None,
                    # NE PAS utiliser startAvailability et endAvailability car ils n'existent pas en base
                    
                    'domaine_principal': getattr(consultant, 'domaine_principal', 'DIGITAL') or 'DIGITAL',
                    'specialite': getattr(consultant, 'specialite', '') or '',
                    'expertise': getattr(consultant, 'expertise', 'D√©butant') or 'D√©butant',
                    'statut': getattr(consultant, 'statut', 'En_attente') or 'En_attente',
                    'is_validated': getattr(consultant, 'is_validated', False),
                    
                    # Champs par d√©faut pour compatibilit√© frontend
                    'skills': '',
                    'annees_experience': 0,
                    'formation_niveau': 'BAC+3',
                    'certifications_count': 0,
                    'projets_realises': 0,
                    'leadership_experience': False,
                    'international_experience': False,
                    'expertise_score': None,
                    'cvFilename': None,
                    'standardizedCvFilename': None,
                    'profileImage': None,
                    
                    # Gestion s√©curis√©e des fichiers
                    'cv': None,
                    'photo': None,
                    
                    # M√©tadonn√©es
                    'created_at': consultant.created_at.isoformat() if consultant.created_at else None,
                    'updated_at': consultant.updated_at.isoformat() if consultant.updated_at else None,
                }
                
                # Gestion s√©curis√©e des fichiers
                try:
                    if hasattr(consultant, 'cv') and consultant.cv:
                        consultant_data['cv'] = consultant.cv.url
                        consultant_data['cvFilename'] = consultant.cv.name.split('/')[-1] if consultant.cv.name else None
                except Exception:
                    pass
                
                try:
                    if hasattr(consultant, 'photo') and consultant.photo:
                        consultant_data['photo'] = consultant.photo.url
                except Exception:
                    pass
                
                # R√©cup√©rer les comp√©tences si possible
                try:
                    competences = Competence.objects.filter(consultant=consultant)
                    skills_list = [comp.nom_competence for comp in competences]
                    consultant_data['skills'] = ', '.join(skills_list)
                except Exception:
                    consultant_data['skills'] = ''
                
                consultants_data.append(consultant_data)
                
            except Exception as e:
                logger.error(f"Erreur traitement consultant {consultant.id}: {str(e)}")
                # Ajouter quand m√™me avec des donn√©es minimales
                consultants_data.append({
                    'id': consultant.id,
                    'nom': getattr(consultant, 'nom', 'Erreur'),
                    'prenom': getattr(consultant, 'prenom', 'Chargement'),
                    'email': getattr(consultant, 'email', ''),
                    'error': str(e)
                })
        
        logger.info(f"Donn√©es finales pr√©par√©es: {len(consultants_data)} consultants")
        
        return Response({
            "success": True,
            "data": consultants_data,
            "count": len(consultants_data),
            "note": "Version corrig√©e sans champs alias"
        })
        
    except Exception as e:
        logger.error(f"Erreur r√©cup√©ration consultants en attente: {str(e)}")
        return Response({
            "success": False,
            "error": str(e),
            "note": "Erreur dans la r√©cup√©ration s√©curis√©e"
        }, status=500)



# FONCTION UTILITAIRE pour v√©rifier la structure de la base de donn√©es
@api_view(['GET'])
def debug_database_structure(request):
    """
    Fonction de d√©bogage pour v√©rifier quels champs existent r√©ellement dans la table
    """
    try:
        from django.db import connection
        
        with connection.cursor() as cursor:
            # Obtenir la structure de la table consultant
            cursor.execute("DESCRIBE consultants_consultant")
            fields_info = cursor.fetchall()
            
            existing_fields = []
            for field_info in fields_info:
                existing_fields.append({
                    'field': field_info[0],
                    'type': field_info[1],
                    'null': field_info[2],
                    'key': field_info[3],
                    'default': field_info[4],
                    'extra': field_info[5]
                })
        
        # Champs que le code essaie d'utiliser
        attempted_fields = [
            'startAvailability', 'endAvailability', 'firstName', 'lastName', 
            'phone', 'country', 'city', 'cvFilename', 'standardizedCvFilename',
            'profileImage', 'annees_experience', 'formation_niveau', 
            'certifications_count', 'projets_realises', 'leadership_experience',
            'international_experience', 'expertise_score'
        ]
        
        field_names = [field['field'] for field in existing_fields]
        missing_fields = [field for field in attempted_fields if field not in field_names]
        
        return Response({
            'existing_fields': existing_fields,
            'missing_fields': missing_fields,
            'attempted_fields': attempted_fields,
            'total_existing': len(existing_fields),
            'total_missing': len(missing_fields)
        })
        
    except Exception as e:
        return Response({'error': str(e)}, status=500)


# CORRECTION POUR LA FONCTION D'INSCRIPTION
@api_view(['POST'])
@parser_classes([MultiPartParser, FormParser])
def consultant_register_corrected(request):
    """
    Enregistre un nouveau consultant - VERSION CORRIG√âE SANS CHAMPS ALIAS
    """
    try:
        logger.info("Donn√©es re√ßues pour l'inscription:", request.data)
        email = request.data.get('email')
        nom = request.data.get('nom') or 'Consultant'
        prenom = request.data.get('prenom') or ''
        password = request.data.get('password') or 'consultant123'
        domaine_principal = request.data.get('domaine_principal') or 'DIGITAL'

        # V√©rifier si un utilisateur avec cet email existe
        existing_user = User.objects.filter(username=email).first()

        if existing_user:
            has_consultant = Consultant.objects.filter(user=existing_user).exists()
            if has_consultant:
                return Response({"error": "Utilisateur existe d√©j√†."}, status=400)
            else:
                logger.info(f"Suppression de l'utilisateur orphelin: {existing_user.username}")
                existing_user.delete()

        with transaction.atomic():
            # Cr√©ation de l'utilisateur
            user = User.objects.create_user(
                username=email,
                email=email,
                password=password,
                nom=nom,
                role='CONSULTANT'
            )
            logger.info(f"Utilisateur cr√©√©: {user.username} (ID: {user.id})")

            # Pr√©paration des donn√©es consultant SANS champs alias
            data = {
                'user': user.id,
                'prenom': prenom,
                'nom': nom,
                'email': email,
                'telephone': request.data.get('telephone', ''),
                'pays': request.data.get('pays', ''),
                'ville': request.data.get('ville', ''),
                'date_debut_dispo': request.data.get('date_debut_dispo'),
                'date_fin_dispo': request.data.get('date_fin_dispo'),
                'domaine_principal': domaine_principal,
                'specialite': request.data.get('specialite', ''),
                'expertise': 'D√©butant',
                'statut': 'En_attente',
                'is_validated': False
            }

            # Traitement des fichiers
            if 'cv' in request.FILES:
                data['cv'] = request.FILES['cv']
            
            if 'photo' in request.FILES:
                data['photo'] = request.FILES['photo']

            # Cr√©ation du consultant
            serializer = ConsultantSerializer(data=data)
            if serializer.is_valid():
                consultant = serializer.save()

                logger.info(f"Consultant cr√©√©: {consultant.nom} {consultant.prenom} (ID: {consultant.id})")

                # Gestion des comp√©tences manuelles
                competences_str = request.data.get('competences')
                if competences_str:
                    competences = [c.strip() for c in competences_str.split(',') if c.strip()]
                    for comp in competences:
                        if not Competence.objects.filter(consultant=consultant, nom_competence__iexact=comp).exists():
                            Competence.objects.create(
                                consultant=consultant,
                                nom_competence=comp,
                                niveau=2
                            )
                    logger.info(f"{len(competences)} comp√©tences manuelles ajout√©es")

                # Envoi d'email de confirmation
                try:
                    email_sent = send_registration_email(consultant)
                    if email_sent:
                        logger.info(f"Email de confirmation envoy√© avec succ√®s √† {consultant.email}")
                    else:
                        logger.warning(f"√âchec de l'envoi d'email √† {consultant.email}")
                except Exception as e:
                    logger.error(f"Erreur lors de l'envoi de l'email de confirmation: {str(e)}")

                # Retour r√©ussi
                logger.info(f"Inscription r√©ussie pour {consultant.nom} (ID: {consultant.id})")
                return Response({
                    "success": True,
                    "message": "Consultant cr√©√© avec succ√®s. Votre compte est en attente de validation par un administrateur.",
                    "consultant_id": consultant.id,
                    "is_validated": False
                }, status=201)

            # Erreur de validation
            logger.error(f"Erreur de validation: {serializer.errors}")
            return Response({
                "success": False,
                "errors": serializer.errors
            }, status=400)

    except Exception as e:
        logger.error(f"Erreur lors de l'inscription: {str(e)}")
        return Response({
            "success": False,
            "error": str(e)
        }, status=500)
@api_view(['POST'])
def consultant_login(request):
    """
    Connecte un consultant existant
    """
    logger.info("Tentative de connexion consultant:", request.data)

    email = request.data.get('email')
    password = request.data.get('password')

    if not email or not password:
        logger.warning(f"Email ou mot de passe manquant: email={email}, password={'*' if password else 'None'}")
        return Response({"error": "Email et mot de passe requis"}, status=400)

    try:
        # Recherche de l'utilisateur par email
        user = User.objects.get(username=email)
        logger.info(f"Utilisateur trouv√©: {user.username} (ID: {user.id})")

        # V√©rification du mot de passe
        if not user.check_password(password):
            logger.warning(f"Mot de passe incorrect pour {email}")
            return Response({"error": "Mot de passe incorrect"}, status=400)

        # V√©rification que l'utilisateur est bien un consultant
        if not hasattr(user, 'consultant_profile'):
            logger.warning(f"L'utilisateur {email} n'est pas un consultant")
            # V√©rifier si un consultant existe pour cet utilisateur
            consultants = Consultant.objects.filter(user=user)
            if consultants.exists():
                consultant = consultants.first()
                logger.info(f"Consultant trouv√© manuellement: {consultant.nom} (ID: {consultant.id})")

                # V√©rifier si le consultant est valid√©
                if not consultant.is_validated:
                    return Response({
                        "error": "Votre compte est en attente de validation par un administrateur",
                        "is_validated": False
                    }, status=403)

                return Response({
                    "consultant_id": consultant.id,
                    "is_validated": True
                })
            else:
                return Response({"error": "Ce compte n'est pas un consultant"}, status=400)

        # R√©cup√©ration du consultant associ√©
        consultant = user.consultant_profile

        # V√©rifier si le consultant est valid√©
        if not consultant.is_validated:
            return Response({
                "error": "Votre compte est en attente de validation par un administrateur",
                "is_validated": False
            }, status=403)

        # Connexion r√©ussie
        logger.info(f"Connexion r√©ussie pour {email} - ID consultant: {consultant.id}")
        return Response({
            "consultant_id": consultant.id,
            "is_validated": True
        })

    except User.DoesNotExist:
        logger.warning(f"Utilisateur non trouv√©: {email}")
        return Response({"error": "Email incorrect"}, status=404)
    except Exception as e:
        logger.error(f"Erreur inattendue lors de la connexion: {e}")
        # Tentative de r√©cup√©ration du consultant par email
        try:
            consultant = Consultant.objects.filter(email=email).first()
            if consultant:
                logger.info(f"Consultant trouv√© par email: {consultant.nom} (ID: {consultant.id})")
                if consultant.user and consultant.user.check_password(password):
                    # V√©rifier si le consultant est valid√©
                    if not consultant.is_validated:
                        return Response({
                            "error": "Votre compte est en attente de validation par un administrateur",
                            "is_validated": False
                        }, status=403)

                    return Response({
                        "consultant_id": consultant.id,
                        "is_validated": True
                    })
        except Exception:
            pass

        return Response({"error": "Erreur lors de la connexion"}, status=500)


from decimal import Decimal
import logging
import re
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from datetime import datetime, timedelta

# Configurer le logging
logger = logging.getLogger(__name__)

# Cache des scores calcul√©s (en m√©moire durant l'ex√©cution)
_score_cache = {}

def get_cache_key(consultant_id, appel_offre_id):
    """
    Cr√©e une cl√© de cache unique pour un consultant et un appel d'offre
    """
    return f"{consultant_id}_{appel_offre_id}"

def calculate_date_match_score_updated(consultant, appel_offre):
    """
    Calcule un score bas√© sur la disponibilit√© du consultant par rapport aux dates du projet
    ADAPT√â pour le nouveau mod√®le AppelOffre scrap√©
    """
    try:
        cache_key = f"date_{get_cache_key(consultant.id, appel_offre.id)}"
        if cache_key in _score_cache:
            return _score_cache[cache_key]
        
        # Pour les appels d'offres scrap√©s, nous n'avons que date_limite
        # Nous devons estimer une dur√©e de projet
        if not appel_offre.date_limite:
            logger.warning(f"Pas de date limite pour l'appel d'offre {appel_offre.id}")
            return 50  # Score neutre si pas de date
        
        if not consultant.date_debut_dispo or not consultant.date_fin_dispo:
            logger.warning(f"Dates de disponibilit√© manquantes pour consultant {consultant.id}")
            return 0
        
        # Estimer la date de d√©but du projet (30 jours avant la date limite)
        date_limite = appel_offre.date_limite
        estimated_project_start = date_limite - timedelta(days=30)
        estimated_project_end = date_limite + timedelta(days=60)  # Estimer 60 jours apr√®s la date limite
        
        consultant_start = consultant.date_debut_dispo
        consultant_end = consultant.date_fin_dispo
        
        # Cas o√π il n'y a pas de chevauchement
        if consultant_end < estimated_project_start or consultant_start > estimated_project_end:
            # V√©rifier la proximit√© des dates (flexibilit√©)
            buffer_days = 15  # Jours de tol√©rance
            
            if consultant_end < estimated_project_start:
                days_gap = (estimated_project_start - consultant_end).days
                if days_gap <= buffer_days:
                    score = max(0, 25 * (1 - days_gap / buffer_days))
                    _score_cache[cache_key] = score
                    return score
            else:
                days_gap = (consultant_start - estimated_project_end).days
                if days_gap <= buffer_days:
                    score = max(0, 25 * (1 - days_gap / buffer_days))
                    _score_cache[cache_key] = score
                    return score
            
            _score_cache[cache_key] = 0
            return 0
        
        # Cas o√π la disponibilit√© du consultant couvre enti√®rement le projet estim√©
        if consultant_start <= estimated_project_start and consultant_end >= estimated_project_end:
            score = 100
            _score_cache[cache_key] = score
            return score
        
        # Calcul du chevauchement partiel
        total_project_days = (estimated_project_end - estimated_project_start).days + 1
        if total_project_days <= 0:
            _score_cache[cache_key] = 50
            return 50
        
        overlap_start = max(consultant_start, estimated_project_start)
        overlap_end = min(consultant_end, estimated_project_end)
        overlap_days = (overlap_end - overlap_start).days + 1
        
        coverage_percentage = (overlap_days / total_project_days) * 100
        
        if coverage_percentage >= 80:
            adjusted_score = 85 + (coverage_percentage - 80) / 2
        else:
            adjusted_score = coverage_percentage * 0.85
        
        score = min(100, adjusted_score)
        _score_cache[cache_key] = score
        return score
        
    except Exception as e:
        logger.error(f"Erreur lors du calcul du score de date: {str(e)}")
        return 0


def detect_domain_from_description(description):
    """
    D√©tecte le domaine principal √† partir de la description de l'appel d'offre
    """
    if not description:
        return 'DIGITAL', 0
    
    description_lower = description.lower()
    domain_scores = {
        'DIGITAL': 0,
        'FINANCE': 0,
        'ENERGIE': 0,
        'INDUSTRIE': 0
    }
    
    # Compter les occurrences de comp√©tences par domaine
    for domain, skills_list in ALL_SKILLS.items():
        for skill in skills_list:
            skill_lower = skill.lower()
            if skill_lower in description_lower:
                domain_scores[domain] += 1
    
    # Ajouter des mots-cl√©s sp√©cifiques par domaine
    domain_keywords = {
        'DIGITAL': ['digital', 'num√©rique', 'informatique', 'web', 'mobile', 'logiciel', 'application', 'syst√®me', 'technologie', 'd√©veloppement'],
        'FINANCE': ['finance', 'financier', 'banque', 'bancaire', 'comptabilit√©', 'audit', 'budget', 'investissement', 'cr√©dit'],
        'ENERGIE': ['√©nergie', '√©nerg√©tique', '√©lectricit√©', 'solaire', '√©olien', 'p√©trole', 'gaz', 'renouvelable', 'transition'],
        'INDUSTRIE': ['industrie', 'industriel', 'usine', 'production', 'fabrication', 'manufacture', 'm√©canique', 'mines']
    }
    
    for domain, keywords in domain_keywords.items():
        for keyword in keywords:
            if keyword in description_lower:
                domain_scores[domain] += 2  # Poids plus fort pour les mots-cl√©s g√©n√©raux
    
    # Retourner le domaine avec le score le plus √©lev√©
    best_domain = max(domain_scores.items(), key=lambda x: x[1])
    return best_domain[0], best_domain[1]


def calculate_skills_match_score_updated(consultant, appel_offre):
    """
    Calcule un score bas√© sur les comp√©tences du consultant par rapport aux crit√®res du projet
    ADAPT√â pour le nouveau mod√®le AppelOffre scrap√©
    """
    try:
        cache_key = get_cache_key(consultant.id, appel_offre.id)
        if cache_key in _score_cache:
            return _score_cache[cache_key]
        
        # R√©cup√©rer les comp√©tences du consultant
        consultant_skills = list(
            Competence.objects.filter(consultant=consultant).values_list('nom_competence', 'niveau')
        )
        
        skills_dict = {skill_name.lower(): niveau for skill_name, niveau in consultant_skills}
        
        if not consultant_skills:
            logger.warning(f"Le consultant {consultant.id} n'a aucune comp√©tence d√©finie")
            _score_cache[cache_key] = 10
            return 10
        
        # --- 1. SCORE DE DOMAINE (15% max) ---
        detected_domain, domain_confidence = detect_domain_from_description(appel_offre.description)
        
        domain_score = 0
        if consultant.domaine_principal == detected_domain:
            domain_score = 15  # Correspondance parfaite
            logger.info(f"Correspondance parfaite de domaine: {detected_domain}")
        elif domain_confidence > 0:
            # Correspondance partielle bas√©e sur la confiance
            domain_score = min(12, 15 * (domain_confidence / 10))
        else:
            domain_score = 8  # Score de base pour domaines diff√©rents
        
        # --- 2. SCORE D'EXPERTISE (15% max) ---
        expertise_score = 0
        if consultant.expertise == "Expert" or consultant.expertise == "Senior":
            expertise_score = 15
        elif consultant.expertise == "Interm√©diaire":
            expertise_score = 10
        else:  # D√©butant
            expertise_score = 5
        
        # --- 3. SCORE DE COMP√âTENCES (70% max) ---
        skills_score = 0
        
        # Utiliser les crit√®res structur√©s s'ils existent
        from .models import CriteresEvaluation
        project_criteria = CriteresEvaluation.objects.filter(appel_offre=appel_offre)
        
        if project_criteria.exists():
            logger.info(f"Utilisation des crit√®res structur√©s pour l'AO {appel_offre.id}")
            
            total_weight = float(sum(float(criteria.poids) for criteria in project_criteria))
            weighted_score = 0
            
            if total_weight > 0:
                for criteria in project_criteria:
                    normalized_weight = float(float(criteria.poids) / total_weight) * 70.0
                    keyword = criteria.nom_critere.lower()
                    
                    best_match_score = 0
                    for skill_name, niveau in skills_dict.items():
                        match_score = get_competence_similarity(keyword, skill_name)
                        if match_score > 0:
                            match_score *= (0.7 + 0.3 * (niveau / 5))
                            best_match_score = max(best_match_score, match_score)
                    
                    weighted_score += normalized_weight * best_match_score
                
                skills_score = min(70, weighted_score)
        else:
            # Analyse bas√©e sur la description de l'appel d'offre
            logger.info(f"Analyse bas√©e sur la description pour l'AO {appel_offre.id}")
            
            if appel_offre.description:
                # Extraire des comp√©tences de la description
                mentioned_skills = extract_skills_from_description(appel_offre.description)
                
                if len(mentioned_skills) >= 3:
                    try:
                        # Utiliser une m√©thode de calcul de similarit√© simple
                        skills_score = calculate_alternative_score_updated(mentioned_skills, skills_dict)
                    except Exception as e:
                        logger.error(f"Erreur calcul comp√©tences: {str(e)}")
                        skills_score = calculate_alternative_score_updated(mentioned_skills, skills_dict)
                else:
                    skills_score = calculate_alternative_score_updated(mentioned_skills, skills_dict)
            else:
                # Pas de description d√©taill√©e
                skills_score = 25  # Score r√©duit sans description
        
        # --- 4. CALCUL DU SCORE FINAL ---
        final_score = domain_score + expertise_score + skills_score
        
        # Bonus pour les matchings tr√®s pertinents
        if final_score > 70:
            final_score = final_score * 1.1
            final_score = min(100, final_score)
        
        # Ajustement pour √©viter les scores trop bas
        if final_score < 15:
            final_score = 15
        
        _score_cache[cache_key] = final_score
        
        logger.info(f"Score final pour consultant {consultant.id}: {final_score:.2f}%")
        return final_score
        
    except Exception as e:
        logger.error(f"Erreur lors du calcul du score de comp√©tences: {str(e)}")
        return 20


def extract_skills_from_description(description):
    """
    Extrait les comp√©tences techniques de la description d'un appel d'offre
    """
    if not description:
        return []
    
    description_lower = description.lower()
    found_skills = []
    
    # Rechercher toutes les comp√©tences connues dans la description
    for domain, skills_list in ALL_SKILLS.items():
        for skill in skills_list:
            skill_lower = skill.lower()
            if skill_lower in description_lower:
                found_skills.append(skill_lower)
    
    # Ajouter des mots techniques g√©n√©raux
    technical_terms = [
        'gestion', 'analyse', 'conseil', 'audit', 'formation', 'expertise',
        'd√©veloppement', 'conception', 'mise en ≈ìuvre', 'optimisation',
        'strat√©gie', 'planification', 'coordination', 'supervision'
    ]
    
    for term in technical_terms:
        if term in description_lower:
            found_skills.append(term)
    
    return list(set(found_skills))  # Supprimer les doublons


def calculate_alternative_score_updated(mentioned_skills, consultant_skills_dict):
    """
    M√©thode alternative de calcul du score quand TF-IDF ne peut pas √™tre utilis√©
    """
    try:
        if not mentioned_skills:
            return 35  # Score neutre
        
        matched_count = 0
        matched_skills = set()
        
        for keyword in mentioned_skills:
            best_match_score = 0
            
            for skill_name, niveau in consultant_skills_dict.items():
                if keyword == skill_name:
                    match_score = 1.0
                elif keyword in skill_name or skill_name in keyword:
                    match_score = 0.8
                else:
                    match_score = get_competence_similarity(keyword, skill_name)
                
                if match_score > 0:
                    match_score *= (0.6 + 0.4 * (niveau / 5))
                    best_match_score = max(best_match_score, match_score)
            
            if best_match_score > 0:
                matched_count += best_match_score
                matched_skills.add(keyword)
        
        match_ratio = matched_count / len(mentioned_skills) if mentioned_skills else 0
        skills_score = min(70, 70 * match_ratio)
        
        logger.info(f"Score par analyse alternative: {skills_score:.2f}%")
        return skills_score
        
    except Exception as e:
        logger.error(f"Erreur dans le calcul alternatif: {str(e)}")
        return 35


def get_competence_similarity(comp1, comp2):
    """
    Calcule la similarit√© entre deux comp√©tences
    """
    import re
    
    comp1 = comp1.lower()
    comp2 = comp2.lower()
    
    if comp1 == comp2:
        return 1.0
    
    if comp1 in comp2 or comp2 in comp1:
        ratio = min(len(comp1), len(comp2)) / max(len(comp1), len(comp2))
        return 0.8 * ratio
    
    # Pour les technos avec version
    base_comp1 = re.sub(r'\s+\d+(\.\d+)*', '', comp1)
    base_comp2 = re.sub(r'\s+\d+(\.\d+)*', '', comp2)
    
    if base_comp1 == base_comp2:
        return 0.9
    
    # Tokens communs
    tokens1 = set(re.findall(r'\b\w+\b', comp1))
    tokens2 = set(re.findall(r'\b\w+\b', comp2))
    
    if not tokens1 or not tokens2:
        return 0
        
    common_tokens = tokens1.intersection(tokens2)
    
    if common_tokens:
        jaccard = len(common_tokens) / len(tokens1.union(tokens2))
        return 0.6 * jaccard
    
    return 0


def get_top_skills_updated(consultant, limit=5):
    """
    R√©cup√®re les comp√©tences principales d'un consultant
    """
    try:
        top_skills = Competence.objects.filter(consultant=consultant).order_by('-niveau')[:limit]
        return [skill.nom_competence for skill in top_skills]
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des comp√©tences: {str(e)}")
        return []
@api_view(['GET'])
def validated_matches_updated(request):
    """
    R√©cup√®re la liste des matchings valid√©s - COMPATIBLE NOUVEAU MOD√àLE
    """
    try:
        matches = MatchingResult.objects.filter(is_validated=True).select_related('consultant', 'appel_offre')

        data = []
        for match in matches:
            # Adaptation pour le nouveau mod√®le AppelOffre
            appel_offre = match.appel_offre
            
            data.append({
                'id': match.id,
                'consultant_id': match.consultant.id,
                'consultant_name': f"{match.consultant.prenom} {match.consultant.nom}",
                'appel_offre_id': appel_offre.id,
                'appel_offre_name': appel_offre.titre,  # Nouveau mod√®le utilise 'titre' au lieu de 'nom_projet'
                'client': appel_offre.client or "Client non sp√©cifi√©",
                'score': float(match.score),
                'date_validation': match.created_at,
                'domaine_principal': match.consultant.domaine_principal,
                'consultant_expertise': match.consultant.expertise,
                'email': match.consultant.email,
                # Nouveaux champs pour diff√©rencier les sources
                'is_scraped_offer': True,  # Tous les appels d'offres du nouveau mod√®le sont scrap√©s
                'source_type': 'scraped',
                # Dates d'appel d'offre (nouveau mod√®le)
                'date_de_publication': appel_offre.date_de_publication.isoformat() if appel_offre.date_de_publication else None,
                'date_limite': appel_offre.date_limite.isoformat() if appel_offre.date_limite else None,
                'is_expired': appel_offre.is_expired,
                'days_remaining': appel_offre.days_remaining
            })

        return Response(data)
    except Exception as e:
        logger.error(f"Erreur dans validated_matches_updated: {str(e)}")
        return Response({"error": str(e)}, status=500)
 
def generate_matching_for_offer_updated(appel_offre_id):
    """
    G√©n√®re des matchings pour un appel d'offre scrap√©
    Version adapt√©e pour le nouveau mod√®le AppelOffre
    """
    try:
        # R√©cup√©rer l'appel d'offre scrap√©
        try:
            appel_offre = AppelOffre.objects.get(id=appel_offre_id)
        except AppelOffre.DoesNotExist:
            logger.error(f"Appel d'offre avec ID {appel_offre_id} introuvable")
            return {
                'success': False,
                'error': f"Appel d'offre avec ID {appel_offre_id} introuvable"
            }

        logger.info(f"G√©n√©ration de matchings pour l'appel d'offre scrap√©: {appel_offre.titre}")
        
        # R√©cup√©rer les consultants valid√©s et disponibles
        consultants = Consultant.objects.filter(
            is_validated=True,
            statut='Actif'
        ).exclude(
            date_debut_dispo=None
        ).exclude(
            date_fin_dispo=None
        )
        
        if not consultants.exists():
            logger.warning(f"Aucun consultant disponible pour le matching")
            return {
                'success': False,
                'error': "Aucun consultant disponible pour le matching"
            }
        
        # Vider les anciens matchings
        MatchingResult.objects.filter(appel_offre=appel_offre).delete()
        logger.info(f"Anciens matchings supprim√©s pour l'AO {appel_offre_id}")
        
        # Vider le cache
        clear_score_cache()
        
        results = []
        score_stats = {"min": 100, "max": 0, "total": 0}
        
        # Calculer les scores pour chaque consultant
        for consultant in consultants:
            try:
                # Calculer les scores avec les nouvelles m√©thodes
                date_score = calculate_date_match_score_updated(consultant, appel_offre)
                skills_score = calculate_skills_match_score_updated(consultant, appel_offre)
                
                # Pond√©ration: 30% date, 70% comp√©tences (plus d'importance aux comp√©tences)
                final_score = (date_score * 0.3) + (skills_score * 0.7)
                final_score = min(100, final_score)
                
                # Mettre √† jour les statistiques
                score_stats["min"] = min(score_stats["min"], final_score)
                score_stats["max"] = max(score_stats["max"], final_score)
                score_stats["total"] += final_score
                
                logger.info(f"Score final pour {consultant.nom}: {final_score:.2f}%")
                
                # Enregistrer le r√©sultat
                matching = MatchingResult.objects.create(
                    consultant=consultant,
                    appel_offre=appel_offre,
                    score=Decimal(str(final_score)),
                    is_validated=False
                )
                
                # Ajouter √† la liste de r√©sultats
                results.append({
                    'id': matching.id,
                    'consultant_id': consultant.id,
                    'consultant_name': f"{consultant.prenom} {consultant.nom}",
                    'consultant_expertise': consultant.expertise or "D√©butant",
                    'email': consultant.email,
                    'domaine_principal': consultant.domaine_principal,
                    'specialite': consultant.specialite or "",
                    'top_skills': get_top_skills_updated(consultant),
                    'date_match_score': float(date_score),
                    'skills_match_score': float(skills_score),
                    'score': float(final_score),
                    'is_validated': False
                })
                
            except Exception as e:
                logger.error(f"Erreur lors du calcul pour le consultant {consultant.id}: {str(e)}")
                continue
        
        # Calculer les statistiques finales
        if results:
            score_stats["avg"] = score_stats["total"] / len(results)
            score_stats["count"] = len(results)
            
            logger.info(f"Matchings g√©n√©r√©s: {len(results)}, score moyen: {score_stats['avg']:.2f}%")
        
        # Trier les r√©sultats par score d√©croissant
        sorted_results = sorted(results, key=lambda x: x['score'], reverse=True)
        
        return {
            'success': True,
            'matches': sorted_results,
            'stats': score_stats,
            'appel_offre_info': {
                'id': appel_offre.id,
                'titre': appel_offre.titre,
                'client': appel_offre.client,
                'has_description': bool(appel_offre.description),
                'has_criteria': CriteresEvaluation.objects.filter(appel_offre=appel_offre).exists()
            }
        }
        
    except Exception as e:
        logger.error(f"Erreur lors de la g√©n√©ration des matchings: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }

def get_top_skills_updated(consultant, limit=5):
    """
    R√©cup√®re les comp√©tences principales d'un consultant
    """
    try:
        top_skills = Competence.objects.filter(consultant=consultant).order_by('-niveau')[:limit]
        return [skill.nom_competence for skill in top_skills]
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des comp√©tences: {str(e)}")
        return []

@api_view(['GET', 'POST'])
def matching_for_offer_updated(request, appel_offre_id):
    """
    Endpoint pour g√©n√©rer ou r√©cup√©rer des matchings pour un appel d'offre scrap√©
    Version adapt√©e pour le nouveau mod√®le AppelOffre
    """
    logger.info(f"Requ√™te {request.method} re√ßue pour appel d'offre scrap√© ID {appel_offre_id}")
    
    try:
        appel_offre_id = int(appel_offre_id)
    except ValueError:
        return Response({
            'success': False,
            'error': f"ID d'appel d'offre invalide: {appel_offre_id}"
        }, status=400)

    if request.method == 'GET':
        try:
            # V√©rifier que l'appel d'offre existe
            try:
                appel_offre = AppelOffre.objects.get(id=appel_offre_id)
            except AppelOffre.DoesNotExist:
                return Response({
                    'success': False,
                    'error': f"Appel d'offre avec ID {appel_offre_id} introuvable"
                }, status=404)

            # R√©cup√©rer tous les matchings pour cet appel d'offre
            matches = MatchingResult.objects.filter(
                appel_offre_id=appel_offre_id
            ).order_by('-score')

            logger.info(f"{matches.count()} matchings trouv√©s pour l'AO scrap√© {appel_offre_id}")

            if not matches.exists():
                return Response({
                    'success': True,
                    'matches': []
                })

            result = []
            for match in matches:
                consultant = match.consultant

                try:
                    # Recalculer les scores si n√©cessaire
                    date_score = calculate_date_match_score_updated(consultant, match.appel_offre)
                    skills_score = calculate_skills_match_score_updated(consultant, match.appel_offre)
                    calculated_score = (date_score * 0.3) + (skills_score * 0.7)
                    stored_score = float(match.score)
                    
                    # V√©rifier la coh√©rence
                    if abs(calculated_score - stored_score) > 5:
                        match.score = Decimal(str(calculated_score))
                        match.save(update_fields=['score'])
                        stored_score = calculated_score
                        
                except Exception as e:
                    logger.error(f"Erreur lors du recalcul des scores: {str(e)}")
                    date_score = 0
                    skills_score = 0
                    stored_score = float(match.score)

                stored_score = max(0, min(100, float(stored_score)))

                result.append({
                    'id': match.id,
                    'consultant_id': consultant.id,
                    'consultant_name': f"{consultant.prenom} {consultant.nom}",
                    'consultant_expertise': consultant.expertise or "D√©butant",
                    'email': consultant.email,
                    'domaine_principal': consultant.domaine_principal,
                    'specialite': consultant.specialite or "",
                    'top_skills': get_top_skills_updated(consultant),
                    'date_match_score': round(date_score, 1),
                    'skills_match_score': round(skills_score, 1),
                    'score': stored_score,
                    'is_validated': match.is_validated
                })

            sorted_result = sorted(result, key=lambda x: x['score'], reverse=True)

            return Response({
                'success': True,
                'matches': sorted_result
            })
            
        except Exception as e:
            logger.error(f"Erreur lors de la r√©cup√©ration des matchings: {str(e)}")
            return Response({
                'success': False,
                'error': f"Erreur lors de la r√©cup√©ration des matchings: {str(e)}"
            }, status=500)

    elif request.method == 'POST':
        try:
            logger.info(f"G√©n√©ration de nouveaux matchings pour l'AO scrap√© {appel_offre_id}")
            
            clear_score_cache()
            
            result = generate_matching_for_offer_updated(appel_offre_id)

            if not isinstance(result, dict):
                return Response({
                    'success': False,
                    'error': "Erreur interne: format de r√©sultat invalide"
                }, status=500)

            if result.get('success') and not result.get('matches', []):
                return Response({
                    'success': False,
                    'error': "Aucun consultant disponible pour le matching"
                }, status=404)
                
            # V√©rification des scores g√©n√©r√©s
            if result.get('success') and result.get('matches'):
                for i, match in enumerate(result['matches'][:5]):
                    logger.info(f"Match g√©n√©r√© {i+1}: {match['consultant_name']}, Score: {match['score']}")
                
                scores = [match['score'] for match in result['matches']]
                unique_scores = set(scores)
                if len(unique_scores) == 1:
                    logger.warning(f"ATTENTION: Tous les consultants ont le m√™me score: {list(unique_scores)[0]}")

            return Response(result)
            
        except Exception as e:
            logger.error(f"Exception lors de la g√©n√©ration des matchings: {str(e)}")
            return Response({
                'success': False,
                'error': f"Erreur lors de la g√©n√©ration des matchings: {str(e)}"
            }, status=500)
    else:
        return Response({
            'success': False,
            'error': f"M√©thode {request.method} non support√©e"
        }, status=405)

def create_notification_for_consultant(consultant, notification_type, title, content, appel_offre=None, match=None):
    """
    Fonction helper pour cr√©er des notifications de mani√®re s√©curis√©e
    """
    try:
        notification = Notification.objects.create(
            consultant=consultant,
            type=notification_type,
            title=title,
            content=content,
            related_appel=appel_offre,
            related_match=match,
            is_read=False
        )
        
        logger.info(f"‚úÖ Notification cr√©√©e: {title} pour {consultant.nom} {consultant.prenom}")
        return notification
        
    except Exception as e:
        logger.error(f"‚ùå Erreur cr√©ation notification: {str(e)}")
        return None

    """
    R√©cup√®re les notifications d'un consultant - VERSION CORRIG√âE
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        
        logger.info(f"üîî R√©cup√©ration des notifications pour consultant {consultant_id}")

        # R√©cup√©rer toutes les notifications avec relations
        notifications = Notification.objects.filter(
            consultant=consultant
        ).select_related('related_appel', 'related_match').order_by('-created_at')

        logger.info(f"üìä Notifications trouv√©es: {notifications.count()}")

        # Formater pour l'API
        results = []
        for notif in notifications:
            try:
                notification_data = {
                    'id': notif.id,
                    'type': notif.type,
                    'title': notif.title,
                    'content': notif.content,
                    'is_read': notif.is_read,
                    'created_at': notif.created_at.isoformat(),
                    'appel_offre_id': notif.related_appel.id if notif.related_appel else None,
                    'appel_offre_nom': notif.related_appel.titre if notif.related_appel else None,  # Nouveau mod√®le
                    'match_id': notif.related_match.id if notif.related_match else None
                }
                
                results.append(notification_data)
                logger.debug(f"‚úÖ Notification format√©e: {notif.title}")
                
            except Exception as notif_error:
                logger.error(f"‚ùå Erreur formatage notification {notif.id}: {str(notif_error)}")
                continue

        # Compter les notifications non lues
        unread_count = notifications.filter(is_read=False).count()
        
        logger.info(f"‚úÖ {len(results)} notifications format√©es, {unread_count} non lues")

        return Response({
            'success': True,
            'notifications': results,
            'unread_count': unread_count,
            'total_count': len(results)
        })
        
    except Exception as e:
        logger.error(f"‚ùå Erreur lors de la r√©cup√©ration des notifications: {str(e)}")
        return Response({
            'success': False,
            'error': str(e),
            'notifications': [],
            'unread_count': 0
        }, status=500)
@api_view(['PUT'])
def mark_notification_read(request, notification_id):
    """
    Marque une notification comme lue - VERSION CORRIG√âE
    """
    try:
        notification = get_object_or_404(Notification, id=notification_id)
        
        # Marquer comme lue seulement si pas d√©j√† lue
        if not notification.is_read:
            notification.is_read = True
            notification.save(update_fields=['is_read'])
            logger.info(f"‚úÖ Notification {notification_id} marqu√©e comme lue")
        
        return Response({
            'success': True,
            'message': 'Notification marqu√©e comme lue',
            'notification_id': notification_id
        })
        
    except Exception as e:
        logger.error(f"‚ùå Erreur lors du marquage de la notification {notification_id}: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)

@api_view(['PUT'])
def validate_match(request, match_id):
    """
    Endpoint pour valider un matching avec cr√©ation automatique de notification
    VERSION CORRIG√âE avec gestion compl√®te des notifications
    """
    try:
        match = get_object_or_404(MatchingResult, id=match_id)
        consultant = match.consultant
        appel_offre = match.appel_offre

        was_validated = match.is_validated

        # Inverser l'√©tat de validation
        match.is_validated = not match.is_validated
        match.save()

        notification_created = False
        mission_created = False
        
        # üî• Si le match vient d'√™tre valid√©, cr√©er notification ET mission
        if not was_validated and match.is_validated:
            logger.info(f"üéØ Validation d'un matching pour {consultant.nom} {consultant.prenom}")
            
            # 1. Cr√©er la mission automatiquement
            try:
                mission = create_mission_from_matching(match)
                if mission:
                    mission_created = True
                    logger.info(f"‚úÖ Mission cr√©√©e automatiquement: ID {mission.id}")
                
            except Exception as mission_error:
                logger.error(f"‚ùå Erreur cr√©ation mission: {str(mission_error)}")
            
            # 2. CORRECTION - Cr√©er la notification avec le nouveau mod√®le
            try:
                from .models import Notification
                
                title = f"üéâ Mission confirm√©e : {getattr(appel_offre, 'titre', 'Nouvelle mission')}"
                content = (
                    f"F√©licitations ! Votre profil a √©t√© s√©lectionn√© pour la mission "
                    f"'{getattr(appel_offre, 'titre', 'la mission')}' chez {getattr(appel_offre, 'client', 'le client')}. "
                    f"Score de matching: {float(match.score):.1f}%. "
                    f"Vous pouvez maintenant consulter les d√©tails dans la section 'Mes Missions'."
                )

                notification = Notification.objects.create(
                    consultant=consultant,
                    type="MATCH_VALID",
                    title=title,
                    content=content,
                    priority="HIGH",  # Priorit√© √©lev√©e pour les validations
                    related_appel=appel_offre,
                    related_match=match,
                    related_mission=mission if mission_created else None,
                    is_read=False,
                    metadata={
                        'match_score': float(match.score),
                        'client': getattr(appel_offre, 'client', ''),
                        'mission_auto_created': mission_created,
                        'validation_date': timezone.now().isoformat()
                    }
                )
                
                notification_created = True
                logger.info(f"‚úÖ Notification cr√©√©e: ID {notification.id}")
                
            except Exception as notif_error:
                logger.error(f"‚ùå Erreur cr√©ation notification: {str(notif_error)}")

        # üî• Si le match est invalid√©, cr√©er une notification d'information
        elif was_validated and not match.is_validated:
            try:
                from .models import Notification
                
                title = f"‚ÑπÔ∏è Statut de mission mis √† jour"
                content = (
                    f"Le statut de votre candidature pour la mission "
                    f"'{getattr(appel_offre, 'titre', 'la mission')}' a √©t√© mis √† jour. "
                    f"Contactez l'administrateur pour plus d'informations."
                )

                notification = Notification.objects.create(
                    consultant=consultant,
                    type="MISSION_UPDATE",
                    title=title,
                    content=content,
                    priority="NORMAL",
                    related_appel=appel_offre,
                    related_match=match,
                    is_read=False,
                    metadata={
                        'action': 'invalidation',
                        'previous_score': float(match.score)
                    }
                )
                
                notification_created = True
                logger.info(f"‚úÖ Notification d'invalidation cr√©√©e")
                
            except Exception as notif_error:
                logger.error(f"‚ùå Erreur notification invalidation: {str(notif_error)}")

        return Response({
            'success': True,
            'is_validated': match.is_validated,
            'mission_created': mission_created,
            'notification_created': notification_created,
            'message': f"Statut de validation mis √† jour avec succ√®s: {'valid√©' if match.is_validated else 'non valid√©'}",
            'notification_status': 'Notification envoy√©e au consultant' if notification_created else 'Aucune notification cr√©√©e'
        })

    except Exception as e:
        logger.error(f"‚ùå Erreur lors de la validation du matching {match_id}: {str(e)}")
        return Response({
            'success': False,
            'error': f"Erreur lors de la validation: {str(e)}"
        }, status=500)


@api_view(['GET'])
def consultant_matches(request, consultant_id):
    """
    Endpoint pour r√©cup√©rer les meilleurs matchings pour un consultant
    """
    limit = int(request.query_params.get('limit', 5))

    try:
        matchings = MatchingResult.objects.filter(
            consultant_id=consultant_id
        ).order_by('-score')[:limit]

        results = []
        for matching in matchings:
            results.append({
                'id': matching.id,
                'appel_offre_id': matching.appel_offre.id,
                'appel_offre_name': matching.appel_offre.nom_projet,
                'client': matching.appel_offre.client,
                'date_debut': matching.appel_offre.date_debut,
                'date_fin': matching.appel_offre.date_fin,
                'score': matching.score,
                'is_validated': matching.is_validated
            })

        return Response({
            'success': True,
            'matches': results
        })

    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des matchings pour le consultant {consultant_id}: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)


@api_view(['POST'])
def admin_login(request):
    """
    Connecte un administrateur existant
    """
    username = request.data.get('username')
    password = request.data.get('password')

    logger.info(f"Tentative de connexion admin: {username}")

    try:
        user = User.objects.get(username=username)

        if not user.check_password(password):
            return Response({"error": "Mot de passe incorrect"}, status=400)

        if user.role != 'ADMIN':
            return Response({"error": "Ce compte n'est pas un administrateur"}, status=403)

        return Response({
            "admin_id": user.id,
            "username": user.username,
            "role": user.role
        })
    except User.DoesNotExist:
        return Response({"error": "Utilisateur non trouv√©"}, status=404)
    except Exception as e:
        logger.error(f"Erreur lors de la connexion admin: {e}")
        return Response({"error": "Une erreur s'est produite lors de la connexion"}, status=500)



@api_view(['GET'])
def appel_offre_detail(request, pk):
    """
    Vue publique pour r√©cup√©rer les d√©tails d'un appel d'offre
    Compatible avec le nouveau mod√®le
    """
    try:
        appel = get_object_or_404(AppelOffre, id=pk)
        
        appel_data = {
            'id': appel.id,
            'titre': appel.titre,
            'date_de_publication': appel.date_de_publication.isoformat() if appel.date_de_publication else None,
            'date_limite': appel.date_limite.isoformat() if appel.date_limite else None,
            'client': appel.client,
            'type_d_appel_d_offre': appel.type_d_appel_d_offre,
            'description': appel.description,
            'critere_evaluation': appel.critere_evaluation,
            'documents': appel.documents,
            'lien_site': appel.lien_site,
            'created_at': appel.created_at.isoformat(),
            'updated_at': appel.updated_at.isoformat(),
            'is_expired': appel.is_expired,
            'days_remaining': appel.days_remaining
        }
        
        # Informations contextuelles pour le frontend
        appel_data['source_info'] = {
            'is_scraped': True,
            'has_external_link': bool(appel.lien_site),
            'publication_date': appel.date_de_publication.isoformat() if appel.date_de_publication else None,
            'deadline_date': appel.date_limite.isoformat() if appel.date_limite else None
        }
        
        # Informations sur l'enrichissement
        appel_data['enrichment_info'] = {
            'description_length': len(appel.description) if appel.description else 0,
            'has_evaluation_criteria': bool(appel.critere_evaluation),
            'criteria_length': len(appel.critere_evaluation) if appel.critere_evaluation else 0,
            'has_documents_info': bool(appel.documents),
            'structured_criteria_count': CriteresEvaluation.objects.filter(appel_offre=appel).count()
        }
        
        return Response(appel_data)
        
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des d√©tails de l'AO {pk}: {str(e)}")
        return Response({"error": str(e)}, status=500)

@api_view(['GET', 'POST', 'DELETE'])
def appel_offre_criteres(request, appel_id):
    """
    G√®re les crit√®res structur√©s d'un appel d'offre - ADAPT√â AU NOUVEAU MOD√àLE
    Compatible avec le syst√®me de crit√®res d'AppelsOffres.tsx
    """
    try:
        appel = get_object_or_404(AppelOffre, id=appel_id)

        if request.method == 'GET':
            # R√©cup√©rer tous les crit√®res structur√©s
            criteres = CriteresEvaluation.objects.filter(appel_offre=appel).order_by('-poids', 'nom_critere')
            
            # Format compatible avec le frontend
            criteres_data = []
            for critere in criteres:
                criteres_data.append({
                    'id': critere.id,
                    'nom_critere': critere.nom_critere,
                    'poids': float(critere.poids),
                    'description': critere.description
                })
            
            # Informations contextuelles
            response_data = {
                'criteres': criteres_data,
                'total_count': len(criteres_data),
                'appel_offre_info': {
                    'id': appel.id,
                    'titre': appel.titre,
                    'has_text_criteria': bool(appel.critere_evaluation),
                    'text_criteria_length': len(appel.critere_evaluation) if appel.critere_evaluation else 0
                }
            }
            
            return Response(response_data)

        elif request.method == 'POST':
            # Ajouter un nouveau crit√®re structur√©
            nom_critere = request.data.get('nom_critere', '').strip()
            poids = request.data.get('poids', 1)
            
            # Validation
            if not nom_critere:
                return Response({
                    'error': 'Le nom du crit√®re est requis'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # V√©rifier les doublons
            if CriteresEvaluation.objects.filter(
                appel_offre=appel, 
                nom_critere__iexact=nom_critere
            ).exists():
                return Response({
                    'error': 'Un crit√®re avec ce nom existe d√©j√†'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Cr√©er le crit√®re
            try:
                critere = CriteresEvaluation.objects.create(
                    appel_offre=appel,
                    nom_critere=nom_critere,
                    poids=poids,
                    description=request.data.get('description', '')
                )
                
                logger.info(f"Crit√®re '{critere.nom_critere}' ajout√© √† l'AO {appel_id}")
                
                return Response({
                    'id': critere.id,
                    'nom_critere': critere.nom_critere,
                    'poids': float(critere.poids),
                    'description': critere.description
                }, status=201)
                
            except Exception as create_error:
                logger.error(f"Erreur cr√©ation crit√®re: {str(create_error)}")
                return Response({
                    'error': f'Erreur lors de la cr√©ation: {str(create_error)}'
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        elif request.method == 'DELETE':
            # Supprimer tous les crit√®res de cet appel d'offre
            deleted_count, _ = CriteresEvaluation.objects.filter(appel_offre=appel).delete()
            
            logger.info(f"{deleted_count} crit√®res supprim√©s pour l'AO {appel_id}")
            
            return Response({
                'message': f'{deleted_count} crit√®res supprim√©s',
                'appel_offre_id': appel_id,
                'deleted_count': deleted_count
            }, status=200)
            
    except Exception as e:
        logger.error(f"Erreur dans appel_offre_criteres pour AO {appel_id}: {str(e)}")
        return Response({
            'error': f'Erreur lors de la gestion des crit√®res: {str(e)}'
        }, status=500)

@api_view(['POST'])
def bulk_enrich_appels_offres(request):
    """
    NOUVEAU: Enrichissement en lot des appels d'offres
    """
    try:
        appel_ids = request.data.get('appel_ids', [])
        enrichment_data = request.data.get('enrichment_data', {})
        
        if not appel_ids:
            return Response({
                'error': 'Liste des IDs d\'appels d\'offres requise'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Champs autoris√©s pour l'enrichissement en lot
        allowed_fields = ['type_d_appel_d_offre']
        update_data = {k: v for k, v in enrichment_data.items() if k in allowed_fields}
        
        if not update_data:
            return Response({
                'error': 'Aucune donn√©e d\'enrichissement valide fournie',
                'allowed_fields': allowed_fields
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Mise √† jour en lot
        updated_count = AppelOffre.objects.filter(id__in=appel_ids).update(**update_data)
        
        logger.info(f"Enrichissement en lot: {updated_count} appels d'offres mis √† jour")
        
        return Response({
            'message': f'{updated_count} appels d\'offres enrichis avec succ√®s',
            'updated_count': updated_count,
            'updated_fields': list(update_data.keys())
        })
        
    except Exception as e:
        logger.error(f"Erreur dans bulk_enrich_appels_offres: {str(e)}")
        return Response({
            'error': f'Erreur lors de l\'enrichissement en lot: {str(e)}'
        }, status=500)
@api_view(['GET'])
def appels_offres_stats(request):
    """
    NOUVEAU: Statistiques sur les appels d'offres scrap√©s
    """
    try:
        from django.db.models import Count, Q
        from django.utils import timezone
        
        # Statistiques de base
        total_appels = AppelOffre.objects.count()
        
        # Appels actifs (date limite non d√©pass√©e ou non d√©finie)
        now = timezone.now().date()
        appels_actifs = AppelOffre.objects.filter(
            Q(date_limite__gte=now) | Q(date_limite__isnull=True)
        ).count()
        
        # Appels expir√©s
        appels_expires = AppelOffre.objects.filter(date_limite__lt=now).count()
        
        # Appels enrichis (avec description et crit√®res)
        appels_enrichis = AppelOffre.objects.filter(
            description__isnull=False,
            critere_evaluation__isnull=False
        ).exclude(description='').exclude(critere_evaluation='').count()
        
        # Appels avec crit√®res structur√©s
        appels_avec_criteres = AppelOffre.objects.annotate(
            criteria_count=Count('criteres')
        ).filter(criteria_count__gt=0).count()
        
        # Statistiques par type
        stats_par_type = AppelOffre.objects.values('type_d_appel_d_offre').annotate(
            count=Count('id')
        ).order_by('-count')
        
        # Statistiques par client
        stats_par_client = AppelOffre.objects.values('client').annotate(
            count=Count('id')
        ).order_by('-count')[:10]  # Top 10 clients
        
        # Appels r√©cents (7 derniers jours)
        week_ago = now - timedelta(days=7)
        appels_recents = AppelOffre.objects.filter(created_at__gte=week_ago).count()
        
        response_data = {
            'totals': {
                'total_appels': total_appels,
                'appels_actifs': appels_actifs,
                'appels_expires': appels_expires,
                'appels_enrichis': appels_enrichis,
                'appels_avec_criteres': appels_avec_criteres,
                'appels_recents': appels_recents
            },
            'percentages': {
                'taux_enrichissement': (appels_enrichis / total_appels * 100) if total_appels > 0 else 0,
                'taux_criteres_structures': (appels_avec_criteres / total_appels * 100) if total_appels > 0 else 0,
                'taux_actifs': (appels_actifs / total_appels * 100) if total_appels > 0 else 0
            },
            'repartition_types': list(stats_par_type),
            'top_clients': list(stats_par_client),
            'periode_analyse': {
                'date_analyse': now.isoformat(),
                'periode_recente': f"{week_ago.isoformat()} - {now.isoformat()}"
            }
        }
        
        return Response(response_data)
        
    except Exception as e:
        logger.error(f"Erreur dans appels_offres_stats: {str(e)}")
        return Response({
            'error': f'Erreur lors du calcul des statistiques: {str(e)}'
        }, status=500)

@api_view(['PUT', 'DELETE'])
def critere_detail(request, critere_id):
    """
    Modifie ou supprime un crit√®re sp√©cifique
    NOUVEAU: Gestion individuelle des crit√®res structur√©s
    """
    try:
        critere = get_object_or_404(CriteresEvaluation, id=critere_id)
        
        if request.method == 'PUT':
            # Modification d'un crit√®re
            serializer = CriteresEvaluationSerializer(critere, data=request.data, partial=True)
            if serializer.is_valid():
                updated_critere = serializer.save()
                logger.info(f"Crit√®re {critere_id} modifi√©: {updated_critere.nom_critere}")
                return Response(serializer.data)
            else:
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
                
        elif request.method == 'DELETE':
            # Suppression d'un crit√®re
            critere_name = critere.nom_critere
            appel_offre_id = critere.appel_offre.id
            critere.delete()
            
            logger.info(f"Crit√®re '{critere_name}' supprim√© de l'AO {appel_offre_id}")
            
            return Response({
                'message': f'Crit√®re "{critere_name}" supprim√© avec succ√®s'
            }, status=status.HTTP_204_NO_CONTENT)
            
    except Exception as e:
        logger.error(f"Erreur dans critere_detail pour crit√®re {critere_id}: {str(e)}")
        return Response({
            'error': f'Erreur lors de la gestion du crit√®re: {str(e)}'
        }, status=500)
@api_view(['GET'])
def validated_matches(request):
    """
    R√©cup√®re la liste des matchings valid√©s
    """
    try:
        matches = MatchingResult.objects.filter(is_validated=True).select_related('consultant', 'appel_offre')

        data = []
        for match in matches:
            data.append({
                'id': match.id,
                'consultant_id': match.consultant.id,
                'consultant_name': f"{match.consultant.prenom} {match.consultant.nom}",
                'appel_offre_id': match.appel_offre.id,
                'appel_offre_name': match.appel_offre.nom_projet,
                'client': match.appel_offre.client,
                'score': match.score,
                'date_validation': match.created_at,
                'domaine_principal': match.consultant.domaine_principal
            })

        return Response(data)
    except Exception as e:
        logger.error(f"Erreur dans validated_matches: {str(e)}")
        return Response({"error": str(e)}, status=500)


@api_view(['GET'])
def consultant_data(request, consultant_id):
    """
    R√©cup√®re les donn√©es d'un consultant
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        logger.info(
            f"R√©cup√©ration des donn√©es pour consultant ID {consultant_id}: {consultant.nom} {consultant.prenom}")

        # R√©cup√©ration des comp√©tences
        competences = Competence.objects.filter(consultant=consultant)
        competences_list = [c.nom_competence for c in competences]
        logger.info(f"Comp√©tences trouv√©es: {competences_list}")

        # R√©ponse avec les donn√©es du consultant
        return Response({
            "firstName": consultant.prenom,
            "lastName": consultant.nom,
            "email": consultant.user.email if hasattr(consultant, 'user') and consultant.user else consultant.email,
            "phone": consultant.telephone,
            "country": consultant.pays,
            "city": consultant.ville,
            "startAvailability": consultant.date_debut_dispo,
            "endAvailability": consultant.date_fin_dispo,
            "skills": ", ".join(competences_list),
            "expertise": consultant.expertise,
            "domaine_principal": consultant.domaine_principal,
            "specialite": consultant.specialite,
            "cvFilename": consultant.cv.name.split('/')[-1] if consultant.cv else None
        })

    except Exception as e:
        logger.error(f"Erreur r√©cup√©ration consultant {consultant_id}: {e}")
        return Response({"error": f"Erreur lors de la r√©cup√©ration des donn√©es: {str(e)}"}, status=500)



    """
    SOLUTION FINALE - Liste ou cr√©e des consultants (acc√®s admin) - VERSION CORRIG√âE
    """
    try:
        if request.method == 'GET':
            from django.db import connection
            
            # R√©cup√©rer seulement les consultants valid√©s
            consultants = Consultant.objects.filter(is_validated=True)
            
            consultants_data = []
            for consultant in consultants:
                try:
                    consultant_data = {
                        'id': consultant.id,
                        'nom': getattr(consultant, 'nom', ''),
                        'prenom': getattr(consultant, 'prenom', ''),
                        'firstName': getattr(consultant, 'prenom', ''),
                        'lastName': getattr(consultant, 'nom', ''),
                        'email': getattr(consultant, 'email', ''),
                        'telephone': getattr(consultant, 'telephone', ''),
                        'phone': getattr(consultant, 'telephone', ''),
                        'pays': getattr(consultant, 'pays', ''),
                        'country': getattr(consultant, 'pays', ''),
                        'ville': getattr(consultant, 'ville', ''),
                        'city': getattr(consultant, 'ville', ''),
                        
                        # SEULEMENT LES CHAMPS QUI EXISTENT
                        'date_debut_dispo': consultant.date_debut_dispo.isoformat() if consultant.date_debut_dispo else None,
                        'date_fin_dispo': consultant.date_fin_dispo.isoformat() if consultant.date_fin_dispo else None,
                        
                        'domaine_principal': getattr(consultant, 'domaine_principal', 'DIGITAL'),
                        'specialite': getattr(consultant, 'specialite', ''),
                        'expertise': getattr(consultant, 'expertise', 'D√©butant'),
                        'statut': getattr(consultant, 'statut', 'En_attente'),
                        'is_validated': getattr(consultant, 'is_validated', False),
                        'skills': '',
                        
                        # Fichiers
                        'cv': None,
                        'photo': None,
                        
                        'created_at': consultant.created_at.isoformat() if consultant.created_at else None,
                        'updated_at': consultant.updated_at.isoformat() if consultant.updated_at else None,
                        
                        # Valeurs par d√©faut pour les nouveaux champs
                        'annees_experience': 0,
                        'formation_niveau': 'BAC+3',
                        'certifications_count': 0,
                        'projets_realises': 0,
                        'leadership_experience': False,
                        'international_experience': False,
                        'expertise_score': None,
                        'cvFilename': None,
                        'standardizedCvFilename': None,
                        'profileImage': None,
                    }
                    
                    # Gestion s√©curis√©e des fichiers
                    try:
                        if hasattr(consultant, 'cv') and consultant.cv:
                            consultant_data['cv'] = consultant.cv.url
                    except Exception:
                        pass
                        
                    try:
                        if hasattr(consultant, 'photo') and consultant.photo:
                            consultant_data['photo'] = consultant.photo.url
                    except Exception:
                        pass
                    
                    # R√©cup√©rer les comp√©tences
                    try:
                        competences = Competence.objects.filter(consultant=consultant)
                        skills_list = [comp.nom_competence for comp in competences]
                        consultant_data['skills'] = ', '.join(skills_list)
                    except Exception:
                        consultant_data['skills'] = ''
                    
                    consultants_data.append(consultant_data)
                    
                except Exception as e:
                    logger.error(f"Erreur traitement consultant {consultant.id}: {str(e)}")
                    continue
            
            return Response({"success": True, "data": consultants_data})

        elif request.method == 'POST':
            # Pour la cr√©ation, utiliser le serializer normal
            serializer = ConsultantSerializer(data=request.data)
            if serializer.is_valid():
                consultant = serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
            
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des consultants: {str(e)}")
        return Response({"success": False, "error": str(e)}, status=500)
    """
    SOLUTION FINALE - Liste ou cr√©e des consultants (acc√®s admin) - VERSION CORRIG√âE
    """
    try:
        if request.method == 'GET':
            from django.db import connection
            
            # V√©rifier les champs existants
            with connection.cursor() as cursor:
                cursor.execute("DESCRIBE consultants_consultant")
                existing_fields = {row[0] for row in cursor.fetchall()}
            
            logger.info(f"Champs existants dans la table: {existing_fields}")
            
            # Champs de base qui doivent exister
            base_fields = [
                'id', 'nom', 'prenom', 'email', 'telephone', 'pays', 'ville',
                'date_debut_dispo', 'date_fin_dispo', 'domaine_principal', 
                'specialite', 'expertise', 'statut', 'is_validated', 
                'created_at', 'updated_at', 'cv', 'photo'
            ]
            
            # Filtrer seulement les champs qui existent
            valid_fields = [field for field in base_fields if field in existing_fields]
            
            # R√©cup√©rer seulement les consultants valid√©s avec les champs existants
            consultants = Consultant.objects.filter(is_validated=True).only(*valid_fields)
            
            consultants_data = []
            for consultant in consultants:
                try:
                    consultant_data = {
                        'id': consultant.id,
                        'nom': getattr(consultant, 'nom', ''),
                        'prenom': getattr(consultant, 'prenom', ''),
                        'firstName': getattr(consultant, 'prenom', ''),
                        'lastName': getattr(consultant, 'nom', ''),
                        'email': getattr(consultant, 'email', ''),
                        'telephone': getattr(consultant, 'telephone', ''),
                        'phone': getattr(consultant, 'telephone', ''),
                        'pays': getattr(consultant, 'pays', ''),
                        'country': getattr(consultant, 'pays', ''),
                        'ville': getattr(consultant, 'ville', ''),
                        'city': getattr(consultant, 'ville', ''),
                        'date_debut_dispo': consultant.date_debut_dispo.isoformat() if consultant.date_debut_dispo else None,
                        'date_fin_dispo': consultant.date_fin_dispo.isoformat() if consultant.date_fin_dispo else None,
                        'startAvailability': consultant.date_debut_dispo.isoformat() if consultant.date_debut_dispo else None,
                        'endAvailability': consultant.date_fin_dispo.isoformat() if consultant.date_fin_dispo else None,
                        'domaine_principal': getattr(consultant, 'domaine_principal', 'DIGITAL'),
                        'specialite': getattr(consultant, 'specialite', ''),
                        'expertise': getattr(consultant, 'expertise', 'D√©butant'),
                        'statut': getattr(consultant, 'statut', 'En_attente'),
                        'is_validated': getattr(consultant, 'is_validated', False),
                        'skills': '',
                        'cv': consultant.cv.url if consultant.cv else None,
                        'photo': consultant.photo.url if consultant.photo else None,
                        'created_at': consultant.created_at.isoformat() if consultant.created_at else None,
                        'updated_at': consultant.updated_at.isoformat() if consultant.updated_at else None,
                        
                        # Valeurs par d√©faut pour les nouveaux champs
                        'annees_experience': 0,
                        'formation_niveau': 'BAC+3',
                        'certifications_count': 0,
                        'projets_realises': 0,
                        'leadership_experience': False,
                        'international_experience': False,
                        'expertise_score': None,
                        'cvFilename': None,
                        'standardizedCvFilename': None,
                        'profileImage': None,
                    }
                    
                    # R√©cup√©rer les comp√©tences
                    try:
                        competences = Competence.objects.filter(consultant=consultant)
                        skills_list = [comp.nom_competence for comp in competences]
                        consultant_data['skills'] = ', '.join(skills_list)
                    except Exception:
                        consultant_data['skills'] = ''
                    
                    consultants_data.append(consultant_data)
                    
                except Exception as e:
                    logger.error(f"Erreur traitement consultant {consultant.id}: {str(e)}")
                    continue
            
            return Response({"success": True, "data": consultants_data})

        elif request.method == 'POST':
            # Pour la cr√©ation, utiliser le serializer normal
            serializer = ConsultantSerializer(data=request.data)
            if serializer.is_valid():
                consultant = serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
            
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des consultants: {str(e)}")
        return Response({"success": False, "error": str(e)}, status=500)
@api_view(['GET'])
def get_pending_consultants_safe(request):
    """
    VERSION FINALE ET S√âCURIS√âE pour r√©cup√©rer les consultants en attente
    Compatible avec PendingConsultants.tsx et g√®re tous les champs possibles
    """
    try:
        logger.info("=== D√âBUT R√âCUP√âRATION CONSULTANTS EN ATTENTE ===")
        
        # R√©cup√©ration des consultants en attente sans .only() pour √©viter les erreurs
        consultants_pending = Consultant.objects.filter(is_validated=False).select_related('user')
        
        logger.info(f"Nombre de consultants en attente trouv√©s: {consultants_pending.count()}")
        
        consultants_data = []
        
        for consultant in consultants_pending:
            try:
                # Construction s√©curis√©e des donn√©es avec tous les champs possibles
                consultant_data = {
                    # IDs et identifiants
                    'id': consultant.id,
                    
                    # Noms (champs principaux + alias pour compatibilit√©)
                    'nom': safe_get_field_value(consultant, 'nom', default=''),
                    'prenom': safe_get_field_value(consultant, 'prenom', default=''),
                    'firstName': safe_get_field_value(consultant, 'firstName', 'prenom', ''),
                    'lastName': safe_get_field_value(consultant, 'lastName', 'nom', ''),
                    
                    # Contact (champs principaux + alias)
                    'email': safe_get_field_value(consultant, 'email', default=''),
                    'telephone': safe_get_field_value(consultant, 'telephone', default=''),
                    'phone': safe_get_field_value(consultant, 'phone', 'telephone', ''),
                    
                    # Localisation (champs principaux + alias)
                    'pays': safe_get_field_value(consultant, 'pays', default=''),
                    'country': safe_get_field_value(consultant, 'country', 'pays', ''),
                    'ville': safe_get_field_value(consultant, 'ville', default=''),
                    'city': safe_get_field_value(consultant, 'city', 'ville', ''),
                    
                    # Disponibilit√© (champs principaux + alias)
                    'date_debut_dispo': safe_get_date_field(consultant, 'date_debut_dispo'),
                    'date_fin_dispo': safe_get_date_field(consultant, 'date_fin_dispo'),
                    'startAvailability': safe_get_date_field(consultant, 'startAvailability', 'date_debut_dispo'),
                    'endAvailability': safe_get_date_field(consultant, 'endAvailability', 'date_fin_dispo'),
                    
                    # Expertise et domaine
                    'domaine_principal': safe_get_field_value(consultant, 'domaine_principal', default='DIGITAL'),
                    'specialite': safe_get_field_value(consultant, 'specialite', default=''),
                    'expertise': safe_get_field_value(consultant, 'expertise', default='D√©butant'),
                    
                    # Nouveaux champs d'expertise (avec valeurs par d√©faut)
                    'annees_experience': safe_get_field_value(consultant, 'annees_experience', default=0),
                    'formation_niveau': safe_get_field_value(consultant, 'formation_niveau', default='BAC+3'),
                    'certifications_count': safe_get_field_value(consultant, 'certifications_count', default=0),
                    'projets_realises': safe_get_field_value(consultant, 'projets_realises', default=0),
                    'leadership_experience': safe_get_field_value(consultant, 'leadership_experience', default=False),
                    'international_experience': safe_get_field_value(consultant, 'international_experience', default=False),
                    'expertise_score': safe_get_field_value(consultant, 'expertise_score', default=None),
                    
                    # Statut et validation
                    'statut': safe_get_field_value(consultant, 'statut', default='En_attente'),
                    'is_validated': safe_get_field_value(consultant, 'is_validated', default=False),
                    
                    # Fichiers (gestion s√©curis√©e)
                    'cv': None,
                    'photo': None,
                    'cvFilename': safe_get_field_value(consultant, 'cvFilename', default=None),
                    'standardizedCvFilename': safe_get_field_value(consultant, 'standardizedCvFilename', default=None),
                    'profileImage': safe_get_field_value(consultant, 'profileImage', default=None),
                    
                    # Comp√©tences (sera rempli plus tard)
                    'skills': '',
                    
                    # M√©tadonn√©es
                    'created_at': safe_get_date_field(consultant, 'created_at'),
                    'updated_at': safe_get_date_field(consultant, 'updated_at'),
                }
                
                # Gestion s√©curis√©e des fichiers
                try:
                    if hasattr(consultant, 'cv') and consultant.cv:
                        consultant_data['cv'] = consultant.cv.url
                        if not consultant_data['cvFilename']:
                            consultant_data['cvFilename'] = consultant.cv.name.split('/')[-1] if consultant.cv.name else None
                except Exception as cv_error:
                    logger.warning(f"Erreur CV pour consultant {consultant.id}: {cv_error}")
                    consultant_data['cv'] = None
                
                try:
                    if hasattr(consultant, 'photo') and consultant.photo:
                        consultant_data['photo'] = consultant.photo.url
                except Exception as photo_error:
                    logger.warning(f"Erreur photo pour consultant {consultant.id}: {photo_error}")
                    consultant_data['photo'] = None
                
                # R√©cup√©rer les comp√©tences
                try:
                    competences = Competence.objects.filter(consultant=consultant)
                    skills_list = [comp.nom_competence for comp in competences]
                    consultant_data['skills'] = ', '.join(skills_list) if skills_list else ''
                except Exception as skills_error:
                    logger.warning(f"Erreur comp√©tences pour consultant {consultant.id}: {skills_error}")
                    consultant_data['skills'] = ''
                
                # Synchroniser les champs alias si les principaux sont vides
                if not consultant_data['startAvailability'] and consultant_data['date_debut_dispo']:
                    consultant_data['startAvailability'] = consultant_data['date_debut_dispo']
                if not consultant_data['endAvailability'] and consultant_data['date_fin_dispo']:
                    consultant_data['endAvailability'] = consultant_data['date_fin_dispo']
                
                consultants_data.append(consultant_data)
                
                logger.debug(f"Consultant {consultant.id} trait√© avec succ√®s")
                
            except Exception as e:
                logger.error(f"Erreur traitement consultant {consultant.id}: {str(e)}")
                # Ajouter quand m√™me avec des donn√©es minimales pour √©viter une page blanche
                consultants_data.append({
                    'id': consultant.id,
                    'nom': getattr(consultant, 'nom', 'Erreur'),
                    'prenom': getattr(consultant, 'prenom', 'Chargement'),
                    'firstName': getattr(consultant, 'prenom', 'Erreur'),
                    'lastName': getattr(consultant, 'nom', 'Chargement'),
                    'email': getattr(consultant, 'email', ''),
                    'telephone': getattr(consultant, 'telephone', ''),
                    'phone': getattr(consultant, 'telephone', ''),
                    'pays': getattr(consultant, 'pays', ''),
                    'country': getattr(consultant, 'pays', ''),
                    'ville': getattr(consultant, 'ville', ''),
                    'city': getattr(consultant, 'ville', ''),
                    'domaine_principal': getattr(consultant, 'domaine_principal', 'DIGITAL'),
                    'specialite': getattr(consultant, 'specialite', ''),
                    'expertise': getattr(consultant, 'expertise', 'D√©butant'),
                    'is_validated': False,
                    'statut': 'En_attente',
                    'error': f"Erreur traitement: {str(e)}"
                })
        
        logger.info(f"=== FIN R√âCUP√âRATION ===")
        logger.info(f"Donn√©es finales pr√©par√©es: {len(consultants_data)} consultants")
        
        return Response({
            "success": True,
            "data": consultants_data,
            "count": len(consultants_data),
            "message": f"{len(consultants_data)} consultants en attente r√©cup√©r√©s"
        })
        
    except Exception as e:
        logger.error(f"‚ùå Erreur critique dans get_pending_consultants_safe: {str(e)}")
        return Response({
            "success": False,
            "error": str(e),
            "message": "Erreur lors de la r√©cup√©ration des consultants en attente",
            "data": []
        }, status=500)
        
@api_view(['GET', 'POST'])
def admin_consultants_fixed(request):
    """
    SOLUTION IMM√âDIATE - Liste ou cr√©e des consultants sans utiliser les champs manquants
    """
    try:
        if request.method == 'GET':
            from django.db import connection
            
            # V√©rifier les champs existants
            with connection.cursor() as cursor:
                cursor.execute("DESCRIBE consultants_consultant")
                existing_fields = {row[0] for row in cursor.fetchall()}
            
            # Champs de base qui doivent exister
            base_fields = [
                'id', 'nom', 'prenom', 'email', 'telephone', 'pays', 'ville',
                'date_debut_dispo', 'date_fin_dispo', 'domaine_principal', 
                'specialite', 'expertise', 'statut', 'is_validated', 
                'created_at', 'updated_at', 'cv', 'photo'
            ]
            
            # Filtrer seulement les champs qui existent
            valid_fields = [field for field in base_fields if field in existing_fields]
            
            # R√©cup√©rer seulement les consultants valid√©s avec les champs existants
            consultants = Consultant.objects.filter(is_validated=True).only(*valid_fields)
            
            consultants_data = []
            for consultant in consultants:
                try:
                    consultant_data = {
                        'id': consultant.id,
                        'nom': getattr(consultant, 'nom', ''),
                        'prenom': getattr(consultant, 'prenom', ''),
                        'firstName': getattr(consultant, 'prenom', ''),
                        'lastName': getattr(consultant, 'nom', ''),
                        'email': getattr(consultant, 'email', ''),
                        'telephone': getattr(consultant, 'telephone', ''),
                        'phone': getattr(consultant, 'telephone', ''),
                        'pays': getattr(consultant, 'pays', ''),
                        'country': getattr(consultant, 'pays', ''),
                        'ville': getattr(consultant, 'ville', ''),
                        'city': getattr(consultant, 'ville', ''),
                        'date_debut_dispo': consultant.date_debut_dispo.isoformat() if consultant.date_debut_dispo else None,
                        'date_fin_dispo': consultant.date_fin_dispo.isoformat() if consultant.date_fin_dispo else None,
                        'startAvailability': consultant.date_debut_dispo.isoformat() if consultant.date_debut_dispo else None,
                        'endAvailability': consultant.date_fin_dispo.isoformat() if consultant.date_fin_dispo else None,
                        'domaine_principal': getattr(consultant, 'domaine_principal', 'DIGITAL'),
                        'specialite': getattr(consultant, 'specialite', ''),
                        'expertise': getattr(consultant, 'expertise', 'D√©butant'),
                        'statut': getattr(consultant, 'statut', 'En_attente'),
                        'is_validated': getattr(consultant, 'is_validated', False),
                        'skills': '',
                        'cv': consultant.cv.url if consultant.cv else None,
                        'photo': consultant.photo.url if consultant.photo else None,
                        'created_at': consultant.created_at.isoformat() if consultant.created_at else None,
                        'updated_at': consultant.updated_at.isoformat() if consultant.updated_at else None,
                    }
                    
                    # R√©cup√©rer les comp√©tences
                    try:
                        competences = Competence.objects.filter(consultant=consultant)
                        skills_list = [comp.nom_competence for comp in competences]
                        consultant_data['skills'] = ', '.join(skills_list)
                    except Exception:
                        consultant_data['skills'] = ''
                    
                    consultants_data.append(consultant_data)
                    
                except Exception as e:
                    logger.error(f"Erreur traitement consultant {consultant.id}: {str(e)}")
                    continue
            
            return Response({"success": True, "data": consultants_data})

        elif request.method == 'POST':
            # Pour la cr√©ation, utiliser le serializer normal
            serializer = ConsultantSerializer(data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
            
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des consultants: {str(e)}")
        return Response({"success": False, "error": str(e)}, status=500)


# SOLUTION TEMPORAIRE : Fonction utilitaire pour v√©rifier les champs existants
def get_existing_fields():
    """Utilitaire pour v√©rifier quels champs existent dans la table consultant"""
    from django.db import connection
    
    try:
        with connection.cursor() as cursor:
            cursor.execute("DESCRIBE consultants_consultant")
            existing_fields = {row[0] for row in cursor.fetchall()}
        return existing_fields
    except Exception as e:
        logger.error(f"Erreur v√©rification champs: {e}")
        return set()
@api_view(['PUT'])
def admin_validate_consultant(request, pk):
    """
    Valide un consultant - VERSION CORRIG√âE FINALE
    """
    try:
        consultant = get_object_or_404(Consultant, pk=pk)
        
        logger.info(f"=== VALIDATION CONSULTANT {pk} ===")
        logger.info(f"Consultant: {consultant.prenom} {consultant.nom}")

        # Changer le statut de validation
        consultant.is_validated = True
        consultant.statut = 'Actif'  # Changer aussi le statut
        consultant.save(update_fields=['is_validated', 'statut'])

        # Envoi d'un email au consultant pour l'informer que son compte a √©t√© valid√©
        try:
            from .email_service import send_validation_email
            send_validation_email(consultant)
            logger.info(f"Email de validation envoy√© √† {consultant.email}")
        except Exception as e:
            logger.error(f"Erreur lors de l'envoi de l'email de validation: {str(e)}")

        logger.info(f"‚úÖ Consultant {consultant.prenom} {consultant.nom} valid√© avec succ√®s")

        return Response({
            "success": True,
            "message": f"Consultant {consultant.prenom} {consultant.nom} valid√© avec succ√®s"
        })
    except Exception as e:
        logger.error(f"‚ùå Erreur lors de la validation du consultant {pk}: {str(e)}")
        return Response({
            "success": False,
            "error": str(e)
        }, status=500)
@api_view(['DELETE'])
def admin_consultant_detail_delete(request, pk):
    """
    Supprime un consultant (rejet) - VERSION CORRIG√âE FINALE
    """
    try:
        consultant = get_object_or_404(Consultant, pk=pk)
        
        logger.info(f"=== D√âBUT SUPPRESSION CONSULTANT {pk} ===")
        logger.info(f"Consultant: {consultant.prenom} {consultant.nom}")
        
        # Stocker les informations avant suppression
        user = consultant.user
        consultant_name = f"{consultant.prenom} {consultant.nom}"
        consultant_id = consultant.id
        consultant_email = consultant.email
        
        # Utiliser une transaction pour garantir la coh√©rence
        with transaction.atomic():
            
            # 1. Supprimer les notifications li√©es
            try:
                notifications_deleted = 0
                if hasattr(consultant, '_meta') and 'Notification' in [model.__name__ for model in consultant._meta.get_fields()]:
                    from .models import Notification
                    notifications = Notification.objects.filter(consultant=consultant)
                    notifications_deleted = notifications.count()
                    notifications.delete()
                    logger.info(f"‚úÖ {notifications_deleted} notifications supprim√©es")
            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Erreur suppression notifications: {e}")
                notifications_deleted = 0

            # 2. Supprimer les matchings li√©s
            try:
                matchings_deleted = 0
                from .models import MatchingResult
                matchings = MatchingResult.objects.filter(consultant=consultant)
                matchings_deleted = matchings.count()
                matchings.delete()
                logger.info(f"‚úÖ {matchings_deleted} matchings supprim√©s")
            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Erreur suppression matchings: {e}")
                matchings_deleted = 0

            # 3. Supprimer les comp√©tences
            try:
                competences = Competence.objects.filter(consultant=consultant)
                competences_deleted = competences.count()
                competences.delete()
                logger.info(f"‚úÖ {competences_deleted} comp√©tences supprim√©es")
            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Erreur suppression comp√©tences: {e}")
                competences_deleted = 0

            # 4. Mettre √† jour les documents GED (ne pas supprimer, juste d√©lier)
            try:
                doc_ged_updated = 0
                # V√©rifier si le mod√®le DocumentGED existe
                try:
                    from .models import DocumentGED
                    documents = DocumentGED.objects.filter(consultant=consultant)
                    doc_ged_updated = documents.update(consultant=None)
                    logger.info(f"‚úÖ {doc_ged_updated} documents GED d√©li√©s")
                except ImportError:
                    logger.info("‚ÑπÔ∏è Mod√®le DocumentGED non disponible")
            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Erreur mise √† jour documents GED: {e}")
                doc_ged_updated = 0

            # 5. Supprimer les missions li√©es
            try:
                missions_deleted = 0
                # V√©rifier si le mod√®le Mission existe
                try:
                    from .models import Mission
                    missions = Mission.objects.filter(consultant=consultant)
                    missions_deleted = missions.count()
                    missions.delete()
                    logger.info(f"‚úÖ {missions_deleted} missions supprim√©es")
                except ImportError:
                    logger.info("‚ÑπÔ∏è Mod√®le Mission non disponible")
            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Erreur suppression missions: {e}")
                missions_deleted = 0

            # 6. Supprimer les fichiers physiques (CV et photo)
            files_deleted = []
            
            # Suppression du CV
            if consultant.cv and consultant.cv.name:
                try:
                    cv_path = consultant.cv.path
                    if os.path.exists(cv_path):
                        os.remove(cv_path)
                        files_deleted.append(f"CV: {cv_path}")
                        logger.info(f"‚úÖ Fichier CV supprim√©: {cv_path}")
                except Exception as e:
                    logger.warning(f"‚ö†Ô∏è Erreur suppression CV: {e}")
            
            # Suppression de la photo
            if consultant.photo and consultant.photo.name:
                try:
                    photo_path = consultant.photo.path
                    if os.path.exists(photo_path):
                        os.remove(photo_path)
                        files_deleted.append(f"Photo: {photo_path}")
                        logger.info(f"‚úÖ Fichier photo supprim√©: {photo_path}")
                except Exception as e:
                    logger.warning(f"‚ö†Ô∏è Erreur suppression photo: {e}")

            # 7. Supprimer les CV standardis√©s/Richat
            try:
                cv_standardises_deleted = 0
                cv_dir = os.path.join(settings.MEDIA_ROOT, 'standardized_cvs')
                if os.path.exists(cv_dir):
                    # Patterns de recherche pour les CV de ce consultant
                    patterns = [
                        f"standardized_cv_{consultant_id}_",
                        f"CV_Richat_{consultant.prenom}_{consultant.nom}",
                        f"CV_Richat_{consultant.nom}_{consultant.prenom}"
                    ]
                    
                    for filename in os.listdir(cv_dir):
                        if filename.endswith('.pdf'):
                            if any(pattern in filename for pattern in patterns if pattern):
                                try:
                                    file_path = os.path.join(cv_dir, filename)
                                    os.remove(file_path)
                                    cv_standardises_deleted += 1
                                    files_deleted.append(f"CV standardis√©: {filename}")
                                    logger.info(f"‚úÖ CV standardis√© supprim√©: {filename}")
                                except Exception as e:
                                    logger.warning(f"‚ö†Ô∏è Erreur suppression CV standardis√© {filename}: {e}")
                
                logger.info(f"‚úÖ {cv_standardises_deleted} CV standardis√©s supprim√©s")
            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Erreur suppression CV standardis√©s: {e}")
                cv_standardises_deleted = 0

            # 8. Supprimer le consultant de la base de donn√©es
            try:
                consultant.delete()
                logger.info(f"‚úÖ Consultant {consultant_id} supprim√© de la base de donn√©es")
            except Exception as e:
                logger.error(f"‚ùå Erreur critique suppression consultant: {e}")
                raise e

            # 9. Supprimer l'utilisateur associ√© (si existe et n'est li√© √† rien d'autre)
            user_deleted = False
            if user:
                try:
                    user_id = user.id
                    user_username = user.username
                    
                    # V√©rifier si l'utilisateur a d'autres relations
                    has_other_relations = False
                    
                    # V√©rifier s'il y a d'autres consultants avec ce user
                    other_consultants = Consultant.objects.filter(user=user).exists()
                    if other_consultants:
                        has_other_relations = True
                        logger.info(f"‚ÑπÔ∏è Utilisateur {user_id} a d'autres consultants, conservation")
                    
                    # Si pas d'autres relations, supprimer l'utilisateur
                    if not has_other_relations:
                        user.delete()
                        user_deleted = True
                        logger.info(f"‚úÖ Utilisateur {user_id} ({user_username}) supprim√©")
                    else:
                        logger.info(f"‚ÑπÔ∏è Utilisateur {user_id} conserv√© (autres relations)")
                        
                except Exception as e:
                    logger.error(f"‚ùå Erreur suppression utilisateur: {e}")
                    # Ne pas faire √©chouer toute la transaction pour l'utilisateur
                    user_deleted = False

        # 10. R√©sum√© des suppressions
        deletion_summary = {
            'consultant_id': consultant_id,
            'consultant_name': consultant_name,
            'consultant_email': consultant_email,
            'notifications_deleted': notifications_deleted,
            'matchings_deleted': matchings_deleted,
            'competences_deleted': competences_deleted,
            'missions_deleted': missions_deleted,
            'documents_ged_unlinked': doc_ged_updated,
            'files_deleted': files_deleted,
            'cv_standardises_deleted': cv_standardises_deleted,
            'user_deleted': user_deleted,
            'total_files_deleted': len(files_deleted)
        }

        logger.info(f"‚úÖ SUPPRESSION TERMIN√âE AVEC SUCC√àS")
        logger.info(f"üìä R√©sum√©: {deletion_summary}")

        return Response({
            'success': True,
            'message': f'Consultant "{consultant_name}" supprim√© avec succ√®s',
            'deletion_summary': deletion_summary
        }, status=200)

    except Exception as e:
        logger.error(f"‚ùå ERREUR CRITIQUE lors de la suppression du consultant {pk}: {str(e)}")
        return Response({
            'success': False,
            'error': f'Erreur lors de la suppression: {str(e)}',
            'consultant_id': pk
        }, status=500)
@api_view(['GET'])
def consultant_notifications(request, consultant_id):
    """
    R√©cup√®re les notifications d'un consultant - VERSION CORRIG√âE COMPL√àTE
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        
        logger.info(f"üîî R√©cup√©ration des notifications pour consultant {consultant_id}")

        # R√©cup√©rer toutes les notifications avec relations ET tri par priorit√© et date
        notifications = Notification.objects.filter(
            consultant=consultant
        ).select_related('related_appel', 'related_match', 'related_mission').order_by(
            '-priority',  # D'abord par priorit√© (HIGH, NORMAL, LOW)
            '-created_at'  # Puis par date (plus r√©cent en premier)
        )

        logger.info(f"üìä Notifications trouv√©es: {notifications.count()}")

        # Formater pour l'API avec plus de d√©tails
        results = []
        for notif in notifications:
            try:
                notification_data = {
                    'id': notif.id,
                    'type': notif.type,
                    'title': notif.title,
                    'content': notif.content,
                    'priority': getattr(notif, 'priority', 'NORMAL'),
                    'is_read': notif.is_read,
                    'created_at': notif.created_at.isoformat(),
                    'read_at': notif.read_at.isoformat() if notif.read_at else None,
                    
                    # Informations sur l'appel d'offre
                    'appel_offre_id': notif.related_appel.id if notif.related_appel else None,
                    'appel_offre_nom': getattr(notif.related_appel, 'titre', None) if notif.related_appel else None,
                    'client': getattr(notif.related_appel, 'client', None) if notif.related_appel else None,
                    
                    # Informations sur le matching
                    'match_id': notif.related_match.id if notif.related_match else None,
                    'match_score': float(notif.related_match.score) if notif.related_match else None,
                    
                    # Informations sur la mission
                    'mission_id': notif.related_mission.id if notif.related_mission else None,
                    'mission_title': getattr(notif.related_mission, 'nom_projet', None) if notif.related_mission else None,
                    
                    # M√©tadonn√©es additionnelles
                    'metadata': getattr(notif, 'metadata', {}),
                    
                    # √Çge de la notification
                    'age_days': (timezone.now() - notif.created_at).days
                }
                
                results.append(notification_data)
                logger.debug(f"‚úÖ Notification format√©e: {notif.title}")
                
            except Exception as notif_error:
                logger.error(f"‚ùå Erreur formatage notification {notif.id}: {str(notif_error)}")
                continue

        # Compter les notifications par type et priorit√©
        unread_count = notifications.filter(is_read=False).count()
        priority_counts = {
            'HIGH': notifications.filter(priority='HIGH', is_read=False).count(),
            'NORMAL': notifications.filter(priority='NORMAL', is_read=False).count(),
            'LOW': notifications.filter(priority='LOW', is_read=False).count(),
        }
        
        logger.info(f"‚úÖ {len(results)} notifications format√©es, {unread_count} non lues")

        return Response({
            'success': True,
            'notifications': results,
            'unread_count': unread_count,
            'total_count': len(results),
            'priority_counts': priority_counts,
            'consultant_name': f"{consultant.prenom} {consultant.nom}"
        })
        
    except Exception as e:
        logger.error(f"‚ùå Erreur lors de la r√©cup√©ration des notifications: {str(e)}")
        return Response({
            'success': False,
            'error': str(e),
            'notifications': [],
            'unread_count': 0
        }, status=500)

@api_view(['GET'])
def consultant_missions(request, consultant_id):
    """
    Version simplifi√©e sans le champ created_at manquant
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        logger.info(f"üîç R√©cup√©ration des missions pour consultant {consultant_id}")

        # R√©cup√©ration des matchings valid√©s seulement (plus simple)
        validated_matches = MatchingResult.objects.filter(
            consultant=consultant,
            is_validated=True
        ).select_related('appel_offre')
        
        logger.info(f"üìä Matchings valid√©s trouv√©s: {validated_matches.count()}")

        missions = []
        
        for match in validated_matches:
            try:
                appel_offre = match.appel_offre
                
                mission_data = {
                    'id': f"match_{match.id}",
                    'appel_offre_id': appel_offre.id,
                    'nom_projet': getattr(appel_offre, 'titre', 'Projet sans nom'),
                    'client': getattr(appel_offre, 'client', 'Client non sp√©cifi√©'),
                    'description': getattr(appel_offre, 'description', ''),
                    'date_debut': appel_offre.date_de_publication.isoformat() if getattr(appel_offre, 'date_de_publication', None) else None,
                    'date_fin': appel_offre.date_limite.isoformat() if getattr(appel_offre, 'date_limite', None) else None,
                    'score': float(match.score),
                    'date_validation': match.created_at.isoformat(),
                    'statut': 'Matching valid√©',
                    'type': 'matching'
                }
                
                missions.append(mission_data)
                logger.info(f"‚úÖ Mission ajout√©e: {mission_data['nom_projet']}")
                
            except Exception as match_error:
                logger.error(f"‚ùå Erreur match {match.id}: {str(match_error)}")
                continue

        # Tri par date de validation
        missions.sort(key=lambda x: x.get('date_validation', ''), reverse=True)
        
        logger.info(f"‚úÖ Total {len(missions)} missions pour consultant {consultant_id}")

        return Response({
            "success": True,
            "missions": missions,
            "total_count": len(missions),
            "stats": {
                "direct_missions": 0,  # Temporairement 0 pour √©viter les erreurs
                "validated_matches": validated_matches.count(),
                "total": len(missions)
            }
        })
        
    except Exception as e:
        logger.error(f"‚ùå Erreur missions: {str(e)}")
        return Response({
            "success": False,
            "error": str(e),
            "missions": [],
            "total_count": 0
        }, status=500)
def create_mission_from_matching(matching_result):
    """
    Cr√©e automatiquement une mission √† partir d'un matching valid√©
    VERSION CORRIG√âE avec notifications
    """
    try:
        from .models import Mission, Notification
        from django.utils import timezone
        from datetime import timedelta
        
        # V√©rifier si une mission n'existe pas d√©j√†
        existing_mission = Mission.objects.filter(
            consultant=matching_result.consultant,
            appel_offre=matching_result.appel_offre
        ).first()
        
        if existing_mission:
            logger.info(f"Mission d√©j√† existante pour ce matching: {existing_mission.id}")
            return existing_mission
        
        # üî• CR√âER LA NOUVELLE MISSION AVEC ADAPTATION AU NOUVEAU MOD√àLE
        appel_offre = matching_result.appel_offre
        
        mission = Mission.objects.create(
            appel_offre=appel_offre,
            consultant=matching_result.consultant,
            # Utiliser les nouveaux champs du mod√®le AppelOffre
            nom_projet=getattr(appel_offre, 'titre', 'Mission depuis matching'),
            client=getattr(appel_offre, 'client', 'Client non sp√©cifi√©'),
            description=getattr(appel_offre, 'description', ''),
            titre=f"Mission: {getattr(appel_offre, 'titre', 'Titre non d√©fini')}",
            # Gestion des dates avec le nouveau mod√®le
            date_debut=getattr(appel_offre, 'date_de_publication', None) or timezone.now().date(),
            date_fin=getattr(appel_offre, 'date_limite', None) or (timezone.now().date() + timedelta(days=30)),
            statut='Valid√©e',
            score=matching_result.score
        )
        
        logger.info(f"‚úÖ Mission cr√©√©e automatiquement depuis matching {matching_result.id}: {mission.id}")
        
        # üî• NOUVEAUT√â - Cr√©er une notification sp√©cifique pour le d√©but de mission
        try:
            start_notification = Notification.objects.create(
                consultant=matching_result.consultant,
                type="MISSION_START",
                title=f"üöÄ Nouvelle mission assign√©e",
                content=(
                    f"Votre mission '{mission.nom_projet}' chez {mission.client} "
                    f"a √©t√© officiellement cr√©√©e. "
                    f"Date de d√©but pr√©vue: {mission.date_debut.strftime('%d/%m/%Y') if mission.date_debut else '√Ä d√©finir'}. "
                    f"Consultez tous les d√©tails dans votre espace missions."
                ),
                priority="HIGH",
                related_appel=appel_offre,
                related_match=matching_result,
                related_mission=mission,
                metadata={
                    'mission_id': mission.id,
                    'mission_title': mission.nom_projet,
                    'client': mission.client,
                    'auto_created': True
                }
            )
            logger.info(f"‚úÖ Notification de d√©but de mission cr√©√©e: {start_notification.id}")
        except Exception as notif_error:
            logger.error(f"‚ùå Erreur notification d√©but mission: {str(notif_error)}")
        
        return mission
        
    except Exception as e:
        logger.error(f"‚ùå Erreur cr√©ation mission depuis matching: {str(e)}")
        return None


@api_view(['GET'])
def debug_consultant_missions(request, consultant_id):
    """
    Endpoint de d√©bogage pour diagnostiquer les probl√®mes de missions
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)

        # R√©cup√©rer tous les matchings pour ce consultant (valid√©s ou non)
        all_matches = MatchingResult.objects.filter(consultant=consultant)
        validated_matches = all_matches.filter(is_validated=True)
        
        # R√©cup√©rer toutes les missions directes
        from .models import Mission
        direct_missions = Mission.objects.filter(consultant=consultant)

        # Informations de debug
        debug_info = {
            "consultant_info": {
                "id": consultant.id,
                "nom": consultant.nom,
                "prenom": consultant.prenom,
                "email": consultant.email,
                "is_validated": consultant.is_validated,
                "statut": consultant.statut
            },
            "missions_directes": {
                "count": direct_missions.count(),
                "missions": [
                    {
                        "id": mission.id,
                        "nom_projet": mission.nom_projet,
                        "appel_offre_id": mission.appel_offre.id,
                        "appel_offre_titre": getattr(mission.appel_offre, 'titre', 'N/A'),
                        "statut": mission.statut,
                        "created_at": mission.created_at.isoformat()
                    } for mission in direct_missions
                ]
            },
            "matchings": {
                "total_count": all_matches.count(),
                "validated_count": validated_matches.count(),
                "all_matchings": [
                    {
                        "id": match.id,
                        "appel_offre_id": match.appel_offre.id,
                        "appel_offre_titre": getattr(match.appel_offre, 'titre', 'N/A'),
                        "is_validated": match.is_validated,
                        "score": float(match.score),
                        "created_at": match.created_at.isoformat()
                    } for match in all_matches
                ]
            },
            "api_test": {
                "url_missions": f"/api/consultant/{consultant_id}/missions/",
                "expected_response": "JSON avec success: true et liste missions"
            }
        }

        logger.info(f"Debug consultant missions - Consultant ID: {consultant_id}")
        logger.info(f"Missions directes: {direct_missions.count()}")
        logger.info(f"Matchings valid√©s: {validated_matches.count()}")

        return Response(debug_info)
        
    except Exception as e:
        logger.error(f"Erreur dans debug_consultant_missions: {str(e)}")
        return Response({"error": str(e)}, status=500)


@api_view(['GET'])
def debug_matching_status(request, consultant_id=None, appel_offre_id=None):
    """
    Endpoint de d√©bogage pour v√©rifier l'√©tat des matchings
    Accessible uniquement en mode DEBUG
    """
    if not settings.DEBUG:
        return Response({"error": "Endpoint disponible uniquement en mode DEBUG"}, status=403)

    try:
        filters = {}
        if consultant_id:
            filters['consultant_id'] = consultant_id
        if appel_offre_id:
            filters['appel_offre_id'] = appel_offre_id

        matchings = MatchingResult.objects.filter(**filters)

        results = []
        for match in matchings:
            results.append({
                'id': match.id,
                'consultant_id': match.consultant.id,
                'consultant_nom': f"{match.consultant.prenom} {match.consultant.nom}",
                'appel_offre_id': match.appel_offre.id,
                'appel_offre_nom': match.appel_offre.nom_projet,
                'score': float(match.score),
                'is_validated': match.is_validated,
                'created_at': match.created_at
            })

        return Response({
            'success': True,
            'count': len(results),
            'matchings': results
        })
    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)


# Gestion des documents GED
# views.py - Section GED corrig√©e avec gestion des erreurs

import os
import logging
from django.db.models import Q
from django.http import FileResponse, JsonResponse
from django.shortcuts import get_object_or_404
from rest_framework import status
from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from rest_framework.response import Response
from django.core.files.storage import default_storage
from django.utils import timezone
import mimetypes

from .models import (
    DocumentGED, DocumentCategory, DocumentVersion, DocumentAccess,
    AppelOffre, Consultant, Mission, Projet
)
from .serializers import DocumentGEDSerializer, DocumentCategorySerializer

logger = logging.getLogger(__name__)


@api_view(['GET', 'POST'])
@parser_classes([MultiPartParser, FormParser, JSONParser])
def documents_list(request):
    """
    Liste tous les documents ou cr√©e un nouveau document
    Version corrig√©e avec gestion des erreurs de fichiers manquants
    """
    if request.method == 'GET':
        try:
            # R√©cup√©rer les param√®tres de filtre
            document_type = request.query_params.get('document_type')
            category_id = request.query_params.get('category_id')
            search = request.query_params.get('search')
            folder_type = request.query_params.get('folder_type')
            appel_offre_id = request.query_params.get('appel_offre_id')

            # Construire la requ√™te
            documents = DocumentGED.objects.all()

            # Appliquer les filtres
            if document_type and document_type not in ['null', 'tous', '']:
                documents = documents.filter(document_type=document_type)

            if category_id and category_id not in ['null', 'tous', '']:
                documents = documents.filter(category_id=category_id)

            if folder_type and folder_type not in ['null', 'tous', '']:
                documents = documents.filter(folder_type=folder_type)

            if appel_offre_id and appel_offre_id not in ['null', 'tous', '']:
                documents = documents.filter(appel_offre_id=appel_offre_id)

            if search:
                documents = documents.filter(
                    Q(title__icontains=search) |
                    Q(description__icontains=search) |
                    Q(tags__icontains=search)
                )

            # Ordonner par date de modification d√©croissante
            documents = documents.order_by('-modified_date')

            # V√©rifier l'existence des fichiers et filtrer ceux qui sont corrompus
            valid_documents = []
            for doc in documents:
                try:
                    # V√©rifier l'existence du fichier
                    if doc.file and doc.file.name:
                        file_exists = doc.check_file_exists()
                        if file_exists or not doc.file_exists:
                            # Inclure le document m√™me si le fichier est manquant
                            # pour permettre √† l'utilisateur de voir l'erreur
                            valid_documents.append(doc)
                    else:
                        # Document sans fichier
                        valid_documents.append(doc)
                except Exception as e:
                    logger.error(f"Erreur lors de la v√©rification du document {doc.id}: {str(e)}")
                    # Inclure quand m√™me le document pour permettre la gestion de l'erreur
                    valid_documents.append(doc)

            # S√©rialiser et retourner les r√©sultats
            serializer = DocumentGEDSerializer(valid_documents, many=True)
            return Response(serializer.data)

        except Exception as e:
            logger.error(f"Erreur lors de la r√©cup√©ration des documents: {str(e)}")
            return Response({
                'error': 'Erreur lors de la r√©cup√©ration des documents',
                'detail': str(e)
            }, status=500)

    elif request.method == 'POST':
        try:
            # Cr√©er un nouveau document
            serializer = DocumentGEDSerializer(data=request.data)
            if serializer.is_valid():
                # V√©rifier si l'utilisateur est authentifi√©
                if request.user.is_authenticated:
                    document = serializer.save(created_by=request.user)
                else:
                    document = serializer.save(created_by=None)

                # Enregistrer l'acc√®s au document
                if request.user.is_authenticated:
                    DocumentAccess.objects.create(
                        document=document,
                        user=request.user,
                        action='EDIT'
                    )

                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            logger.error(f"Erreur lors de la cr√©ation du document: {str(e)}")
            return Response({
                'error': 'Erreur lors de la cr√©ation du document',
                'detail': str(e)
            }, status=500)


@api_view(['GET'])
def documents_by_appel_offre(request, appel_offre_id):
    """
    R√©cup√®re les documents li√©s √† un appel d'offre, organis√©s par type de dossier
    Version corrig√©e avec les nouveaux types de dossiers
    """
    try:
        # V√©rifier que l'appel d'offre existe
        appel_offre = get_object_or_404(AppelOffre, id=appel_offre_id)

        # R√©cup√©rer tous les documents de cet appel d'offre
        documents = DocumentGED.objects.filter(appel_offre=appel_offre)

        # Organiser les documents par type de dossier
        result = {}
        
        # D√©finir les types de dossiers selon le contexte
        # Pour les appels d'offres (A.O)
        ao_folder_types = [
            ('AO_ADMIN', 'Dossier administratif'),
            ('AO_TECHNIQUE', 'Dossier technique'),
            ('AO_FINANCE', 'Dossier financier'),
        ]
        
        # Pour les AMI
        ami_folder_types = [
            ('AMI_CONTEXTE', 'Contexte'),
            ('AMI_OUTREACH', 'Outreach'),
        ]
        
        # Dossier g√©n√©ral
        general_folder_types = [
            ('GENERAL', 'G√©n√©ral'),
        ]
        
        # Combiner tous les types
        all_folder_types = ao_folder_types + ami_folder_types + general_folder_types
        
        for folder_type, folder_label in all_folder_types:
            docs = documents.filter(folder_type=folder_type)
            
            # V√©rifier l'existence des fichiers pour chaque document
            valid_docs = []
            for doc in docs:
                try:
                    doc.check_file_exists()
                    valid_docs.append(doc)
                except Exception as e:
                    logger.warning(f"Erreur lors de la v√©rification du document {doc.id}: {str(e)}")
                    valid_docs.append(doc)  # Inclure quand m√™me pour afficher l'erreur
            
            if valid_docs:
                result[folder_type] = {
                    'label': folder_label,
                    'documents': DocumentGEDSerializer(valid_docs, many=True).data
                }

        return Response({
            'appel_offre': {
                'id': appel_offre.id,
                'nom': appel_offre.nom_projet,
                'client': appel_offre.client
            },
            'documents_by_folder': result
        })
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des documents par appel d'offre: {str(e)}")
        return Response({
            'error': 'Erreur lors de la r√©cup√©ration des documents',
            'detail': str(e)
        }, status=500)


@api_view(['GET', 'PUT', 'DELETE'])
@parser_classes([MultiPartParser, FormParser])
def document_detail(request, pk):
    """
    R√©cup√®re, modifie ou supprime un document sp√©cifique
    Version corrig√©e avec gestion des erreurs
    """
    try:
        document = get_object_or_404(DocumentGED, pk=pk)
    except Exception as e:
        return Response({
            'error': 'Document non trouv√©',
            'detail': str(e)
        }, status=404)

    if request.method == 'GET':
        try:
            # V√©rifier l'existence du fichier
            document.check_file_exists()
            
            # Enregistrer l'acc√®s au document
            if request.user.is_authenticated:
                DocumentAccess.objects.create(
                    document=document,
                    user=request.user,
                    action='VIEW'
                )

            serializer = DocumentGEDSerializer(document)
            return Response(serializer.data)
        except Exception as e:
            logger.error(f"Erreur lors de la r√©cup√©ration du document {pk}: {str(e)}")
            return Response({
                'error': 'Erreur lors de la r√©cup√©ration du document',
                'detail': str(e)
            }, status=500)

    elif request.method == 'PUT':
        try:
            # Mettre √† jour le document
            serializer = DocumentGEDSerializer(document, data=request.data, partial=True)
            if serializer.is_valid():
                updated_doc = serializer.save()

                # Si un nouveau fichier est envoy√©, cr√©er une nouvelle version
                if 'file' in request.FILES:
                    from datetime import datetime
                    version_number = f"{datetime.now().strftime('%Y%m%d')}-{document.versions.count() + 1}"

                    if request.user.is_authenticated:
                        DocumentVersion.objects.create(
                            document=document,
                            version_number=version_number,
                            file=document.file,  # Ancien fichier
                            comments=f"Version cr√©√©e lors de la mise √† jour du {datetime.now().strftime('%d/%m/%Y')}",
                            created_by=request.user
                        )
                    else:
                        DocumentVersion.objects.create(
                            document=document,
                            version_number=version_number,
                            file=document.file,  # Ancien fichier
                            comments=f"Version cr√©√©e lors de la mise √† jour du {datetime.now().strftime('%d/%m/%Y')}",
                            created_by=None
                        )

                # Enregistrer l'acc√®s au document
                if request.user.is_authenticated:
                    DocumentAccess.objects.create(
                        document=document,
                        user=request.user,
                        action='EDIT'
                    )

                return Response(serializer.data)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            logger.error(f"Erreur lors de la mise √† jour du document {pk}: {str(e)}")
            return Response({
                'error': 'Erreur lors de la mise √† jour du document',
                'detail': str(e)
            }, status=500)

    elif request.method == 'DELETE':
        try:
            # Enregistrer l'acc√®s avant la suppression
            if request.user.is_authenticated:
                DocumentAccess.objects.create(
                    document=document,
                    user=request.user,
                    action='DELETE'
                )

            # Supprimer le fichier physique s'il existe
            try:
                if document.file and document.file.name:
                    file_path = document.file.path
                    if os.path.exists(file_path):
                        os.remove(file_path)
                        logger.info(f"Fichier supprim√©: {file_path}")
            except Exception as file_error:
                logger.warning(f"Erreur lors de la suppression du fichier physique: {str(file_error)}")

            # Supprimer le document de la base de donn√©es
            document.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as e:
            logger.error(f"Erreur lors de la suppression du document {pk}: {str(e)}")
            return Response({
                'error': 'Erreur lors de la suppression du document',
                'detail': str(e)
            }, status=500)


@api_view(['GET'])
def document_download(request, pk):
    """
    T√©l√©charge un document avec gestion des erreurs
    """
    try:
        document = get_object_or_404(DocumentGED, pk=pk)

        # V√©rifier que le fichier existe
        if not document.file or not document.file.name:
            return Response({
                'error': 'Aucun fichier associ√© √† ce document'
            }, status=404)

        # V√©rifier l'existence physique du fichier
        try:
            file_path = document.file.path
            if not os.path.exists(file_path):
                logger.error(f"Fichier physique manquant: {file_path}")
                document.file_exists = False
                document.last_file_check = timezone.now()
                document.save(update_fields=['file_exists', 'last_file_check'])
                
                return Response({
                    'error': 'Le fichier physique est manquant sur le serveur',
                    'file_path': file_path
                }, status=404)
        except Exception as path_error:
            logger.error(f"Erreur lors de l'acc√®s au chemin du fichier: {str(path_error)}")
            return Response({
                'error': 'Erreur lors de l\'acc√®s au fichier',
                'detail': str(path_error)
            }, status=500)

        # Enregistrer l'acc√®s au document
        if request.user.is_authenticated:
            DocumentAccess.objects.create(
                document=document,
                user=request.user,
                action='DOWNLOAD'
            )

        # Pr√©parer le t√©l√©chargement
        file_name = os.path.basename(file_path)

        # D√©terminer le type MIME
        content_type, encoding = mimetypes.guess_type(file_path)
        content_type = content_type or 'application/octet-stream'

        # Retourner le fichier
        try:
            response = FileResponse(open(file_path, 'rb'), content_type=content_type)
            response['Content-Disposition'] = f'attachment; filename="{file_name}"'
            
            # Marquer le fichier comme existant
            document.file_exists = True
            document.save(update_fields=['file_exists'])
            
            return response
        except Exception as download_error:
            logger.error(f"Erreur lors du t√©l√©chargement: {str(download_error)}")
            return Response({
                'error': 'Erreur lors du t√©l√©chargement du fichier',
                'detail': str(download_error)
            }, status=500)

    except Exception as e:
        logger.error(f"Erreur g√©n√©rale lors du t√©l√©chargement du document {pk}: {str(e)}")
        return Response({
            'error': 'Erreur lors du t√©l√©chargement',
            'detail': str(e)
        }, status=500)


@api_view(['GET', 'POST'])
def document_categories(request):
    """
    Liste ou cr√©e des cat√©gories de documents
    """
    try:
        if request.method == 'GET':
            categories = DocumentCategory.objects.all().order_by('name')
            serializer = DocumentCategorySerializer(categories, many=True)
            return Response(serializer.data)

        elif request.method == 'POST':
            serializer = DocumentCategorySerializer(data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    except Exception as e:
        logger.error(f"Erreur dans document_categories: {str(e)}")
        return Response({
            'error': 'Erreur lors de la gestion des cat√©gories',
            'detail': str(e)
        }, status=500)


@api_view(['PUT', 'DELETE'])
def document_category_detail(request, pk):
    """
    Modifie ou supprime une cat√©gorie
    """
    try:
        category = get_object_or_404(DocumentCategory, pk=pk)

        if request.method == 'PUT':
            serializer = DocumentCategorySerializer(category, data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        elif request.method == 'DELETE':
            category.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
    except Exception as e:
        logger.error(f"Erreur dans document_category_detail: {str(e)}")
        return Response({
            'error': 'Erreur lors de la gestion de la cat√©gorie',
            'detail': str(e)
        }, status=500)


@api_view(['GET'])
def document_versions(request, document_id):
    """
    R√©cup√®re les versions d'un document
    """
    try:
        document = get_object_or_404(DocumentGED, pk=document_id)
        versions = document.versions.all().order_by('-created_at')

        data = []
        for version in versions:
            data.append({
                'id': version.id,
                'version_number': version.version_number,
                'created_by': version.created_by.username if version.created_by else "Inconnu",
                'created_at': version.created_at,
                'comments': version.comments,
                'file_url': request.build_absolute_uri(version.file.url) if version.file else None
            })

        return Response(data)
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des versions: {str(e)}")
        return Response({
            'error': 'Erreur lors de la r√©cup√©ration des versions',
            'detail': str(e)
        }, status=500)


@api_view(['GET'])
def document_stats(request):
    """
    Statistiques sur les documents dans la GED
    """
    try:
        total_documents = DocumentGED.objects.count()
        documents_by_type = {}

        for doc_type, label in DocumentGED.DOCUMENT_TYPES:
            count = DocumentGED.objects.filter(document_type=doc_type).count()
            documents_by_type[doc_type] = {
                'label': label,
                'count': count
            }

        recent_documents = DocumentGED.objects.order_by('-upload_date')[:5]
        
        # V√©rifier l'existence des fichiers pour les documents r√©cents
        valid_recent_docs = []
        for doc in recent_documents:
            try:
                doc.check_file_exists()
                valid_recent_docs.append(doc)
            except Exception as e:
                logger.warning(f"Erreur lors de la v√©rification du document r√©cent {doc.id}: {str(e)}")
                valid_recent_docs.append(doc)  # Inclure quand m√™me
        
        recent_docs_data = DocumentGEDSerializer(valid_recent_docs, many=True).data

        return Response({
            'total_documents': total_documents,
            'documents_by_type': documents_by_type,
            'recent_documents': recent_docs_data
        })
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des statistiques: {str(e)}")
        return Response({
            'error': 'Erreur lors de la r√©cup√©ration des statistiques',
            'detail': str(e)
        }, status=500)


@api_view(['POST'])
def cleanup_missing_files(request):
    """
    Endpoint utilitaire pour nettoyer les documents avec des fichiers manquants
    """
    try:
        if not request.user.is_authenticated or not request.user.is_staff:
            return Response({
                'error': 'Permission refus√©e'
            }, status=403)

        missing_files = []
        fixed_documents = []
        
        documents = DocumentGED.objects.all()
        
        for doc in documents:
            try:
                if doc.file and doc.file.name:
                    file_path = doc.file.path
                    if not os.path.exists(file_path):
                        missing_files.append({
                            'id': doc.id,
                            'title': doc.title,
                            'file_path': file_path
                        })
                        
                        # Marquer comme manquant
                        doc.file_exists = False
                        doc.last_file_check = timezone.now()
                        doc.save(update_fields=['file_exists', 'last_file_check'])
                        fixed_documents.append(doc.id)
                    else:
                        # Marquer comme existant
                        if not doc.file_exists:
                            doc.file_exists = True
                            doc.save(update_fields=['file_exists'])
                            fixed_documents.append(doc.id)
            except Exception as e:
                logger.error(f"Erreur lors de la v√©rification du document {doc.id}: {str(e)}")

        return Response({
            'success': True,
            'missing_files_count': len(missing_files),
            'missing_files': missing_files,
            'fixed_documents': fixed_documents
        })
    except Exception as e:
        logger.error(f"Erreur lors du nettoyage: {str(e)}")
        return Response({
            'error': 'Erreur lors du nettoyage',
            'detail': str(e)
        }, status=500)
    """
    Statistiques sur les documents dans la GED
    """
    total_documents = DocumentGED.objects.count()
    documents_by_type = {}

    for doc_type, label in DocumentGED.DOCUMENT_TYPES:
        count = DocumentGED.objects.filter(document_type=doc_type).count()
        documents_by_type[doc_type] = {
            'label': label,
            'count': count
        }

    recent_documents = DocumentGED.objects.order_by('-upload_date')[:5]
    recent_docs_data = DocumentGEDSerializer(recent_documents, many=True).data

    return Response({
        'total_documents': total_documents,
        'documents_by_type': documents_by_type,
        'recent_documents': recent_docs_data
    })
    
# Ajoutez cette fonction de d√©bogage
def analyze_matching_performance():
    """
    Analyse la performance des algorithmes de matching
    en comparant les pr√©dictions avec les validations manuelles.
    Cette fonction peut √™tre ex√©cut√©e p√©riodiquement pour ajuster les param√®tres.
    """
    try:
        # R√©cup√©rer tous les matchings avec leurs statuts de validation
        matchings = MatchingResult.objects.all()
        
        # S√©parer les matchings valid√©s et non valid√©s
        validated = matchings.filter(is_validated=True)
        non_validated = matchings.filter(is_validated=False)
        
        # Si pas assez de donn√©es pour l'analyse, sortir
        if validated.count() < 5:
            logger.info("Pas assez de matchings valid√©s pour analyser la performance")
            return {
                'success': False,
                'message': "Donn√©es insuffisantes pour l'analyse"
            }
        
        # Analyse des scores moyens
        avg_validated_score = validated.aggregate(models.Avg('score'))['score__avg']
        avg_non_validated_score = non_validated.aggregate(models.Avg('score'))['score__avg']
        
        # R√©cup√©rer les appels d'offre avec au moins un matching valid√©
        appels_with_validated = AppelOffre.objects.filter(
            matchings__is_validated=True
        ).distinct()
        
        performance_by_domain = {}
        performance_by_expertise = {}
        
        # Analyse par domaine
        for domain, _ in Consultant.SPECIALITES_CHOICES:
            domain_validated = validated.filter(consultant__domaine_principal=domain).count()
            domain_total = matchings.filter(consultant__domaine_principal=domain).count()
            
            if domain_total > 0:
                domain_ratio = domain_validated / domain_total
                performance_by_domain[domain] = {
                    'validated': domain_validated,
                    'total': domain_total,
                    'ratio': domain_ratio
                }
        
        # Analyse par niveau d'expertise
        expertise_levels = ['D√©butant', 'Interm√©diaire', 'Expert']
        for level in expertise_levels:
            level_validated = validated.filter(consultant__expertise=level).count()
            level_total = matchings.filter(consultant__expertise=level).count()
            
            if level_total > 0:
                level_ratio = level_validated / level_total
                performance_by_expertise[level] = {
                    'validated': level_validated,
                    'total': level_total,
                    'ratio': level_ratio
                }
        
        # Analyse d√©taill√©e des matchings valid√©s vs. non valid√©s
        detailed_analysis = []
        
        for appel in appels_with_validated:
            appel_validated = validated.filter(appel_offre=appel)
            appel_non_validated = non_validated.filter(appel_offre=appel)
            
            # Score moyen des valid√©s pour cet appel
            avg_score_validated = appel_validated.aggregate(models.Avg('score'))['score__avg'] or 0
            
            # Comparer les comp√©tences des consultants valid√©s vs. non valid√©s
            validated_consultants = Consultant.objects.filter(id__in=appel_validated.values_list('consultant_id', flat=True))
            non_validated_consultants = Consultant.objects.filter(id__in=appel_non_validated.values_list('consultant_id', flat=True))
            
            # Comp√©tences communes chez les valid√©s
            common_skills = {}
            
            for consultant in validated_consultants:
                skills = Competence.objects.filter(consultant=consultant).values_list('nom_competence', flat=True)
                for skill in skills:
                    skill_lower = skill.lower()
                    if skill_lower in common_skills:
                        common_skills[skill_lower] += 1
                    else:
                        common_skills[skill_lower] = 1
            
            # Trouver les comp√©tences les plus fr√©quentes chez les valid√©s
            top_skills = sorted(common_skills.items(), key=lambda x: x[1], reverse=True)[:5]
            
            detailed_analysis.append({
                'appel_offre_id': appel.id,
                'appel_offre_name': appel.nom_projet,
                'validated_count': appel_validated.count(),
                'non_validated_count': appel_non_validated.count(),
                'avg_score_validated': avg_score_validated,
                'top_skills': [skill for skill, count in top_skills]
            })
        
        # Recommandations bas√©es sur l'analyse
        recommendations = []
        
        # Si l'√©cart entre les scores valid√©s et non valid√©s est faible
        if avg_validated_score and avg_non_validated_score:
            score_diff = avg_validated_score - avg_non_validated_score
            
            if score_diff < 10:
                recommendations.append(
                    "L'√©cart entre les scores des matchings valid√©s et non valid√©s est faible. "
                    "Il est recommand√© d'ajuster les pond√©rations de l'algorithme de matching."
                )
        
        # Recommandations par domaine
        for domain, perf in performance_by_domain.items():
            if perf['ratio'] < 0.2 and perf['total'] > 10:
                domain_name = dict(Consultant.SPECIALITES_CHOICES)[domain]
                recommendations.append(
                    f"Faible taux de validation pour le domaine {domain_name} ({perf['ratio']*100:.1f}%). "
                    "Les crit√®res sp√©cifiques √† ce domaine pourraient n√©cessiter une r√©vision."
                )
        
        # Recommandations par expertise
        for level, perf in performance_by_expertise.items():
            if perf['ratio'] < 0.2 and perf['total'] > 10:
                recommendations.append(
                    f"Faible taux de validation pour les consultants de niveau {level} ({perf['ratio']*100:.1f}%). "
                    "Ajustez potentiellement l'impact du niveau d'expertise dans le calcul du score."
                )
        
        result = {
            'success': True,
            'stats': {
                'total_matchings': matchings.count(),
                'validated_matchings': validated.count(),
                'validation_rate': validated.count() / matchings.count() if matchings.count() > 0 else 0,
                'avg_validated_score': avg_validated_score,
                'avg_non_validated_score': avg_non_validated_score
            },
            'performance_by_domain': performance_by_domain,
            'performance_by_expertise': performance_by_expertise,
            'detailed_analysis': detailed_analysis,
            'recommendations': recommendations
        }
        
        logger.info(f"Analyse de performance du matching compl√©t√©e: {len(recommendations)} recommandations g√©n√©r√©es")
        return result
    except Exception as e:
        logger.error(f"Erreur lors de l'analyse de performance: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }
def analyze_matching_performance():
    """
    Analyse la performance des algorithmes de matching
    en comparant les pr√©dictions avec les validations manuelles.
    Cette fonction peut √™tre ex√©cut√©e p√©riodiquement pour ajuster les param√®tres.
    """
    try:
        # R√©cup√©rer tous les matchings avec leurs statuts de validation
        matchings = MatchingResult.objects.all()
        
        # S√©parer les matchings valid√©s et non valid√©s
        validated = matchings.filter(is_validated=True)
        non_validated = matchings.filter(is_validated=False)
        
        # Si pas assez de donn√©es pour l'analyse, sortir
        if validated.count() < 5:
            logger.info("Pas assez de matchings valid√©s pour analyser la performance")
            return {
                'success': False,
                'message': "Donn√©es insuffisantes pour l'analyse"
            }
        
        # Analyse des scores moyens
        avg_validated_score = validated.aggregate(models.Avg('score'))['score__avg']
        avg_non_validated_score = non_validated.aggregate(models.Avg('score'))['score__avg']
        
        # R√©cup√©rer les appels d'offre avec au moins un matching valid√©
        appels_with_validated = AppelOffre.objects.filter(
            matchings__is_validated=True
        ).distinct()
        
        performance_by_domain = {}
        performance_by_expertise = {}
        
        # Analyse par domaine
        for domain, _ in Consultant.SPECIALITES_CHOICES:
            domain_validated = validated.filter(consultant__domaine_principal=domain).count()
            domain_total = matchings.filter(consultant__domaine_principal=domain).count()
            
            if domain_total > 0:
                domain_ratio = domain_validated / domain_total
                performance_by_domain[domain] = {
                    'validated': domain_validated,
                    'total': domain_total,
                    'ratio': domain_ratio
                }
        
        # Analyse par niveau d'expertise
        expertise_levels = ['D√©butant', 'Interm√©diaire', 'Expert']
        for level in expertise_levels:
            level_validated = validated.filter(consultant__expertise=level).count()
            level_total = matchings.filter(consultant__expertise=level).count()
            
            if level_total > 0:
                level_ratio = level_validated / level_total
                performance_by_expertise[level] = {
                    'validated': level_validated,
                    'total': level_total,
                    'ratio': level_ratio
                }
        
        # Analyse d√©taill√©e des matchings valid√©s vs. non valid√©s
        detailed_analysis = []
        
        for appel in appels_with_validated:
            appel_validated = validated.filter(appel_offre=appel)
            appel_non_validated = non_validated.filter(appel_offre=appel)
            
            # Score moyen des valid√©s pour cet appel
            avg_score_validated = appel_validated.aggregate(models.Avg('score'))['score__avg'] or 0
            
            # Comparer les comp√©tences des consultants valid√©s vs. non valid√©s
            validated_consultants = Consultant.objects.filter(id__in=appel_validated.values_list('consultant_id', flat=True))
            non_validated_consultants = Consultant.objects.filter(id__in=appel_non_validated.values_list('consultant_id', flat=True))
            
            # Comp√©tences communes chez les valid√©s
            common_skills = {}
            
            for consultant in validated_consultants:
                skills = Competence.objects.filter(consultant=consultant).values_list('nom_competence', flat=True)
                for skill in skills:
                    skill_lower = skill.lower()
                    if skill_lower in common_skills:
                        common_skills[skill_lower] += 1
                    else:
                        common_skills[skill_lower] = 1
            
            # Trouver les comp√©tences les plus fr√©quentes chez les valid√©s
            top_skills = sorted(common_skills.items(), key=lambda x: x[1], reverse=True)[:5]
            
            detailed_analysis.append({
                'appel_offre_id': appel.id,
                'appel_offre_name': appel.nom_projet,
                'validated_count': appel_validated.count(),
                'non_validated_count': appel_non_validated.count(),
                'avg_score_validated': avg_score_validated,
                'top_skills': [skill for skill, count in top_skills]
            })
        
        # Recommandations bas√©es sur l'analyse
        recommendations = []
        
        # Si l'√©cart entre les scores valid√©s et non valid√©s est faible
        if avg_validated_score and avg_non_validated_score:
            score_diff = avg_validated_score - avg_non_validated_score
            
            if score_diff < 10:
                recommendations.append(
                    "L'√©cart entre les scores des matchings valid√©s et non valid√©s est faible. "
                    "Il est recommand√© d'ajuster les pond√©rations de l'algorithme de matching."
                )
        
        # Recommandations par domaine
        for domain, perf in performance_by_domain.items():
            if perf['ratio'] < 0.2 and perf['total'] > 10:
                domain_name = dict(Consultant.SPECIALITES_CHOICES)[domain]
                recommendations.append(
                    f"Faible taux de validation pour le domaine {domain_name} ({perf['ratio']*100:.1f}%). "
                    "Les crit√®res sp√©cifiques √† ce domaine pourraient n√©cessiter une r√©vision."
                )
        
        # Recommandations par expertise
        for level, perf in performance_by_expertise.items():
            if perf['ratio'] < 0.2 and perf['total'] > 10:
                recommendations.append(
                    f"Faible taux de validation pour les consultants de niveau {level} ({perf['ratio']*100:.1f}%). "
                    "Ajustez potentiellement l'impact du niveau d'expertise dans le calcul du score."
                )
        
        result = {
            'success': True,
            'stats': {
                'total_matchings': matchings.count(),
                'validated_matchings': validated.count(),
                'validation_rate': validated.count() / matchings.count() if matchings.count() > 0 else 0,
                'avg_validated_score': avg_validated_score,
                'avg_non_validated_score': avg_non_validated_score
            },
            'performance_by_domain': performance_by_domain,
            'performance_by_expertise': performance_by_expertise,
            'detailed_analysis': detailed_analysis,
            'recommendations': recommendations
        }
        
        logger.info(f"Analyse de performance du matching compl√©t√©e: {len(recommendations)} recommandations g√©n√©r√©es")
        return result
    except Exception as e:
        logger.error(f"Erreur lors de l'analyse de performance: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }
@api_view(['GET'])
def debug_skills_match(request, consultant_id, appel_offre_id):
    """
    Endpoint de d√©bogage pour analyser le calcul du score des comp√©tences
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        appel_offre = get_object_or_404(AppelOffre, id=appel_offre_id)
        
        # R√©cup√©rer les comp√©tences du consultant
        consultant_skills = list(
            Competence.objects.filter(consultant=consultant).values_list('nom_competence', 'niveau')
        )
        
        # Extraire des mots-cl√©s de la description
        description_lower = appel_offre.description.lower() if appel_offre.description else ""
        words = re.findall(r'\b(\w{4,})\b', description_lower)
        stop_words = {'dans', 'pour', 'avec', 'cette', 'votre', 'notre', 'leur', 'vous', 'nous', '√™tre', 'avoir'}
        keywords = [word for word in words if word not in stop_words and not word.isdigit()]
        
        # V√©rifier les comp√©tences par domaine
        domain_scores = {}
        domain_matches = {}
        for domain, skills_list in ALL_SKILLS.items():
            domain_scores[domain] = 0
            domain_matches[domain] = []
            
            for skill in skills_list:
                skill_lower = skill.lower()
                if skill_lower in description_lower:
                    domain_scores[domain] += 1
                    domain_matches[domain].append(skill)
        
        # Calculer le score
        date_score = calculate_date_match_score(consultant, appel_offre)
        skills_score = calculate_skills_match_score(consultant, appel_offre)
        global_score = (date_score * 0.4) + (skills_score * 0.6)
        
        return Response({
            'consultant_skills': consultant_skills,
            'appel_offre_keywords': keywords,
            'domain_scores': domain_scores,
            'domain_matches': domain_matches,
            'date_score': date_score,
            'skills_score': skills_score,
            'global_score': global_score,
            'description_length': len(description_lower)
        })
    except Exception as e:
        logger.error(f"Erreur lors du d√©bogage des comp√©tences: {str(e)}")
        return Response({'error': str(e)}, status=500)
    
    
def calculate_enhanced_match_score(consultant, appel_offre):
    """
    Version am√©lior√©e du calcul de score de matching avec:
    - Poids du domaine principal r√©duit (15% au lieu de 25%)
    - Prise en compte de l'expertise (15%)
    - Analyse s√©mantique am√©lior√©e pour les comp√©tences (70%)
    """
    # Score du domaine (15%)
    domain_score = calculate_domain_match_score(consultant, appel_offre)
    
    # Score d'expertise (15%)
    expertise_score = calculate_expertise_score(consultant)
    
    # Score de comp√©tences (70%)
    skills_score = calculate_enhanced_skills_score(consultant, appel_offre)
    
    # Score final combin√©
    combined_score = (domain_score * 0.15) + (expertise_score * 0.15) + (skills_score * 0.70)
    
    # Ajustement sigmo√Øde pour favoriser les scores √©lev√©s
    adjusted_score = sigmoid_adjustment(combined_score)
    
    return adjusted_score
def calculate_semantic_skills_score(consultant, appel_offre):
    """
    Analyse s√©mantique am√©lior√©e pour les comp√©tences
    """
    # Importer spaCy pour l'analyse s√©mantique
    try:
        import spacy
        nlp = spacy.load("fr_core_news_md")  # Mod√®le fran√ßais moyen
    except:
        # Fallback si spaCy n'est pas disponible
        return calculate_alternative_skills_score(consultant, appel_offre)
    
    # Cr√©er des documents spaCy pour analyse s√©mantique
    # R√©cup√©rer comp√©tences du consultant
    consultant_skills = list(
        Competence.objects.filter(consultant=consultant).values_list('nom_competence', flat=True)
    )
    
    # R√©cup√©rer comp√©tences demand√©es dans l'appel d'offre
    ao_skills = []
    
    # D'abord via les crit√®res explicites s'ils existent
    criteria = CriteresEvaluation.objects.filter(appel_offre=appel_offre)
    if criteria.exists():
        ao_skills = [c.nom_critere for c in criteria]
    else:
        # Sinon extraire des mots-cl√©s de la description
        if appel_offre.description:
            doc = nlp(appel_offre.description)
            # Extraire les noms et adjectifs significatifs
            ao_skills = [token.text for token in doc if token.pos_ in ["NOUN", "ADJ"] and len(token.text) > 3]
            
    # Calculer les similarit√©s entre les ensembles de comp√©tences
    max_similarities = []
    
    for ao_skill in ao_skills:
        ao_doc = nlp(ao_skill.lower())
        best_similarity = 0
        
        for consultant_skill in consultant_skills:
            cons_doc = nlp(consultant_skill.lower())
            # Calculer la similarit√© s√©mantique
            similarity = ao_doc.similarity(cons_doc)
            best_similarity = max(best_similarity, similarity)
            
        max_similarities.append(best_similarity)
    
    # Si aucune comp√©tence trouv√©e, score moyen
    if not max_similarities:
        return 35
        
    # Calculer le score moyen des meilleures correspondances
    avg_similarity = sum(max_similarities) / len(max_similarities)
    
    # Convertir en score sur 70
    final_score = avg_similarity * 70
    
    return final_score

def get_top_skills(consultant, limit=5):
    """
    R√©cup√®re les comp√©tences principales d'un consultant
    """
    try:
        top_skills = Competence.objects.filter(consultant=consultant).order_by('-niveau')[:limit]
        return [skill.nom_competence for skill in top_skills]
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des comp√©tences: {str(e)}")
        return []
    
@api_view(['PUT'])
@parser_classes([MultiPartParser, FormParser, JSONParser])
def update_consultant_profile(request, consultant_id):
    """
    Met √† jour le profil d'un consultant - VERSION CORRIG√âE FINALE
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        
        logger.info(f"D√©but mise √† jour profil consultant {consultant_id}")
        logger.info(f"Donn√©es re√ßues: {request.data}")
        
        # Pr√©parer les donn√©es pour la mise √† jour
        data = {}
        updated_fields = []
        
        # Mapping des champs frontend vers backend
        field_mapping = {
            'firstName': 'prenom',
            'lastName': 'nom',
            'email': 'email',
            'phone': 'telephone',
            'country': 'pays',
            'city': 'ville',
            'startAvailability': 'date_debut_dispo',
            'endAvailability': 'date_fin_dispo',
            'domaine_principal': 'domaine_principal',
            'specialite': 'specialite',
            'expertise': 'expertise'
        }
        
        # Traiter chaque champ
        for frontend_field, backend_field in field_mapping.items():
            if frontend_field in request.data:
                value = request.data[frontend_field]
                if value is not None and value != '':
                    data[backend_field] = value
                    updated_fields.append(backend_field)
        
        # Traiter les fichiers
        if 'photo' in request.FILES:
            data['photo'] = request.FILES['photo']
            updated_fields.append('photo')
            logger.info(f"Photo re√ßue: {request.FILES['photo'].name}")
            
        if 'cv' in request.FILES:
            data['cv'] = request.FILES['cv']
            updated_fields.append('cv')
            logger.info(f"CV re√ßu: {request.FILES['cv'].name}")
        
        logger.info(f"Champs √† mettre √† jour: {updated_fields}")
        
        # Utiliser le serializer pour la validation et la mise √† jour
        serializer = ConsultantSerializer(consultant, data=data, partial=True)
        
        if serializer.is_valid():
            try:
                updated_consultant = serializer.save()
                
                # Mettre √† jour l'utilisateur associ√© si l'email a chang√©
                if 'email' in data and consultant.user:
                    try:
                        consultant.user.email = data['email']
                        consultant.user.username = data['email']
                        consultant.user.save()
                        logger.info(f"Utilisateur mis √† jour pour consultant {consultant_id}")
                    except Exception as user_error:
                        logger.error(f"Erreur mise √† jour utilisateur: {user_error}")
                
                # R√©cup√©rer les comp√©tences pour la r√©ponse
                competences = Competence.objects.filter(consultant=updated_consultant)
                competences_list = [c.nom_competence for c in competences]
                
                logger.info(f"Consultant {consultant_id} mis √† jour avec succ√®s")
                
                # Pr√©parer la r√©ponse avec les donn√©es mises √† jour
                response_data = {
                    "firstName": updated_consultant.prenom,
                    "lastName": updated_consultant.nom,
                    "email": updated_consultant.email,
                    "phone": updated_consultant.telephone,
                    "country": updated_consultant.pays,
                    "city": updated_consultant.ville,
                    "startAvailability": updated_consultant.date_debut_dispo,
                    "endAvailability": updated_consultant.date_fin_dispo,
                    "expertise": updated_consultant.expertise,
                    "domaine_principal": updated_consultant.domaine_principal,
                    "specialite": updated_consultant.specialite,
                    "skills": ", ".join(competences_list),
                    "cvFilename": updated_consultant.cv.name.split('/')[-1] if updated_consultant.cv else None,
                    "photo": updated_consultant.photo.url if updated_consultant.photo else None
                }
                
                return Response({
                    "success": True,
                    "message": "Profil mis √† jour avec succ√®s",
                    "data": response_data
                }, status=200)
                
            except Exception as save_error:
                logger.error(f"Erreur lors de la sauvegarde du profil: {save_error}")
                return Response({
                    "success": False,
                    "error": f"Erreur lors de la sauvegarde: {str(save_error)}"
                }, status=500)
        else:
            logger.error(f"Erreurs de validation pour profil {consultant_id}: {serializer.errors}")
            return Response({
                "success": False,
                "errors": serializer.errors,
                "message": "Erreurs de validation"
            }, status=400)
        
    except Exception as e:
        logger.error(f"Erreur g√©n√©rale lors de la mise √† jour du profil consultant {consultant_id}: {str(e)}")
        return Response({
            "success": False,
            "error": f"Erreur: {str(e)}"
        }, status=500)


@api_view(['GET'])
def download_standardized_cv(request, consultant_id):
    """
    Endpoint pour t√©l√©charger le CV standardis√© d'un consultant
    """
    try:
        # Chemin du r√©pertoire standardized_cvs
        standardized_dir = os.path.join(settings.MEDIA_ROOT, "standardized_cvs")
        
        # Cr√©er le r√©pertoire s'il n'existe pas
        os.makedirs(standardized_dir, exist_ok=True)
        
        # Chercher le CV standardis√© le plus r√©cent pour ce consultant
        pattern = f"standardized_cv_{consultant_id}_*.pdf"
        
        try:
            files = default_storage.listdir("standardized_cvs")[1]  # [1] pour les fichiers
        except Exception as e:
            logger.error(f"Erreur lors de la lecture du r√©pertoire: {str(e)}")
            # V√©rifier si le r√©pertoire existe, sinon le cr√©er
            os.makedirs(os.path.join(settings.MEDIA_ROOT, "standardized_cvs"), exist_ok=True)
            files = []
        
        matching_files = [f for f in files if f.startswith(f"standardized_cv_{consultant_id}_")]
        
        if not matching_files:
            return Response({"error": "Aucun CV standardis√© trouv√© pour ce consultant"}, 
                          status=status.HTTP_404_NOT_FOUND)
        
        # Trier par date de cr√©ation (plus r√©cent en premier)
        matching_files.sort(reverse=True)
        latest_cv = matching_files[0]
        
        # Chemin complet du fichier
        file_path = os.path.join(standardized_dir, latest_cv)
        
        # V√©rifier que le fichier existe
        if not os.path.exists(file_path):
            logger.error(f"Fichier introuvable: {file_path}")
            return Response({"error": "Fichier CV standardis√© introuvable sur le serveur"}, 
                          status=status.HTTP_404_NOT_FOUND)
        
        # Ouvrir et lire le fichier
        try:
            with open(file_path, 'rb') as f:
                file_data = f.read()
        except Exception as e:
            logger.error(f"Erreur lors de la lecture du fichier: {str(e)}")
            return Response({"error": f"Erreur lors de l'acc√®s au fichier: {str(e)}"}, 
                          status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        # Renvoyer le fichier
        response = HttpResponse(file_data, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="CV_Standardis√©_{consultant_id}.pdf"'
        
        return response
        
    except Exception as e:
        logger.error(f"Erreur lors du t√©l√©chargement du CV standardis√©: {str(e)}")
        return Response({"error": f"Erreur lors du t√©l√©chargement du CV standardis√©: {str(e)}"}, 
                      status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def check_standardized_cv(request, consultant_id):
    """
    Endpoint pour v√©rifier si un CV standardis√© existe pour un consultant
    """
    try:
        # Chemin du r√©pertoire standardized_cvs
        standardized_dir = os.path.join(settings.MEDIA_ROOT, "standardized_cvs")
        
        # Cr√©er le r√©pertoire s'il n'existe pas
        os.makedirs(standardized_dir, exist_ok=True)
        
        # Chercher tous les CV standardis√©s pour ce consultant
        pattern = f"standardized_cv_{consultant_id}_*.pdf"
        matching_files = []
        
        # V√©rifier si le r√©pertoire existe et liste les fichiers
        try:
            files = default_storage.listdir("standardized_cvs")[1]  # [1] pour les fichiers
            matching_files = [f for f in files if f.startswith(f"standardized_cv_{consultant_id}_")]
        except Exception as e:
            logger.error(f"Erreur lors de la lecture du r√©pertoire: {str(e)}")
            matching_files = []
        
        if not matching_files:
            return Response({"available": False, "message": "Aucun CV standardis√© trouv√©"})
        
        # Trier par date de cr√©ation (plus r√©cent en premier)
        matching_files.sort(reverse=True)
        latest_cv = matching_files[0]
        
        # Chemin complet du fichier
        file_path = os.path.join(standardized_dir, latest_cv)
        
        # V√©rifier que le fichier existe
        if not os.path.exists(file_path):
            return Response({"available": False, "message": "Fichier CV standardis√© introuvable"})
        
        return Response({
            "available": True,
            "filename": latest_cv,
            "created_at": datetime.fromtimestamp(os.path.getctime(file_path)).strftime('%Y-%m-%d %H:%M:%S')
        })
        
    except Exception as e:
        logger.error(f"Erreur lors de la v√©rification du CV standardis√©: {str(e)}")
        return Response({"available": False, "error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    
@api_view(['POST'])
@parser_classes([MultiPartParser, FormParser])
def process_cv(request):
    try:
        cv_file = request.FILES.get('cv')
        if not cv_file:
            return Response({"error": "Aucun fichier fourni"}, status=400)

        consultant_id = request.data.get('consultant_id', 'temp')
        processor = CVProcessor(cv_file)

        if not processor.extract_text():
            return Response({"error": "Erreur lors de l'extraction du texte du CV"}, status=400)

        processor.parse_cv()
        pdf_data = processor.generate_richat_cv()  # ‚úÖ C'est la bonne m√©thode

        # Enregistrement du fichier PDF
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        filename = f"standardized_cv_{consultant_id}_{timestamp}.pdf"
        path = default_storage.save(f"standardized_cvs/{filename}", ContentFile(pdf_data))
        url = default_storage.url(path)

        return Response({
            "success": True,
            "cv_url": url,
            "filename": filename,
            "extracted_data": processor.extracted_data
        })
    except Exception as e:
        return Response({"error": str(e)}, status=500)
    
def clean_request_data(data, allowed_fields=None):
    """
    Nettoie les donn√©es de la requ√™te en supprimant les valeurs vides et nulles
    """
    cleaned = {}
    
    for key, value in data.items():
        # Ignorer les champs non autoris√©s si sp√©cifi√©s
        if allowed_fields and key not in allowed_fields:
            continue
            
        # Ignorer les valeurs vides ou nulles
        if value is not None and value != '' and value != 'null':
            cleaned[key] = value
    
    return cleaned


# Fonction utilitaire pour valider les dates
def validate_date_range(start_date, end_date, field_prefix="date"):
    """
    Valide qu'une date de fin est post√©rieure √† une date de d√©but
    """
    from datetime import datetime
    
    try:
        if isinstance(start_date, str):
            start = datetime.strptime(start_date, '%Y-%m-%d').date()
        else:
            start = start_date
            
        if isinstance(end_date, str):
            end = datetime.strptime(end_date, '%Y-%m-%d').date()
        else:
            end = end_date
            
        if start >= end:
            return {
                f'{field_prefix}_fin': f'La date de fin doit √™tre post√©rieure √† la date de d√©but'
            }
    except ValueError as e:
        return {
            f'{field_prefix}': f'Format de date invalide: {str(e)}'
        }
    
    return None


# D√©corateur pour logger les requ√™tes
def log_request(func):
    """
    D√©corateur pour logger automatiquement les requ√™tes
    """
    def wrapper(request, *args, **kwargs):
        logger.info(f"=== D√âBUT {func.__name__} ===")
        logger.info(f"M√©thode: {request.method}")
        logger.info(f"Param√®tres URL: {kwargs}")
        logger.info(f"Donn√©es: {request.data if hasattr(request, 'data') else 'Aucune'}")
        
        try:
            response = func(request, *args, **kwargs)
            logger.info(f"=== FIN {func.__name__} - Succ√®s ===")
            return response
        except Exception as e:
            logger.error(f"=== FIN {func.__name__} - Erreur: {str(e)} ===")
            raise
    
    return wrapper

# views.py - Ajout des endpoints pour le CV Richat complet

from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
from io import BytesIO
import json
from datetime import datetime

# Mod√®le de donn√©es pour le CV Richat complet
class RichatCVData:
    def __init__(self, data):
        self.titre = data.get('titre', 'Mr.')
        self.nom_expert = data.get('nom_expert', '')
        self.date_naissance = data.get('date_naissance', '')
        self.pays_residence = data.get('pays_residence', '')
        self.titre_professionnel = data.get('titre_professionnel', '')
        self.resume_profil = data.get('resume_profil', '')
        self.formations = data.get('formations', [])
        self.experiences = data.get('experiences', [])
        self.langues = data.get('langues', [])
        self.missions_reference = data.get('missions_reference', [])
        self.certifications = data.get('certifications', [])
        self.adhesions_professionnelles = data.get('adhesions_professionnelles', 'N/A')

def generate_richat_cv_pdf(cv_data, consultant_info):
    """
    G√©n√®re un PDF au format CV Richat standardis√©
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                          topMargin=0.5*inch, bottomMargin=0.5*inch,
                          leftMargin=0.5*inch, rightMargin=0.5*inch)
    
    # Styles
    styles = getSampleStyleSheet()
    
    # Styles personnalis√©s Richat
    title_style = ParagraphStyle(
        'RichatTitle',
        parent=styles['Title'],
        fontSize=16,
        textColor=colors.HexColor('#1e40af'),
        spaceAfter=20,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    header_style = ParagraphStyle(
        'RichatHeader',
        parent=styles['Heading1'],
        fontSize=14,
        textColor=colors.HexColor('#1e40af'),
        spaceAfter=12,
        spaceBefore=20,
        fontName='Helvetica-Bold'
    )
    
    subheader_style = ParagraphStyle(
        'RichatSubHeader',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=colors.HexColor('#374151'),
        spaceAfter=8,
        spaceBefore=10,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'RichatBody',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.black,
        spaceAfter=6,
        alignment=TA_JUSTIFY,
        leftIndent=0,
        fontName='Helvetica'
    )
    
    table_style = ParagraphStyle(
        'RichatTable',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.black,
        fontName='Helvetica'
    )
    
    story = []
    
    # En-t√™te avec nom et pagination
    story.append(Paragraph(f"{cv_data.nom_expert} - {cv_data.titre_professionnel}", title_style))
    story.append(Paragraph("1/12", ParagraphStyle('PageNum', fontSize=8, alignment=TA_RIGHT)))
    story.append(Spacer(1, 0.2*inch))
    
    # Titre principal
    story.append(Paragraph("CURRICULUM VITAE (CV)", header_style))
    story.append(Spacer(1, 0.1*inch))
    
    # Tableau des informations personnelles
    personal_data = [
        ['Titre', cv_data.titre],
        ['Nom de l\'expert', cv_data.nom_expert],
        ['Date de naissance :', cv_data.date_naissance or '02-01-1978'],
        ['Pays de citoyennet√©/r√©sidence', cv_data.pays_residence]
    ]
    
    personal_table = Table(personal_data, colWidths=[2*inch, 4*inch])
    personal_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    
    story.append(personal_table)
    story.append(Spacer(1, 0.2*inch))
    
    # Titre professionnel centr√©
    story.append(Paragraph(cv_data.titre_professionnel, 
                          ParagraphStyle('ProfTitle', fontSize=14, alignment=TA_CENTER, 
                                       fontName='Helvetica-Bold', spaceAfter=15)))
    
    # R√©sum√© du Profil
    story.append(Paragraph("R√©sum√© du Profil", subheader_style))
    story.append(Paragraph(cv_data.resume_profil, body_style))
    story.append(Spacer(1, 0.15*inch))
    
    # √âducation
    if cv_data.formations:
        story.append(Paragraph("√âducation :", subheader_style))
        
        # Tableau des formations
        formation_data = [['Nom √âcole/Universit√©', 'P√©riode d\'√©tude', 'Dipl√¥me obtenu | Sp√©cialisation']]
        
        for formation in cv_data.formations:
            if formation.get('nom_ecole'):
                formation_data.append([
                    formation.get('nom_ecole', ''),
                    formation.get('periode_etude', ''),
                    f"{formation.get('diplome_obtenu', '')} ({formation.get('specialisation', '')})"
                ])
        
        if len(formation_data) > 1:
            formation_table = Table(formation_data, colWidths=[2.5*inch, 1.5*inch, 2.5*inch])
            formation_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f3f4f6')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (2, 0), 'Helvetica-Bold'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            
            story.append(formation_table)
            story.append(Spacer(1, 0.15*inch))
    
    # Exp√©rience professionnelle
    if cv_data.experiences:
        story.append(Paragraph("Exp√©rience professionnelle :", subheader_style))
        
        # Tableau d'exp√©rience
        exp_data = [['P√©riode', 'Nom de l\'employeur, Titre professionnel', 'Pays', 'R√©sum√© des activit√©s men√©es dans le cadre de cette mission']]
        
        for exp in cv_data.experiences:
            if exp.get('nom_employeur'):
                exp_data.append([
                    exp.get('periode', ''),
                    f"{exp.get('nom_employeur', '')}\n{exp.get('titre_professionnel', '')}",
                    exp.get('pays', ''),
                    exp.get('activites', '')
                ])
        
        if len(exp_data) > 1:
            exp_table = Table(exp_data, colWidths=[1*inch, 2*inch, 1*inch, 2.5*inch])
            exp_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f3f4f6')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            
            story.append(exp_table)
            story.append(Spacer(1, 0.15*inch))
    
    # Adh√©sions professionnelles
    story.append(Paragraph("Adh√©sion √† des associations professionnelles et √† des publications :", subheader_style))
    story.append(Paragraph(cv_data.adhesions_professionnelles, body_style))
    story.append(Spacer(1, 0.1*inch))
    
    # Langues
    if cv_data.langues:
        story.append(Paragraph("Langues parl√©es (n'indiquez que les langues dans lesquelles vous pouvez travailler) :", subheader_style))
        
        # Tableau des langues
        langue_data = [['', 'Parler', 'Lecture', '√âditorial']]
        
        for langue in cv_data.langues:
            if langue.get('langue'):
                langue_data.append([
                    langue.get('langue', ''),
                    langue.get('parler', ''),
                    langue.get('lecture', ''),
                    langue.get('editorial', '')
                ])
        
        if len(langue_data) > 1:
            langue_table = Table(langue_data, colWidths=[1.5*inch, 1.5*inch, 1.5*inch, 1.5*inch])
            langue_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f3f4f6')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            
            story.append(langue_table)
            story.append(Spacer(1, 0.15*inch))
    
    # Ad√©quation √† la mission (Missions de r√©f√©rence)
    if cv_data.missions_reference:
        story.append(Paragraph("Ad√©quation √† la mission :", subheader_style))
        story.append(Paragraph("R√©f√©rence √† des travaux ou missions ant√©rieurs illustrant la capacit√© de l'expert √† mener √† bien les t√¢ches qui lui sont confi√©es.", body_style))
        
        for mission in cv_data.missions_reference:
            if mission.get('nom_projet'):
                # Tableau pour chaque mission
                mission_data = [
                    ['Nom du projet :', mission.get('nom_projet', '')],
                    ['Date :', mission.get('date', '')],
                    ['Soci√©t√© :', mission.get('societe', '')],
                    ['Poste occup√© :', mission.get('poste_occupe', '')],
                    ['Lieu :', mission.get('lieu', '')],
                    ['Client / Bailleur :', mission.get('client_bailleur', '')],
                    ['Br√®ve description du projet ‚Äì Objectifs du projet :', mission.get('description_projet', '')],
                    ['Type ou secteur d\'activit√© :', mission.get('type_secteur', '')],
                    ['Activit√©s et responsabilit√©s sur le projet :', mission.get('activites_responsabilites', '')]
                ]
                
                mission_table = Table(mission_data, colWidths=[2*inch, 4*inch])
                mission_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.white),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                
                story.append(mission_table)
                story.append(Spacer(1, 0.1*inch))
    
    # Certifications
    if cv_data.certifications:
        story.append(Paragraph("Certifications:", subheader_style))
        
        for cert in cv_data.certifications:
            story.append(Paragraph(f"‚Ä¢ {cert}", body_style))
        
        story.append(Spacer(1, 0.1*inch))
    
    # Construction du PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

@api_view(['POST'])
@parser_classes([JSONParser])
def generate_richat_cv_complete(request, consultant_id):
    """
    G√©n√®re un CV complet au format Richat √† partir des donn√©es saisies
    """
    try:
        # V√©rifier que le consultant existe et appartient √† l'utilisateur
        consultant = get_object_or_404(Consultant, id=consultant_id)
        
        # R√©cup√©rer les donn√©es JSON du CV
        cv_data_json = request.data
        
        if not cv_data_json:
            return Response({
                'success': False,
                'error': 'Aucune donn√©e CV fournie'
            }, status=400)
        
        # Cr√©er l'objet de donn√©es CV
        cv_data = RichatCVData(cv_data_json)
        
        # Valider les donn√©es essentielles
        if not cv_data.nom_expert:
            return Response({
                'success': False,
                'error': 'Le nom de l\'expert est requis'
            }, status=400)
        
        # Informations du consultant pour le contexte
        consultant_info = {
            'id': consultant.id,
            'email': consultant.email,
            'phone': consultant.telephone,
            'domaine_principal': consultant.domaine_principal
        }
        
        # G√©n√©rer le PDF
        pdf_buffer = generate_richat_cv_pdf(cv_data, consultant_info)
        
        # Cr√©er le nom de fichier
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"CV_Richat_{cv_data.nom_expert.replace(' ', '_')}_{timestamp}.pdf"
        
        # Sauvegarder le fichier
        try:
            # Cr√©er le r√©pertoire s'il n'existe pas
            cv_dir = os.path.join(settings.MEDIA_ROOT, 'standardized_cvs')
            os.makedirs(cv_dir, exist_ok=True)
            
            # Chemin complet du fichier
            file_path = os.path.join(cv_dir, filename)
            
            # √âcrire le PDF
            with open(file_path, 'wb') as f:
                f.write(pdf_buffer.getvalue())
            
            # Construire l'URL de t√©l√©chargement
            file_url = f"{settings.MEDIA_URL}standardized_cvs/{filename}"
            
            # Optionnel: Sauvegarder les m√©tadonn√©es du CV en base
            try:
                # Cr√©er un enregistrement de CV g√©n√©r√© si vous avez un mod√®le pour cela
                # CVRichatGenerated.objects.create(
                #     consultant=consultant,
                #     filename=filename,
                #     cv_data=json.dumps(cv_data_json),
                #     generated_at=datetime.now()
                # )
                pass
            except Exception as e:
                logger.warning(f"Impossible de sauvegarder les m√©tadonn√©es CV: {e}")
            
            logger.info(f"CV Richat g√©n√©r√© avec succ√®s pour {consultant.nom}: {filename}")
            
            return Response({
                'success': True,
                'message': 'CV Richat g√©n√©r√© avec succ√®s',
                'cv_url': file_url,
                'filename': filename,
                'file_size': len(pdf_buffer.getvalue()),
                'generated_at': datetime.now().isoformat(),
                'consultant_name': cv_data.nom_expert,
                'format': 'richat_professional_standard'
            }, status=200)
            
        except Exception as e:
            logger.error(f"Erreur lors de la sauvegarde du CV: {str(e)}")
            return Response({
                'success': False,
                'error': f'Erreur lors de la sauvegarde: {str(e)}'
            }, status=500)
    
    except Exception as e:
        logger.error(f"Erreur lors de la g√©n√©ration du CV Richat: {str(e)}")
        return Response({
            'success': False,
            'error': f'Erreur lors de la g√©n√©ration du CV: {str(e)}'
        }, status=500)

@api_view(['GET'])
def get_richat_cv_template(request, consultant_id):
    """
    Retourne un template pr√©-rempli pour le CV Richat - VERSION CORRIG√âE
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        
        # R√©cup√©rer les comp√©tences du consultant
        competences = Competence.objects.filter(consultant=consultant)
        skills_list = [comp.nom_competence for comp in competences]
        skills_string = ', '.join(skills_list) if skills_list else 'Non sp√©cifi√©es'
        
        # Construire le template avec les donn√©es existantes
        template_data = {
            'titre': 'Mr.' if consultant.nom else 'Mme.',
            'nom_expert': f"{consultant.prenom or ''} {consultant.nom or ''}".strip(),
            'date_naissance': '',
            'pays_residence': f"{consultant.pays or 'Mauritanie'} - {consultant.ville or 'Nouakchott'}",
            'titre_professionnel': consultant.specialite or "Consultant",
            'resume_profil': f"Expert en {consultant.domaine_principal or 'Digital'} avec sp√©cialisation en {consultant.specialite or 'consulting'}. " +
                           f"Consultant exp√©riment√© avec un niveau d'expertise {consultant.expertise or 'Interm√©diaire'}. " +
                           f"Comp√©tences principales: {skills_string}.",
            'formations': [{
                'nom_ecole': '',
                'periode_etude': '',
                'diplome_obtenu': '',
                'specialisation': consultant.domaine_principal or ''
            }],
            'experiences': [{
                'periode': '',
                'nom_employeur': '',
                'titre_professionnel': consultant.specialite or '',
                'pays': consultant.pays or 'Mauritanie',
                'activites': ''
            }],
            'langues': [
                {
                    'langue': 'Fran√ßais',
                    'parler': '',
                    'lecture': '',
                    'editorial': ''
                },
                {
                    'langue': 'Anglais',
                    'parler': '',
                    'lecture': '',
                    'editorial': ''
                },
                {
                    'langue': 'Arabe',
                    'parler': 'Native speaker',
                    'lecture': 'Native speaker',
                    'editorial': 'Native speaker'
                }
            ],
            'missions_reference': [{
                'nom_projet': '',
                'date': '',
                'societe': '',
                'poste_occupe': '',
                'lieu': consultant.pays or 'Mauritanie',
                'client_bailleur': '',
                'description_projet': '',
                'type_secteur': consultant.domaine_principal or '',
                'activites_responsabilites': ''
            }],
            'certifications': [],
            'adhesions_professionnelles': 'N/A'
        }
        
        return Response({
            'success': True,
            'template_data': template_data,
            'consultant_info': {
                'nom': consultant.nom,
                'prenom': consultant.prenom,
                'email': consultant.email,
                'domaine_principal': consultant.domaine_principal,
                'specialite': consultant.specialite
            }
        })
        
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration du template CV: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)


@api_view(['GET'])
def list_richat_cvs(request, consultant_id):
    """
    Liste tous les CV Richat g√©n√©r√©s pour un consultant
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        
        # Chemin du r√©pertoire des CV
        cv_dir = os.path.join(settings.MEDIA_ROOT, 'standardized_cvs')
        
        if not os.path.exists(cv_dir):
            return Response({
                'success': True,
                'cvs': [],
                'total_count': 0
            })
        
        # Rechercher tous les CV de ce consultant
        cv_files = []
        consultant_name_patterns = [
            f"CV_Richat_{consultant.prenom}_{consultant.nom}",
            f"CV_Richat_{consultant.nom}_{consultant.prenom}",
            f"standardized_cv_{consultant.id}",
            f"CV_Richat_{consultant.prenom or ''}{consultant.nom or ''}"
        ]
        
        for filename in os.listdir(cv_dir):
            if filename.endswith('.pdf'):
                # V√©rifier si le fichier correspond √† ce consultant
                matches_consultant = any(pattern.replace(' ', '_') in filename for pattern in consultant_name_patterns if pattern)
                
                if matches_consultant:
                    file_path = os.path.join(cv_dir, filename)
                    if os.path.exists(file_path):
                        file_stats = os.stat(file_path)
                        
                        cv_files.append({
                            'filename': filename,
                            'file_size': file_stats.st_size,
                            'created_at': datetime.fromtimestamp(file_stats.st_ctime).isoformat(),
                            'modified_at': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
                            'download_url': f"{settings.MEDIA_URL}standardized_cvs/{filename}",
                            'file_type': 'richat_professional_cv'
                        })
        
        # Trier par date de cr√©ation (plus r√©cent en premier)
        cv_files.sort(key=lambda x: x['created_at'], reverse=True)
        
        return Response({
            'success': True,
            'cvs': cv_files,
            'total_count': len(cv_files),
            'consultant_name': f"{consultant.prenom} {consultant.nom}"
        })
        
    except Exception as e:
        logger.error(f"Erreur lors de la liste des CV Richat: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)

# Mod√®le optionnel pour sauvegarder les m√©tadonn√©es des CV g√©n√©r√©s
# √Ä ajouter dans models.py si souhait√©

# views.py - Fonctions suppl√©mentaires pour le CV Richat complet
@api_view(['POST'])
@parser_classes([JSONParser])
def validate_richat_cv(request, consultant_id):
    """
    Valide les donn√©es du CV Richat avant g√©n√©ration
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        cv_data = request.data
        
        # Validation basique
        errors = []
        warnings = []
        score = 100
        
        # V√©rifications essentielles
        if not cv_data.get('nom_expert'):
            errors.append('Nom de l\'expert manquant')
            score -= 20
        
        if not cv_data.get('titre_professionnel'):
            errors.append('Titre professionnel manquant')
            score -= 20
            
        if not cv_data.get('resume_profil') or len(cv_data.get('resume_profil', '')) < 50:
            warnings.append('R√©sum√© du profil trop court')
            score -= 10
            
        # V√©rifier les formations
        formations = cv_data.get('formations', [])
        if not any(f.get('nom_ecole') for f in formations):
            warnings.append('Aucune formation renseign√©e')
            score -= 10
            
        # V√©rifier les exp√©riences
        experiences = cv_data.get('experiences', [])
        if not any(e.get('nom_employeur') for e in experiences):
            warnings.append('Aucune exp√©rience renseign√©e')
            score -= 15
        
        validation_result = {
            'valid': len(errors) == 0,
            'errors': errors,
            'warnings': warnings,
            'score': max(0, score),
            'recommendations': []
        }
        
        if score < 70:
            validation_result['recommendations'].append('Compl√©ter les informations manquantes')
        if len(experiences) < 2:
            validation_result['recommendations'].append('Ajouter plus d\'exp√©riences professionnelles')
            
        return Response({
            'success': True,
            'validation': validation_result
        })
        
    except Exception as e:
        logger.error(f"Erreur validation CV: {str(e)}")
        return Response({
            'success': False,
            'error': str(e),
            'validation': {
                'valid': False,
                'errors': [f'Erreur syst√®me: {str(e)}'],
                'score': 0
            }
        }, status=500)

@api_view(['GET'])
def download_specific_richat_cv(request, consultant_id, filename):
    """
    T√©l√©charge un CV Richat sp√©cifique par nom de fichier
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)

        # Chemin du fichier
        cv_dir = os.path.join(settings.MEDIA_ROOT, 'standardized_cvs')
        file_path = os.path.join(cv_dir, filename)

        # V√©rifier existence du fichier
        if not os.path.exists(file_path):
            return Response({'success': False, 'error': 'Fichier CV introuvable.'}, status=404)

        # V√©rifier appartenance au consultant
        identifiers = [
            str(consultant.id),
            consultant.nom.lower() if consultant.nom else '',
            consultant.prenom.lower() if consultant.prenom else ''
        ]
        filename_lower = filename.lower()
        if not any(identifier in filename_lower for identifier in identifiers if identifier):
            return Response({'success': False, 'error': 'Acc√®s non autoris√© √† ce fichier.'}, status=403)

        # Renvoyer le fichier
        response = FileResponse(open(file_path, 'rb'), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

    except Exception as e:
        logger.error(f"Erreur lors du t√©l√©chargement: {str(e)}")
        return Response({'success': False, 'error': 'Erreur serveur.'}, status=500)  
 
@api_view(['GET'])
def check_richat_cv_status(request, consultant_id):
    """
    V√©rifie si un CV Richat existe pour un consultant - FONCTION CORRIG√âE
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        
        # Chemin du r√©pertoire des CV standardis√©s
        cv_dir = os.path.join(settings.MEDIA_ROOT, 'standardized_cvs')
        
        # Cr√©er le r√©pertoire s'il n'existe pas
        os.makedirs(cv_dir, exist_ok=True)
        
        # Patterns de recherche pour les fichiers CV de ce consultant
        search_patterns = [
            f"CV_Richat_{consultant.prenom}_{consultant.nom}",
            f"CV_Richat_{consultant.nom}_{consultant.prenom}",
            f"standardized_cv_{consultant_id}",
            f"CV_Richat"  # Pattern plus g√©n√©ral
        ]
        
        # Chercher les fichiers CV
        found_files = []
        
        if os.path.exists(cv_dir):
            try:
                all_files = os.listdir(cv_dir)
                for filename in all_files:
                    if filename.endswith('.pdf'):
                        # V√©rifier si le fichier correspond √† ce consultant
                        filename_lower = filename.lower()
                        consultant_match = False
                        
                        # V√©rification par ID de consultant (plus fiable)
                        if f"_{consultant_id}_" in filename or f"standardized_cv_{consultant_id}" in filename:
                            consultant_match = True
                        
                        # V√©rification par nom si pr√©sent
                        elif consultant.nom and consultant.prenom:
                            name_variants = [
                                f"{consultant.prenom}_{consultant.nom}".lower(),
                                f"{consultant.nom}_{consultant.prenom}".lower(),
                                f"{consultant.prenom}{consultant.nom}".lower()
                            ]
                            if any(variant in filename_lower for variant in name_variants):
                                consultant_match = True
                        
                        if consultant_match:
                            file_path = os.path.join(cv_dir, filename)
                            if os.path.exists(file_path):
                                file_stats = os.stat(file_path)
                                found_files.append({
                                    'filename': filename,
                                    'file_size': file_stats.st_size,
                                    'created_at': datetime.fromtimestamp(file_stats.st_ctime).isoformat(),
                                    'file_path': file_path
                                })
                
            except Exception as list_error:
                logger.error(f"Erreur lors de la lecture du r√©pertoire CV: {str(list_error)}")
        
        # Trier les fichiers par date de cr√©ation (plus r√©cent en premier)
        found_files.sort(key=lambda x: x['created_at'], reverse=True)
        
        if found_files:
            latest_cv = found_files[0]
            return Response({
                'available': True,
                'filename': latest_cv['filename'],
                'created_at': latest_cv['created_at'],
                'file_size': latest_cv['file_size'],
                'download_url': f"/api/consultant/{consultant_id}/download-cv/"
            })
        else:
            return Response({
                'available': False,
                'message': 'Aucun CV Richat trouv√© pour ce consultant'
            })
            
    except Exception as e:
        logger.error(f"Erreur lors de la v√©rification du CV Richat pour consultant {consultant_id}: {str(e)}")
        return Response({
            'available': False,
            'error': str(e)
        }, status=500)


@api_view(['GET'])
def download_richat_cv(request, consultant_id):
    """
    T√©l√©charge le CV Richat le plus r√©cent d'un consultant - FONCTION CORRIG√âE
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        
        # Chemin du r√©pertoire des CV
        cv_dir = os.path.join(settings.MEDIA_ROOT, 'standardized_cvs')
        
        if not os.path.exists(cv_dir):
            return Response({
                'success': False,
                'error': 'R√©pertoire CV introuvable'
            }, status=404)
        
        # Chercher le CV le plus r√©cent pour ce consultant
        found_files = []
        
        try:
            all_files = os.listdir(cv_dir)
            for filename in all_files:
                if filename.endswith('.pdf'):
                    # V√©rifier si le fichier correspond √† ce consultant
                    filename_lower = filename.lower()
                    consultant_match = False
                    
                    # V√©rification par ID (plus fiable)
                    if f"_{consultant_id}_" in filename or f"standardized_cv_{consultant_id}" in filename:
                        consultant_match = True
                    
                    # V√©rification par nom
                    elif consultant.nom and consultant.prenom:
                        name_variants = [
                            f"{consultant.prenom}_{consultant.nom}".lower(),
                            f"{consultant.nom}_{consultant.prenom}".lower()
                        ]
                        if any(variant in filename_lower for variant in name_variants):
                            consultant_match = True
                    
                    if consultant_match:
                        file_path = os.path.join(cv_dir, filename)
                        if os.path.exists(file_path):
                            file_stats = os.stat(file_path)
                            found_files.append({
                                'filename': filename,
                                'created_at': file_stats.st_ctime,
                                'file_path': file_path
                            })
        
        except Exception as e:
            logger.error(f"Erreur lors de la recherche de fichiers CV: {str(e)}")
            return Response({
                'success': False,
                'error': 'Erreur lors de la recherche des fichiers CV'
            }, status=500)
        
        if not found_files:
            return Response({
                'success': False,
                'error': 'Aucun CV Richat trouv√© pour ce consultant'
            }, status=404)
        
        # Trier par date de cr√©ation (plus r√©cent en premier)
        found_files.sort(key=lambda x: x['created_at'], reverse=True)
        latest_cv = found_files[0]
        
        # V√©rifier que le fichier existe toujours
        if not os.path.exists(latest_cv['file_path']):
            return Response({
                'success': False,
                'error': 'Fichier CV introuvable sur le serveur'
            }, status=404)
        
        # Retourner le fichier
        try:
            response = FileResponse(
                open(latest_cv['file_path'], 'rb'),
                content_type='application/pdf'
            )
            response['Content-Disposition'] = f'attachment; filename="CV_Richat_{consultant.prenom}_{consultant.nom}.pdf"'
            return response
            
        except Exception as e:
            logger.error(f"Erreur lors du t√©l√©chargement: {str(e)}")
            return Response({
                'success': False,
                'error': 'Erreur lors du t√©l√©chargement du fichier'
            }, status=500)
            
    except Exception as e:
        logger.error(f"Erreur g√©n√©rale lors du t√©l√©chargement du CV Richat: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)
        


def recalculate_all_expertise_levels():
    """
    Fonction utilitaire pour recalculer tous les niveaux d'expertise
    √Ä ex√©cuter apr√®s la migration du nouveau syst√®me
    """
    consultants = Consultant.objects.all()
    updated_count = 0
    
    for consultant in consultants:
        if consultant.update_expertise_level():
            updated_count += 1
    
    print(f"Expertise recalcul√©e pour {updated_count}/{consultants.count()} consultants")
    
    return updated_count


# Nouvelle fonction pour la modification des comp√©tences
@api_view(['POST'])
def add_consultant_competence(request, consultant_id):
    """
    Ajoute une comp√©tence √† un consultant - VERSION AM√âLIOR√âE
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        nom_competence = request.data.get('nom_competence')
        niveau = int(request.data.get('niveau', 1))

        if not nom_competence:
            return Response({"error": "Nom de comp√©tence requis"}, status=400)

        # V√©rifier si la comp√©tence existe d√©j√†
        if Competence.objects.filter(consultant=consultant, nom_competence__iexact=nom_competence).exists():
            return Response({"error": "Cette comp√©tence existe d√©j√† pour ce consultant"}, status=400)

        # Cr√©er la comp√©tence
        competence = Competence.objects.create(
            consultant=consultant,
            nom_competence=nom_competence,
            niveau=niveau
        )

        # NOUVEAU : Recalcul automatique de l'expertise
        consultant.update_expertise_level()
        
        # Obtenir les d√©tails de l'expertise pour feedback
        expertise_details = consultant.get_expertise_details()

        return Response({
            'success': True,
            'competence': CompetenceSerializer(competence).data,
            'consultant_expertise': {
                'niveau': consultant.expertise,
                'details': expertise_details
            }
        }, status=201)
        
    except Exception as e:
        return Response({"error": str(e)}, status=500)
    
# views.py - Correction rapide pour g√©rer les champs manquants

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
import logging

logger = logging.getLogger(__name__)

def safe_get_field(obj, field_name, default=None):
    """R√©cup√®re un champ de mani√®re s√©curis√©e"""
    try:
        return getattr(obj, field_name, default)
    except:
        return default

def safe_get_date_field(obj, field_name, fallback_field=None):
    """R√©cup√®re un champ date de mani√®re s√©curis√©e avec fallback"""
    try:
        value = getattr(obj, field_name, None)
        if value:
            return value.isoformat() if hasattr(value, 'isoformat') else str(value)
        
        # Essayer le champ de fallback
        if fallback_field:
            fallback_value = getattr(obj, fallback_field, None)
            if fallback_value:
                return fallback_value.isoformat() if hasattr(fallback_value, 'isoformat') else str(fallback_value)
        
        return None
    except Exception as e:
        logger.warning(f"Erreur lors de la r√©cup√©ration du champ {field_name}: {e}")
        return None

@api_view(['GET'])
@permission_classes([AllowAny])
def get_consultants_pending(request):
    """
    R√©cup√®re la liste des consultants en attente de validation
    VERSION S√âCURIS√âE avec gestion des champs manquants
    """
    try:
        from .models import Consultant
        
        # Utiliser only() pour sp√©cifier exactement les champs √† r√©cup√©rer
        consultants_pending = Consultant.objects.filter(
            is_validated=False
        ).select_related('user')
        
        consultants_data = []
        for consultant in consultants_pending:
            try:
                # Construction s√©curis√©e des donn√©es
                consultant_data = {
                    'id': consultant.id,
                    'nom': safe_get_field(consultant, 'nom', ''),
                    'prenom': safe_get_field(consultant, 'prenom', ''),
                    'firstName': safe_get_field(consultant, 'firstName') or safe_get_field(consultant, 'prenom', ''),
                    'lastName': safe_get_field(consultant, 'lastName') or safe_get_field(consultant, 'nom', ''),
                    'email': safe_get_field(consultant, 'email', ''),
                    'telephone': safe_get_field(consultant, 'telephone', ''),
                    'phone': safe_get_field(consultant, 'phone') or safe_get_field(consultant, 'telephone', ''),
                    'pays': safe_get_field(consultant, 'pays', ''),
                    'country': safe_get_field(consultant, 'country') or safe_get_field(consultant, 'pays', ''),
                    'ville': safe_get_field(consultant, 'ville', ''),
                    'city': safe_get_field(consultant, 'city') or safe_get_field(consultant, 'ville', ''),
                    
                    # Gestion s√©curis√©e des dates avec fallbacks
                    'date_debut_dispo': safe_get_date_field(consultant, 'date_debut_dispo'),
                    'date_fin_dispo': safe_get_date_field(consultant, 'date_fin_dispo'),
                    'startAvailability': safe_get_date_field(consultant, 'startAvailability', 'date_debut_dispo'),
                    'endAvailability': safe_get_date_field(consultant, 'endAvailability', 'date_fin_dispo'),
                    
                    # Autres champs avec valeurs par d√©faut
                    'domaine_principal': safe_get_field(consultant, 'domaine_principal', 'DIGITAL'),
                    'specialite': safe_get_field(consultant, 'specialite', ''),
                    'expertise': safe_get_field(consultant, 'expertise', 'D√©butant'),
                    'expertise_score': safe_get_field(consultant, 'expertise_score'),
                    'annees_experience': safe_get_field(consultant, 'annees_experience', 0),
                    'formation_niveau': safe_get_field(consultant, 'formation_niveau', 'BAC+3'),
                    'certifications_count': safe_get_field(consultant, 'certifications_count', 0),
                    'projets_realises': safe_get_field(consultant, 'projets_realises', 0),
                    'leadership_experience': safe_get_field(consultant, 'leadership_experience', False),
                    'international_experience': safe_get_field(consultant, 'international_experience', False),
                    'skills': safe_get_field(consultant, 'skills', ''),
                    'statut': safe_get_field(consultant, 'statut', 'En_attente'),
                    'is_validated': safe_get_field(consultant, 'is_validated', False),
                    
                    # Fichiers
                    'cv': consultant.cv.url if safe_get_field(consultant, 'cv') else None,
                    'cvFilename': safe_get_field(consultant, 'cvFilename') or (consultant.cv.name.split('/')[-1] if safe_get_field(consultant, 'cv') else None),
                    'standardizedCvFilename': safe_get_field(consultant, 'standardizedCvFilename'),
                    'photo': consultant.photo.url if safe_get_field(consultant, 'photo') else None,
                    'profileImage': safe_get_field(consultant, 'profileImage'),
                    
                    # M√©tadonn√©es
                    'created_at': safe_get_date_field(consultant, 'created_at'),
                    'updated_at': safe_get_date_field(consultant, 'updated_at'),
                }
                consultants_data.append(consultant_data)
                
            except Exception as e:
                logger.error(f"Erreur lors du traitement du consultant {consultant.id}: {e}")
                continue
        
        return JsonResponse({
            'success': True,
            'data': consultants_data,
            'count': len(consultants_data)
        })
        
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des consultants en attente: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': f"Erreur lors de la r√©cup√©ration des consultants en attente: {str(e)}"
        }, status=500)


@api_view(['GET'])
@permission_classes([AllowAny])
def get_all_consultants(request):
    """
    R√©cup√®re la liste de tous les consultants
    VERSION S√âCURIS√âE avec gestion des champs manquants
    """
    try:
        from .models import Consultant
        
        # Param√®tres de filtrage
        statut_filter = request.GET.get('statut', None)
        validated_filter = request.GET.get('validated', None)
        search = request.GET.get('search', '')
        
        # Query de base
        consultants_query = Consultant.objects.all().select_related('user')
        
        # Filtres
        if statut_filter:
            consultants_query = consultants_query.filter(statut=statut_filter)
            
        if validated_filter is not None:
            validated_bool = validated_filter.lower() == 'true'
            consultants_query = consultants_query.filter(is_validated=validated_bool)
            
        if search:
            from django.db.models import Q
            consultants_query = consultants_query.filter(
                Q(nom__icontains=search) |
                Q(prenom__icontains=search) |
                Q(email__icontains=search)
            )
        
        consultants_data = []
        for consultant in consultants_query:
            try:
                # Utiliser la m√™me logique s√©curis√©e que get_consultants_pending
                consultant_data = {
                    'id': consultant.id,
                    'nom': safe_get_field(consultant, 'nom', ''),
                    'prenom': safe_get_field(consultant, 'prenom', ''),
                    'firstName': safe_get_field(consultant, 'firstName') or safe_get_field(consultant, 'prenom', ''),
                    'lastName': safe_get_field(consultant, 'lastName') or safe_get_field(consultant, 'nom', ''),
                    'email': safe_get_field(consultant, 'email', ''),
                    'telephone': safe_get_field(consultant, 'telephone', ''),
                    'phone': safe_get_field(consultant, 'phone') or safe_get_field(consultant, 'telephone', ''),
                    'pays': safe_get_field(consultant, 'pays', ''),
                    'country': safe_get_field(consultant, 'country') or safe_get_field(consultant, 'pays', ''),
                    'ville': safe_get_field(consultant, 'ville', ''),
                    'city': safe_get_field(consultant, 'city') or safe_get_field(consultant, 'ville', ''),
                    
                    # Gestion s√©curis√©e des dates
                    'date_debut_dispo': safe_get_date_field(consultant, 'date_debut_dispo'),
                    'date_fin_dispo': safe_get_date_field(consultant, 'date_fin_dispo'),
                    'startAvailability': safe_get_date_field(consultant, 'startAvailability', 'date_debut_dispo'),
                    'endAvailability': safe_get_date_field(consultant, 'endAvailability', 'date_fin_dispo'),
                    
                    # Autres champs
                    'domaine_principal': safe_get_field(consultant, 'domaine_principal', 'DIGITAL'),
                    'specialite': safe_get_field(consultant, 'specialite', ''),
                    'expertise': safe_get_field(consultant, 'expertise', 'D√©butant'),
                    'expertise_score': safe_get_field(consultant, 'expertise_score'),
                    'statut': safe_get_field(consultant, 'statut', 'En_attente'),
                    'is_validated': safe_get_field(consultant, 'is_validated', False),
                    'skills': safe_get_field(consultant, 'skills', ''),
                    
                    # Fichiers
                    'cv': consultant.cv.url if safe_get_field(consultant, 'cv') else None,
                    'photo': consultant.photo.url if safe_get_field(consultant, 'photo') else None,
                    
                    # M√©tadonn√©es
                    'created_at': safe_get_date_field(consultant, 'created_at'),
                    'updated_at': safe_get_date_field(consultant, 'updated_at'),
                }
                consultants_data.append(consultant_data)
                
            except Exception as e:
                logger.error(f"Erreur lors du traitement du consultant {consultant.id}: {e}")
                continue
        
        return JsonResponse({
            'success': True,
            'data': consultants_data,
            'count': len(consultants_data)
        })
        
    except Exception as e:
        logger.error(f"Erreur lors de la r√©cup√©ration des consultants: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': f"Erreur lors de la r√©cup√©ration des consultants: {str(e)}"
        }, status=500)
        
def check_field_exists(table_name, field_name):
    """
    V√©rifie si un champ existe dans une table
    """
    try:
        with connection.cursor() as cursor:
            cursor.execute(f"SHOW COLUMNS FROM {table_name}")
            existing_fields = {row[0] for row in cursor.fetchall()}
            return field_name in existing_fields
    except Exception as e:
        logger.error(f"Erreur lors de la v√©rification du champ {field_name}: {e}")
        return False


def safe_get_field_value(obj, field_name, fallback_field=None, default=None):
    """
    R√©cup√®re la valeur d'un champ de mani√®re s√©curis√©e avec fallback
    """
    try:
        # V√©rifier si le champ principal existe
        if hasattr(obj, field_name):
            value = getattr(obj, field_name, None)
            if value is not None:
                return value
        
        # Essayer le champ de fallback
        if fallback_field and hasattr(obj, fallback_field):
            value = getattr(obj, fallback_field, None)
            if value is not None:
                return value
        
        return default
    except Exception as e:
        logger.warning(f"Erreur lors de la r√©cup√©ration du champ {field_name}: {e}")
        return default




    """
    Extrait les comp√©tences du CV d'un consultant en utilisant competences_data.py
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        
        # V√©rifier que le consultant a un CV
        if not consultant.cv or not consultant.cv.name:
            return Response({
                'success': False,
                'error': 'Aucun CV trouv√© pour ce consultant'
            }, status=400)
        
        # Chemin du fichier CV
        cv_path = consultant.cv.path
        
        if not os.path.exists(cv_path):
            return Response({
                'success': False,
                'error': 'Fichier CV introuvable sur le serveur'
            }, status=404)
        
        logger.info(f"Extraction des comp√©tences pour le consultant {consultant_id} depuis {cv_path}")
        
        # Utiliser la fonction d'extraction existante
        competences, primary_domain = extract_competences_from_cv(cv_path)
        
        if not competences:
            return Response({
                'success': False,
                'error': 'Aucune comp√©tence d√©tect√©e dans le CV',
                'skills': []
            }, status=200)
        
        # Mettre √† jour le domaine principal si d√©tect√©
        if primary_domain and primary_domain != consultant.domaine_principal:
            consultant.domaine_principal = primary_domain
            consultant.save(update_fields=['domaine_principal'])
            logger.info(f"Domaine principal mis √† jour : {primary_domain}")
        
        # Sauvegarder les nouvelles comp√©tences (√©viter les doublons)
        added_skills = []
        existing_skills = set(
            Competence.objects.filter(consultant=consultant)
            .values_list('nom_competence', flat=True)
        )
        
        for competence_name in competences:
            competence_clean = competence_name.strip().title()
            
            # √âviter les doublons (insensible √† la casse)
            if competence_clean.lower() not in [skill.lower() for skill in existing_skills]:
                Competence.objects.create(
                    consultant=consultant,
                    nom_competence=competence_clean,
                    niveau=2  # Niveau par d√©faut pour les comp√©tences extraites
                )
                added_skills.append(competence_clean)
                existing_skills.add(competence_clean)
        
        # Mettre √† jour le niveau d'expertise du consultant
        total_competences = Competence.objects.filter(consultant=consultant).count()
        
        # Ancien syst√®me de calcul (√† remplacer par le nouveau syst√®me d'expertise)
        if total_competences >= 10:
            consultant.expertise = "Expert"
        elif total_competences >= 5:
            consultant.expertise = "Interm√©diaire"
        else:
            consultant.expertise = "D√©butant"
        
        consultant.save(update_fields=['expertise'])
        
        # Mettre √† jour le champ skills pour l'affichage
        all_skills = list(existing_skills)
        consultant.skills = ', '.join(all_skills) if hasattr(consultant, 'skills') else None
        if consultant.skills is not None:
            consultant.save(update_fields=['skills'])
        
        logger.info(f"Extraction termin√©e : {len(added_skills)} nouvelles comp√©tences ajout√©es")
        
        return Response({
            'success': True,
            'message': f'{len(added_skills)} nouvelles comp√©tences extraites',
            'skills': all_skills,
            'new_skills': added_skills,
            'total_skills': len(all_skills),
            'primary_domain': primary_domain,
            'expertise_level': consultant.expertise
        })
        
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction des comp√©tences : {str(e)}")
        return Response({
            'success': False,
            'error': f'Erreur lors de l\'extraction : {str(e)}'
        }, status=500)
        
# views.py - CORRECTIONS pour l'extraction de comp√©tences

import logging
import os
import re
from decimal import Decimal
from datetime import datetime

from django.shortcuts import get_object_or_404
from django.http import JsonResponse
from rest_framework.decorators import api_view
from rest_framework.response import Response

from .models import Consultant, Competence
from .competences_data import ALL_SKILLS

logger = logging.getLogger(__name__)

def extract_competences(text, domain_principale=None):
    """
    Version finale am√©lior√©e de l'extraction de comp√©tences avec filtrage intelligent
    """
    if not text or len(text.strip()) < 10:
        logger.warning("Texte vide ou trop court pour l'extraction")
        return [], 'DIGITAL'
    
    # Nettoyer le texte
    text_clean = text.lower().strip()
    competences_trouvees = set()
    domain_scores = {}
    
    logger.info(f"D√©but extraction sur {len(text)} caract√®res de texte")
    
    # Liste √©largie des mots √† ignorer (faux positifs fr√©quents)
    mots_ignores = {
        # Mots trop courts ou ambigus
        'ar', 'or', 'ai', 'go', 'r', 'c', 's3', 'ux', 'ui', 'db', 'os', 'it', 'is', 'in', 'on', 
        'at', 'to', 'of', 'as', 'by', 'an', 'be', 'we', 'he', 'me', 'my', 'so', 'up', 'if',
        # Acronymes Mauritaniens ou contextuels non-techniques
        'ets', 'eu ets', 'ges', 'gsm', 'cad', 'bp mauritanie', 'or mauritanie', 'bp', 'mr',
        # Mots fran√ßais courants
        'et', 'ou', 'le', 'la', 'les', 'un', 'une', 'des', 'du', 'de', 'dans', 'sur', 'avec',
        'pour', 'par', 'sans', 'sous', 'vers', 'chez', 'entre', 'depuis', 'pendant',
        # Pr√©positions et conjonctions anglaises
        'and', 'but', 'for', 'nor', 'yet', 'the', 'a', 'an', 'this', 'that', 'these', 'those'
    }
    
    # Comp√©tences techniques valides courtes (exceptions √† garder)
    competences_courtes_valides = {
        'sql', 'php', 'css', 'xml', 'api', 'aws', 'gcp', 'iot', 'erp', 'crm', 'seo', 'sap',
        'html', 'java', 'git', 'npm', 'vue', 'tdd', 'bdd', 'mvp', 'agile', 'scrum', 'ci/cd',
        'rest', 'soap', 'json', 'ajax', 'cors', 'jwt', 'ssl', 'tls', 'ssh', 'ftp', 'http',
        'tcp', 'udp', 'dns', 'dhcp', 'vpn', 'lan', 'wan', 'cdn', 'orm', 'mvc', 'spa'
    }
    
    # Parcourir chaque domaine et ses comp√©tences
    for domain, skills_list in ALL_SKILLS.items():
        domain_score = 0
        domain_competences = []
        
        for skill in skills_list:
            skill_lower = skill.lower().strip()
            
            # Ignorer les mots de la liste noire
            if skill_lower in mots_ignores:
                continue
            
            # Pour les comp√©tences courtes, v√©rifier si c'est une exception valide
            if len(skill_lower) <= 3:
                if skill_lower not in competences_courtes_valides:
                    continue
            
            # Recherche exacte avec v√©rification de contexte
            if skill_lower in text_clean:
                if is_valid_technical_context(text_clean, skill_lower):
                    competences_trouvees.add(skill)
                    domain_competences.append(skill)
                    domain_score += 1
                    logger.debug(f"Comp√©tence trouv√©e (exacte): {skill}")
                    continue
            
            # Recherche avec variations pour les comp√©tences de plus de 3 caract√®res
            if len(skill_lower) > 3:
                skill_variants = [
                    skill_lower.replace(' ', ''),
                    skill_lower.replace(' ', '-'),
                    skill_lower.replace(' ', '.'),
                    skill_lower.replace('-', ' '),
                    skill_lower.replace('.', ' '),
                    skill_lower.replace('_', ' '),
                    skill_lower.replace('/', ' ')
                ]
                
                for variant in skill_variants:
                    if variant in text_clean and is_valid_technical_context(text_clean, variant):
                        competences_trouvees.add(skill)
                        domain_competences.append(skill)
                        domain_score += 0.8
                        logger.debug(f"Comp√©tence trouv√©e (variante): {skill}")
                        break
        
        domain_scores[domain] = {
            'score': domain_score,
            'competences': domain_competences,
            'pourcentage': (domain_score / len(skills_list)) * 100 if skills_list else 0
        }
        
        logger.info(f"Domaine {domain}: {domain_score} comp√©tences trouv√©es ({len(domain_competences)} uniques)")
    
    # D√©terminer le domaine principal
    if domain_scores:
        domain_principal = max(domain_scores.items(), key=lambda x: x[1]['score'])[0]
        best_score = domain_scores[domain_principal]['score']
        
        if best_score == 0:
            domain_principal = domain_principale or 'DIGITAL'
        
        logger.info(f"Domaine principal d√©tect√©: {domain_principal} avec {best_score} comp√©tences")
    else:
        domain_principal = domain_principale or 'DIGITAL'
    
    # Filtrage final des comp√©tences
    competences_filtrees = list(competences_trouvees)
    
    logger.info(f"Total comp√©tences extraites (apr√®s filtrage): {len(competences_filtrees)}")
    
    return competences_filtrees, domain_principal


def is_valid_technical_context(text, skill):
    """
    V√©rifie si une comp√©tence est dans un contexte technique valide
    Version am√©lior√©e avec plus de mots-cl√©s de contexte
    """
    # Mots-cl√©s qui indiquent un contexte technique valide
    context_keywords = [
        # D√©veloppement et programmation
        'd√©veloppement', 'programmation', 'langage', 'framework', 'technologie',
        'comp√©tence', 'skill', 'experience', 'projet', 'utilisation', 'ma√Ætrise',
        'connaissance', 'expertise', 'formation', 'certification', 'outils',
        'logiciel', 'software', 'development', 'programming', 'language',
        'stack', 'tech', 'tools', 'environment', 'platform', 'api', 'database',
        
        # Contexte professionnel
        'projet', 'mission', 'travail', 'emploi', 'poste', 'responsabilit√©',
        'r√©alisation', 'conception', 'd√©velopp√©', 'impl√©ment√©', 'utilis√©',
        'g√©r√©', 'administr√©', 'configur√©', 'd√©ploy√©', 'maintenu',
        
        # Formation et certifications
        'formation', 'cours', 'apprentissage', 'certification', 'dipl√¥me',
        '√©tude', 'universit√©', '√©cole', 'institut', 'acad√©mie', 'bootcamp',
        
        # Contexte technique sp√©cifique
        'web', 'mobile', 'cloud', 'data', 'analyse', 'syst√®me', 'r√©seau',
        's√©curit√©', 'base', 'donn√©es', 'serveur', 'client', 'backend', 'frontend'
    ]
    
    # Chercher le skill dans le texte avec un contexte de 100 caract√®res avant et apr√®s
    import re
    pattern = rf'.{{0,100}}{re.escape(skill)}.{{0,100}}'
    matches = re.finditer(pattern, text, re.IGNORECASE)
    
    for match in matches:
        context = match.group().lower()
        # Si on trouve un mot-cl√© de contexte technique, c'est valide
        if any(keyword in context for keyword in context_keywords):
            return True
    
    # Pour les comp√©tences tr√®s sp√©cifiques (plus de 5 caract√®res), 
    # on peut √™tre moins strict sur le contexte
    if len(skill) > 5:
        return True
    
    # Pour les comp√©tences courtes mais tr√®s techniques
    technical_skills = {
        'sql', 'php', 'css', 'html', 'java', 'python', 'react', 'angular', 'vue',
        'git', 'docker', 'aws', 'azure', 'api', 'rest', 'json', 'xml'
    }
    
    if skill in technical_skills:
        return True
    
    return False
    """
    Version corrig√©e et intelligente de l'extraction de comp√©tences
    """
    if not text or len(text.strip()) < 10:
        logger.warning("Texte vide ou trop court pour l'extraction")
        return [], 'DIGITAL'
    
    # Nettoyer le texte
    text_clean = text.lower().strip()
    competences_trouvees = set()
    domain_scores = {}
    
    logger.info(f"D√©but extraction sur {len(text)} caract√®res de texte")
    
    # Parcourir chaque domaine et ses comp√©tences
    for domain, skills_list in ALL_SKILLS.items():
        domain_score = 0
        domain_competences = []
        
        for skill in skills_list:
            skill_lower = skill.lower().strip()
            
            # Recherche exacte
            if skill_lower in text_clean:
                competences_trouvees.add(skill)
                domain_competences.append(skill)
                domain_score += 1
                logger.debug(f"Comp√©tence trouv√©e (exacte): {skill}")
                continue
            
            # Recherche avec variations (espaces, tirets, points)
            skill_variants = [
                skill_lower.replace(' ', ''),
                skill_lower.replace(' ', '-'),
                skill_lower.replace(' ', '.'),
                skill_lower.replace('-', ' '),
                skill_lower.replace('.', ' '),
                skill_lower.replace('_', ' '),
                skill_lower.replace('/', ' ')
            ]
            
            if any(variant in text_clean for variant in skill_variants):
                competences_trouvees.add(skill)
                domain_competences.append(skill)
                domain_score += 0.8
                logger.debug(f"Comp√©tence trouv√©e (variante): {skill}")
                continue
            
            # Recherche par mots-cl√©s partiels
            skill_words = skill_lower.split()
            if len(skill_words) > 1:
                # Si tous les mots de la comp√©tence sont dans le texte
                if all(word in text_clean for word in skill_words if len(word) > 2):
                    competences_trouvees.add(skill)
                    domain_competences.append(skill)
                    domain_score += 0.6
                    logger.debug(f"Comp√©tence trouv√©e (mots-cl√©s): {skill}")
        
        domain_scores[domain] = {
            'score': domain_score,
            'competences': domain_competences,
            'pourcentage': (domain_score / len(skills_list)) * 100 if skills_list else 0
        }
        
        logger.info(f"Domaine {domain}: {domain_score} comp√©tences trouv√©es ({len(domain_competences)} uniques)")
    
    # D√©terminer le domaine principal
    if domain_scores:
        domain_principal = max(domain_scores.items(), key=lambda x: x[1]['score'])[0]
        best_score = domain_scores[domain_principal]['score']
        
        if best_score == 0:
            domain_principal = domain_principale or 'DIGITAL'
        
        logger.info(f"Domaine principal d√©tect√©: {domain_principal} avec {best_score} comp√©tences")
    else:
        domain_principal = domain_principale or 'DIGITAL'
    
    competences_list = list(competences_trouvees)
    logger.info(f"Total comp√©tences extraites: {len(competences_list)}")
    
    return competences_list, domain_principal
    """
    Version corrig√©e de l'extraction de comp√©tences
    """
    if not text or len(text.strip()) < 10:
        logger.warning("Texte vide ou trop court pour l'extraction")
        return [], 'DIGITAL'
    
    # Nettoyer le texte
    text_clean = text.lower().strip()
    competences_trouvees = set()
    domain_scores = {}
    
    logger.info(f"D√©but extraction sur {len(text)} caract√®res de texte")
    
    # Parcourir chaque domaine et ses comp√©tences
    for domain, skills_list in ALL_SKILLS.items():
        domain_score = 0
        domain_competences = []
        
        for skill in skills_list:
            skill_lower = skill.lower().strip()
            
            # Recherche exacte
            if skill_lower in text_clean:
                competences_trouvees.add(skill)
                domain_competences.append(skill)
                domain_score += 1
                logger.debug(f"Comp√©tence trouv√©e (exacte): {skill}")
            
            # Recherche avec variations (espaces, tirets, points)
            elif any(variant in text_clean for variant in [
                skill_lower.replace(' ', ''),
                skill_lower.replace(' ', '-'),
                skill_lower.replace(' ', '.'),
                skill_lower.replace('-', ' '),
                skill_lower.replace('.', ' ')
            ]):
                competences_trouvees.add(skill)
                domain_competences.append(skill)
                domain_score += 0.8
                logger.debug(f"Comp√©tence trouv√©e (variante): {skill}")
        
        domain_scores[domain] = {
            'score': domain_score,
            'competences': domain_competences,
            'pourcentage': (domain_score / len(skills_list)) * 100 if skills_list else 0
        }
        
        logger.info(f"Domaine {domain}: {domain_score} comp√©tences trouv√©es ({len(domain_competences)} uniques)")
    
    # D√©terminer le domaine principal
    if domain_scores:
        domain_principal = max(domain_scores.items(), key=lambda x: x[1]['score'])[0]
        best_score = domain_scores[domain_principal]['score']
        
        if best_score == 0:
            domain_principal = domain_principale or 'DIGITAL'
        
        logger.info(f"Domaine principal d√©tect√©: {domain_principal} avec {best_score} comp√©tences")
    else:
        domain_principal = domain_principale or 'DIGITAL'
    
    competences_list = list(competences_trouvees)
    logger.info(f"Total comp√©tences extraites: {len(competences_list)}")
    
    return competences_list, domain_principal

def extract_competences_from_cv(file_path):
    """
    Version corrig√©e de l'extraction depuis un fichier CV avec gestion robuste
    """
    logger.info(f"Extraction depuis le fichier: {file_path}")
    
    if not os.path.exists(file_path):
        logger.error(f"Fichier introuvable: {file_path}")
        return [], 'DIGITAL'
    
    try:
        # V√©rifier l'extension du fichier
        file_ext = os.path.splitext(file_path)[1].lower()
        text = ""
        
        if file_ext == '.pdf':
            # Extraction PDF avec PyMuPDF
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    page_text = page.get_text("text")
                    text += page_text + "\n"
                
                doc.close()
                logger.info(f"Texte extrait du PDF: {len(text)} caract√®res")
                
                # Si le PDF n'a pas de texte, essayer OCR
                if len(text.strip()) < 50:
                    logger.info("PDF sans texte d√©tect√©, tentative OCR...")
                    try:
                        import pytesseract
                        from PIL import Image, ImageEnhance
                        
                        doc = fitz.open(file_path)
                        text = ""
                        for page_num in range(min(3, len(doc))):  # Max 3 pages pour OCR
                            page = doc[page_num]
                            # Am√©liorer la qualit√© pour OCR
                            pix = page.get_pixmap(dpi=300)
                            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                            
                            # Am√©lioration d'image pour OCR
                            img = img.convert("L")  # Convertir en niveaux de gris
                            img = ImageEnhance.Contrast(img).enhance(2.0)
                            img = ImageEnhance.Sharpness(img).enhance(1.5)
                            
                            # OCR multilingue
                            page_text = pytesseract.image_to_string(img, lang='fra+eng')
                            text += page_text + "\n"
                        
                        doc.close()
                        logger.info(f"Texte extrait par OCR: {len(text)} caract√®res")
                        
                    except Exception as ocr_error:
                        logger.error(f"Erreur OCR: {ocr_error}")
                        # Essayer une extraction basique
                        text = "PDF sans texte d√©tect√©"
                
            except Exception as pdf_error:
                logger.error(f"Erreur PyMuPDF: {pdf_error}")
                return [], 'DIGITAL'
        
        elif file_ext in ['.doc', '.docx']:
            # Extraction Word
            try:
                if file_ext == '.docx':
                    import docx
                    doc = docx.Document(file_path)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                else:
                    # Pour .doc, utiliser python-docx2txt
                    try:
                        import docx2txt
                        text = docx2txt.process(file_path)
                    except ImportError:
                        logger.error("Module docx2txt non disponible pour fichiers .doc")
                        return [], 'DIGITAL'
                
                logger.info(f"Texte extrait du document Word: {len(text)} caract√®res")
                
            except Exception as word_error:
                logger.error(f"Erreur extraction Word: {word_error}")
                return [], 'DIGITAL'
        
        else:
            logger.error(f"Format de fichier non support√©: {file_ext}")
            return [], 'DIGITAL'
        
        # V√©rifier que nous avons du contenu
        if not text or len(text.strip()) < 20:
            logger.warning(f"Texte extrait insuffisant: {len(text)} caract√®res")
            # Essayer une extraction de secours bas√©e sur le nom du fichier
            filename_skills = extract_skills_from_filename(file_path)
            if filename_skills:
                return filename_skills, 'DIGITAL'
            return [], 'DIGITAL'
        
        # Extraire les comp√©tences du texte
        competences, domain = extract_competences(text)
        
        logger.info(f"Extraction termin√©e: {len(competences)} comp√©tences trouv√©es")
        return competences, domain
        
    except Exception as e:
        logger.error(f"Erreur g√©n√©rale d'extraction: {str(e)}")
        return [], 'DIGITAL'
    """
    Version corrig√©e de l'extraction depuis un fichier CV
    """
    logger.info(f"Extraction depuis le fichier: {file_path}")
    
    if not os.path.exists(file_path):
        logger.error(f"Fichier introuvable: {file_path}")
        return [], 'DIGITAL'
    
    try:
        # V√©rifier l'extension du fichier
        file_ext = os.path.splitext(file_path)[1].lower()
        text = ""
        
        if file_ext == '.pdf':
            # Extraction PDF avec PyMuPDF
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    page_text = page.get_text("text")
                    text += page_text + "\n"
                
                doc.close()
                logger.info(f"Texte extrait du PDF: {len(text)} caract√®res")
                
            except Exception as pdf_error:
                logger.error(f"Erreur PyMuPDF: {pdf_error}")
                
                # Fallback avec OCR si pas de texte
                if len(text.strip()) < 50:
                    logger.info("Tentative OCR avec Tesseract...")
                    try:
                        import pytesseract
                        from PIL import Image
                        
                        doc = fitz.open(file_path)
                        for page_num in range(min(3, len(doc))):  # Max 3 pages pour OCR
                            page = doc[page_num]
                            pix = page.get_pixmap(dpi=200)
                            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                            page_text = pytesseract.image_to_string(img, lang='fra+eng')
                            text += page_text + "\n"
                        
                        doc.close()
                        logger.info(f"Texte extrait par OCR: {len(text)} caract√®res")
                        
                    except Exception as ocr_error:
                        logger.error(f"Erreur OCR: {ocr_error}")
                        return [], 'DIGITAL'
        
        elif file_ext in ['.doc', '.docx']:
            # Extraction Word
            try:
                if file_ext == '.docx':
                    import docx
                    doc = docx.Document(file_path)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                else:
                    # Pour .doc, utiliser python-docx2txt ou autre
                    try:
                        import docx2txt
                        text = docx2txt.process(file_path)
                    except ImportError:
                        logger.error("Module docx2txt non disponible pour fichiers .doc")
                        return [], 'DIGITAL'
                
                logger.info(f"Texte extrait du document Word: {len(text)} caract√®res")
                
            except Exception as word_error:
                logger.error(f"Erreur extraction Word: {word_error}")
                return [], 'DIGITAL'
        
        else:
            logger.error(f"Format de fichier non support√©: {file_ext}")
            return [], 'DIGITAL'
        
        # V√©rifier que nous avons du contenu
        if not text or len(text.strip()) < 20:
            logger.warning(f"Texte extrait insuffisant: {len(text)} caract√®res")
            return [], 'DIGITAL'
        
        # Extraire les comp√©tences
        competences, domain = extract_competences(text)
        
        logger.info(f"Extraction termin√©e: {len(competences)} comp√©tences trouv√©es")
        return competences, domain
        
    except Exception as e:
        logger.error(f"Erreur g√©n√©rale d'extraction: {str(e)}")
        return [], 'DIGITAL'
def extract_skills_from_filename(file_path):
    """
    Extraction de secours bas√©e sur le nom du fichier
    """
    filename = os.path.basename(file_path).lower()
    found_skills = []
    
    # Comp√©tences communes √† rechercher dans le nom du fichier
    common_skills = [
        'python', 'java', 'javascript', 'php', 'sql', 'html', 'css',
        'react', 'angular', 'vue', 'node', 'django', 'flask', 'spring',
        'mysql', 'postgresql', 'mongodb', 'git', 'docker', 'aws', 'azure'
    ]
    
    for skill in common_skills:
        if skill in filename:
            found_skills.append(skill.title())
    
    return found_skills
@api_view(['POST'])
def extract_consultant_competences(request, consultant_id):
    """
    VERSION CORRIG√âE FINALE de l'endpoint d'extraction de comp√©tences
    Supprime la duplication et utilise l'intelligence des competences_data.py
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        
        logger.info(f"=== D√âBUT EXTRACTION COMP√âTENCES ===")
        logger.info(f"Consultant ID: {consultant_id}")
        logger.info(f"Consultant: {consultant.nom} {consultant.prenom}")
        
        # V√©rifier que le consultant a un CV
        if not consultant.cv or not consultant.cv.name:
            logger.error("Aucun CV trouv√©")
            return Response({
                'success': False,
                'error': 'Aucun CV trouv√© pour ce consultant. Veuillez d\'abord t√©l√©charger votre CV.',
                'skills': []
            }, status=400)
        
        # Chemin du fichier CV
        cv_path = consultant.cv.path
        
        if not os.path.exists(cv_path):
            logger.error(f"Fichier introuvable: {cv_path}")
            return Response({
                'success': False,
                'error': 'Fichier CV introuvable sur le serveur. Veuillez t√©l√©charger √† nouveau votre CV.',
                'skills': []
            }, status=404)
        
        logger.info(f"Fichier CV trouv√©: {cv_path}")
        
        # Utiliser la fonction d'extraction corrig√©e
        competences, primary_domain = extract_competences_from_cv(cv_path)
        
        logger.info(f"Extraction termin√©e: {len(competences)} comp√©tences trouv√©es")
        
        if not competences:
            # Essayer une extraction plus permissive avec recherche directe
            logger.info("Premi√®re extraction vide, tentative avec m√©thode alternative...")
            
            try:
                # Lire le fichier directement et chercher des patterns
                with open(cv_path, 'rb') as f:
                    content = f.read()
                
                # Conversion en texte pour recherche
                text_content = str(content, errors='ignore').lower()
                
                # Recherche directe dans competences_data.py
                competences_alternatives = []
                
                # Utiliser ALL_SKILLS pour une recherche exhaustive
                for domain, skills_list in ALL_SKILLS.items():
                    for skill in skills_list:
                        skill_lower = skill.lower()
                        if skill_lower in text_content:
                            competences_alternatives.append(skill)
                
                # Supprimer les doublons
                competences_alternatives = list(set(competences_alternatives))
                
                if competences_alternatives:
                    competences = competences_alternatives
                    logger.info(f"Extraction alternative r√©ussie: {len(competences)} comp√©tences")
                else:
                    logger.warning("Aucune comp√©tence d√©tect√©e m√™me avec m√©thode alternative")
                    return Response({
                        'success': False,
                        'error': 'Aucune comp√©tence d√©tect√©e dans votre CV. V√©rifiez que votre CV contient bien vos comp√©tences techniques.',
                        'skills': [],
                        'debug_info': {
                            'file_size': len(content),
                            'file_type': cv_path.split('.')[-1],
                            'content_preview': text_content[:200] + "..." if len(text_content) > 200 else text_content
                        }
                    }, status=200)
                    
            except Exception as fallback_error:
                logger.error(f"Erreur dans l'extraction alternative: {fallback_error}")
                return Response({
                    'success': False,
                    'error': 'Erreur lors de l\'analyse du CV. V√©rifiez que le fichier n\'est pas corrompu.',
                    'skills': []
                }, status=500)
        
        # Mettre √† jour le domaine principal si d√©tect√©
        if primary_domain and primary_domain != consultant.domaine_principal:
            consultant.domaine_principal = primary_domain
            consultant.save(update_fields=['domaine_principal'])
            logger.info(f"Domaine principal mis √† jour : {primary_domain}")
        
        # Sauvegarder les nouvelles comp√©tences (√©viter les doublons)
        added_skills = []
        existing_skills = set(
            Competence.objects.filter(consultant=consultant)
            .values_list('nom_competence', flat=True)
        )
        
        logger.info(f"Comp√©tences existantes: {len(existing_skills)}")
        
        for competence_name in competences:
            competence_clean = competence_name.strip().title()
            
            # √âviter les doublons (insensible √† la casse)
            if competence_clean.lower() not in [skill.lower() for skill in existing_skills]:
                try:
                    Competence.objects.create(
                        consultant=consultant,
                        nom_competence=competence_clean,
                        niveau=2  # Niveau par d√©faut pour les comp√©tences extraites
                    )
                    added_skills.append(competence_clean)
                    existing_skills.add(competence_clean)
                    logger.debug(f"Comp√©tence ajout√©e: {competence_clean}")
                except Exception as comp_error:
                    logger.error(f"Erreur lors de l'ajout de la comp√©tence {competence_clean}: {comp_error}")
        
        # Mettre √† jour le niveau d'expertise du consultant
        total_competences = Competence.objects.filter(consultant=consultant).count()
        
        # Calcul de l'expertise bas√© sur le nombre de comp√©tences
        if total_competences >= 15:
            consultant.expertise = "Expert"
        elif total_competences >= 8:
            consultant.expertise = "Interm√©diaire"
        elif total_competences >= 3:
            consultant.expertise = "Junior"
        else:
            consultant.expertise = "D√©butant"
        
        consultant.save(update_fields=['expertise'])
        
        # Mettre √† jour le champ skills pour l'affichage dans ConsultantWelcome
        all_skills = list(existing_skills)
        if hasattr(consultant, 'skills'):
            consultant.skills = ', '.join(all_skills)
            consultant.save(update_fields=['skills'])
        
        logger.info(f"‚úì Extraction termin√©e avec succ√®s")
        logger.info(f"‚úì {len(added_skills)} nouvelles comp√©tences ajout√©es")
        logger.info(f"‚úì Total comp√©tences: {len(all_skills)}")
        logger.info(f"‚úì Niveau expertise: {consultant.expertise}")
        
        # Message de succ√®s d√©taill√©
        message = f"{len(added_skills)} nouvelles comp√©tences extraites avec succ√®s" if added_skills else "Aucune nouvelle comp√©tence trouv√©e (toutes d√©j√† pr√©sentes)"
        
        return Response({
            'success': True,
            'message': message,
            'skills': all_skills,  # Toutes les comp√©tences du consultant
            'new_skills': added_skills,  # Seulement les nouvelles
            'total_skills': len(all_skills),
            'primary_domain': primary_domain,
            'expertise_level': consultant.expertise,
            'extraction_details': {
                'extracted_count': len(competences),
                'new_count': len(added_skills),
                'existing_count': len(existing_skills) - len(added_skills),
                'domain_detected': primary_domain
            }
        })
        
    except Exception as e:
        logger.error(f"‚ùå Erreur lors de l'extraction des comp√©tences : {str(e)}")
        return Response({
            'success': False,
            'error': f'Erreur lors de l\'extraction : {str(e)}',
            'skills': []
        }, status=500)

# Endpoint manquant pour l'analyse d'expertise
@api_view(['GET'])
def get_consultant_expertise_analysis(request, consultant_id):
    """
    Analyse d√©taill√©e de l'expertise d'un consultant - VERSION FINALE
    Utilise les m√©thodes du mod√®le Consultant
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        
        logger.info(f"Analyse expertise pour consultant {consultant_id}: {consultant.nom} {consultant.prenom}")
        
        # Utiliser la m√©thode du mod√®le pour obtenir les d√©tails d'expertise
        expertise_details = consultant.get_expertise_details()
        
        logger.info(f"Expertise calcul√©e: {expertise_details['niveau_calcule']} (Score: {expertise_details['score_total']})")
        
        return Response({
            'success': True,
            'consultant': {
                'id': consultant.id,
                'nom': f"{consultant.prenom} {consultant.nom}",
                'expertise_actuelle': consultant.expertise,
                'analyse': expertise_details
            }
        })
        
    except Exception as e:
        logger.error(f"Erreur analyse expertise pour consultant {consultant_id}: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)


@api_view(['POST'])
def update_consultant_expertise_info(request, consultant_id):
    """
    Met √† jour les informations d'expertise d'un consultant - VERSION FINALE
    """
    try:
        consultant = get_object_or_404(Consultant, id=consultant_id)
        
        logger.info(f"Mise √† jour expertise pour consultant {consultant_id}: {consultant.nom} {consultant.prenom}")
        logger.info(f"Donn√©es re√ßues: {request.data}")
        
        # R√©cup√©rer les donn√©es du formulaire
        data = request.data
        
        # Mise √† jour des champs d'expertise
        updated_fields = []
        
        if 'annees_experience' in data:
            consultant.annees_experience = int(data['annees_experience'])
            updated_fields.append('annees_experience')
            logger.info(f"Ann√©es d'exp√©rience: {consultant.annees_experience}")
        
        if 'formation_niveau' in data:
            consultant.formation_niveau = data['formation_niveau']
            updated_fields.append('formation_niveau')
            logger.info(f"Niveau de formation: {consultant.formation_niveau}")
        
        if 'certifications_count' in data:
            consultant.certifications_count = int(data['certifications_count'])
            updated_fields.append('certifications_count')
            logger.info(f"Nombre de certifications: {consultant.certifications_count}")
        
        if 'projets_realises' in data:
            consultant.projets_realises = int(data['projets_realises'])
            updated_fields.append('projets_realises')
            logger.info(f"Projets r√©alis√©s: {consultant.projets_realises}")
        
        if 'leadership_experience' in data:
            consultant.leadership_experience = data['leadership_experience'] in [True, 'true', 'True']
            updated_fields.append('leadership_experience')
            logger.info(f"Exp√©rience leadership: {consultant.leadership_experience}")
        
        if 'international_experience' in data:
            consultant.international_experience = data['international_experience'] in [True, 'true', 'True']
            updated_fields.append('international_experience')
            logger.info(f"Exp√©rience internationale: {consultant.international_experience}")
        
        # Sauvegarder les modifications
        if updated_fields:
            consultant.save(update_fields=updated_fields)
            logger.info(f"Champs mis √† jour: {updated_fields}")
        
        # Recalculer automatiquement l'expertise (la m√©thode save() du mod√®le le fait d√©j√†)
        consultant.update_expertise_level()
        
        # Obtenir les d√©tails mis √† jour
        expertise_details = consultant.get_expertise_details()
        
        logger.info(f"Nouveau niveau d'expertise: {consultant.expertise}")
        logger.info(f"Nouveau score: {consultant.expertise_score}")
        
        return Response({
            'success': True,
            'message': 'Niveau d\'expertise mis √† jour avec succ√®s',
            'consultant': {
                'id': consultant.id,
                'nom': f"{consultant.prenom} {consultant.nom}",
                'expertise_level': consultant.expertise,
                'expertise_score': consultant.expertise_score,
                'expertise_details': expertise_details
            }
        })
        
    except Exception as e:
        logger.error(f"Erreur mise √† jour expertise pour consultant {consultant_id}: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)